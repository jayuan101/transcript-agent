#!/usr/bin/env python3
"""
Transcript Agent
Formats:  .mp3 .wav .m4a .ogg .aac (audio)
          .mp4 .mov .avi .mkv .webm (video)
          .srt .vtt (subtitles)
          .pdf .docx .txt .md (documents)
Outputs:  clean transcript, speaker-labelled dialogue, summary, key points,
          speech analytics (WPM + accent detection)
"""

import os
import sys

# ── Force UTF-8 everywhere so no Unicode character can crash transcription ──────
# Set env var BEFORE any library imports so subprocesses inherit it too
os.environ.setdefault("PYTHONIOENCODING", "utf-8:replace")
os.environ.setdefault("PYTHONUTF8", "1")

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

def _safe_print(*args, **kwargs):
    """print() that never raises on encoding errors."""
    try:
        print(*args, **kwargs)
    except (UnicodeEncodeError, UnicodeDecodeError):
        safe = " ".join(str(a).encode("ascii", errors="replace").decode("ascii") for a in args)
        try:
            print(safe, **{k: v for k, v in kwargs.items() if k != "file"})
        except Exception:
            pass
import json
import re
import tempfile
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
# ── LLM client abstraction ───────────────────────────────────────────────────

class LLMClient:
    """Thin wrapper normalising Anthropic and OpenAI-compatible provider SDKs."""

    def __init__(self, provider: str, api_key: str, model: str, base_url: str = None,
                 use_gpu: bool = True):
        self.provider = provider  # "anthropic" | "openai" | "openai_compat"
        self.model = model
        self.use_gpu = use_gpu
        # Detect Ollama by base_url so we can pass num_gpu options
        self._is_ollama = bool(base_url and ("11434" in base_url or "ollama" in base_url.lower()))
        if provider == "anthropic":
            import anthropic as _ant
            self._client = _ant.Anthropic(api_key=api_key) if api_key else _ant.Anthropic()
        else:
            import openai as _oai
            # Always pass api_key — OpenAI SDK v2.x requires it even with a
            # custom base_url. Use "none" for local models (Ollama) that ignore it.
            kw = {"api_key": api_key if api_key else "none"}
            if base_url:
                kw["base_url"] = base_url
            self._client = _oai.OpenAI(**kw)

    def chat(self, system: str, user: str, max_tokens: int,
             thinking: bool = False, on_usage=None) -> str:
        """on_usage(input_tokens, output_tokens) — called after each API response."""
        if self.provider == "anthropic":
            import time as _time
            kw = {}
            if thinking:
                kw["thinking"] = {"type": "adaptive"}
            _backoff = [5, 15, 30, 60]   # seconds between retries on overloaded
            _max_attempts = 5
            for _attempt in range(_max_attempts):
                try:
                    with self._client.messages.stream(
                        model=self.model,
                        max_tokens=max_tokens,
                        system=system,
                        messages=[{"role": "user", "content": user}],
                        **kw,
                    ) as stream:
                        resp = stream.get_final_message()
                    if on_usage and hasattr(resp, "usage"):
                        on_usage(resp.usage.input_tokens, resp.usage.output_tokens)
                    result = next((b.text for b in resp.content if b.type == "text"), "")
                    if result:
                        return result
                    if _attempt == 1:
                        kw.pop("thinking", None)
                except Exception as _e:
                    _err_str = str(_e).lower()
                    _is_overloaded = (
                        "overloaded" in _err_str
                        or "529" in _err_str
                        or getattr(_e, "status_code", None) == 529
                    )
                    _is_rate_limit = (
                        "rate" in _err_str or "429" in _err_str
                        or getattr(_e, "status_code", None) == 429
                    )
                    if _attempt >= _max_attempts - 1:
                        if _is_overloaded:
                            raise RuntimeError(
                                "Claude is temporarily at capacity. "
                                "Please wait a moment and try again."
                            ) from _e
                        raise
                    wait = _backoff[min(_attempt, len(_backoff)-1)]
                    if _is_overloaded or _is_rate_limit:
                        _time.sleep(wait)
                    else:
                        _time.sleep(3)
            return ""
        else:
            import re as _re2, time as _t2
            for _attempt in range(4):
                try:
                    _kw = dict(
                        model=self.model,
                        max_tokens=max_tokens,
                        messages=[
                            {"role": "system", "content": system},
                            {"role": "user", "content": user},
                        ],
                    )
                    if self._is_ollama:
                        # num_gpu=-1 = all layers on GPU; 0 = CPU only
                        _kw["extra_body"] = {"options": {"num_gpu": -1 if self.use_gpu else 0}}
                    resp = self._client.chat.completions.create(**_kw)
                    if on_usage and resp.usage:
                        on_usage(resp.usage.prompt_tokens, resp.usage.completion_tokens)
                    return resp.choices[0].message.content or ""
                except Exception as e:
                    err = str(e)

                    if "authentication" in err.lower() or "api_key" in err.lower() or "credentials" in err.lower() or "401" in err:
                        raise ValueError(
                            f"API key rejected by {self.provider} ({self.model}). "
                            "Check that you pasted the correct key for the selected provider."
                        )

                    if "429" in err or "RESOURCE_EXHAUSTED" in err or "quota" in err.lower() or "rate" in err.lower():
                        # Daily quota — cannot recover by waiting
                        if "PerDay" in err or ("limit: 0" in err and "PerDay" in err):
                            alt = " Try switching to a different model (e.g. gemini-2.5-flash)." if "gemini" in self.model else ""
                            raise ValueError(
                                f"Daily quota exhausted for {self.model}. "
                                f"Your free-tier daily limit is used up.{alt}"
                            )
                        # Per-minute limit — parse retry delay and wait
                        delay_m = _re2.search(r"retry in (\d+(?:\.\d+)?)s", err, _re2.IGNORECASE)
                        wait = float(delay_m.group(1)) + 2 if delay_m else (15 * (2 ** _attempt))
                        if wait > 120 or _attempt == 3:
                            raise ValueError(
                                f"Rate limit exceeded for {self.model}. "
                                + (f"Retry in {wait:.0f}s, or switch to a less busy model." if wait <= 120
                                   else "Too many retries — switch to a different model.")
                            )
                        _t2.sleep(wait)
                        continue

                    raise
            return ""

    def describe_images(self, image_paths: list, on_log=None) -> str:
        """Send images to the vision API and return combined textual descriptions."""
        import base64 as _b64
        import mimetypes as _mt
        _PROMPT = (
            "Describe this image in detail for transcription context. "
            "Extract all visible text, slide content, whiteboard notes, diagrams, "
            "code, and any speaker or agenda information."
        )
        descriptions = []
        for path in image_paths:
            name = Path(path).name
            mime = _mt.guess_type(path)[0] or "image/png"
            with open(path, "rb") as _f:
                b64 = _b64.b64encode(_f.read()).decode()
            if on_log:
                on_log(f"Analyzing image: {name}…")
            if self.provider == "anthropic":
                resp = self._client.messages.create(
                    model=self.model,
                    max_tokens=1024,
                    messages=[{"role": "user", "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}},
                        {"type": "text", "text": _PROMPT},
                    ]}],
                )
                desc = next((b.text for b in resp.content if b.type == "text"), "")
            else:
                resp = self._client.chat.completions.create(
                    model=self.model,
                    max_tokens=1024,
                    messages=[{"role": "user", "content": [
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                        {"type": "text", "text": _PROMPT},
                    ]}],
                )
                desc = resp.choices[0].message.content or ""
            descriptions.append(f"[Image: {name}]\n{desc}")
            if on_log:
                on_log(f"✅ Image analyzed: {name}")
        return "\n\n".join(descriptions)


# ── resolve bundled ffmpeg (works on Windows without a system ffmpeg install) ──
import subprocess as _sp
import shutil as _shutil
import numpy as _np

def _resolve_ffmpeg() -> str:
    # 1. imageio_ffmpeg bundles a binary — use it if the file actually exists
    try:
        import imageio_ffmpeg as _iff
        _candidate = _iff.get_ffmpeg_exe()
        if _candidate and __import__("os").path.isfile(_candidate):
            return _candidate
    except Exception:
        pass
    # 2. system ffmpeg on PATH
    _sys_ffmpeg = _shutil.which("ffmpeg")
    if _sys_ffmpeg:
        return _sys_ffmpeg
    # 3. bare name — last resort (will raise FileNotFoundError at call time
    #    with a clear message rather than a silent WinError 2)
    return "ffmpeg"

FFMPEG_EXE = _resolve_ffmpeg()

# patch Whisper's internal audio loader to use the resolved binary
try:
    import whisper.audio as _wa

    def _load_audio(file: str, sr: int = _wa.SAMPLE_RATE) -> _np.ndarray:
        cmd = [FFMPEG_EXE, "-nostdin", "-threads", "0", "-i", file,
               "-f", "s16le", "-ac", "1", "-acodec", "pcm_s16le", "-ar", str(sr), "-"]
        out = _sp.run(cmd, capture_output=True, check=True).stdout
        return _np.frombuffer(out, _np.int16).flatten().astype(_np.float32) / 32768.0

    _wa.load_audio = _load_audio
except Exception:
    pass

# ── live Whisper progress: patch tqdm so we can read % in real-time ───────────
import threading as _threading
_progress_lock  = _threading.Lock()
_progress_cb    = None   # set to a callable(float) before each transcribe call

try:
    import sys as _sys
    import tqdm as _tqdm_mod

    class _TrackingTqdm(_tqdm_mod.tqdm):
        def update(self, n=1):
            # Do NOT call super().update() — it renders Unicode progress chars
            # (e.g. ▼ U+25BC) to stderr which crashes ASCII-encoded terminals.
            if self.n is None:
                self.n = 0
            self.n += n
            with _progress_lock:
                cb = _progress_cb
            if cb and self.total and self.total > 0:
                cb(min(self.n / self.total, 1.0))

        def close(self):
            # Fire a final 100% when tqdm closes so the UI never stalls at 99%.
            with _progress_lock:
                cb = _progress_cb
            if cb:
                cb(1.0)
            super().close()

    # whisper.transcribe calls tqdm.tqdm(...) where tqdm is the MODULE.
    # Patching the function object (_wt = whisper.transcribe the function) does
    # nothing. We must patch tqdm.tqdm on the real submodule via sys.modules.
    import whisper.transcribe as _wt_unused  # ensures module is loaded
    _real_wt = _sys.modules["whisper.transcribe"]
    _real_wt.tqdm.tqdm = _TrackingTqdm
except Exception:
    pass


# ── optional dependency imports ───────────────────────────────────────────────

try:
    import whisper as openai_whisper
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def extract_profile_text(file_path: str) -> str:
    """Extract plain text from a candidate resume/profile (PDF, DOCX, or TXT)."""
    p = Path(file_path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        if not PDF_AVAILABLE:
            return ""
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n".join(parts)
    elif ext in (".docx", ".doc"):
        if not DOCX_AVAILABLE:
            return ""
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    else:
        return p.read_text(encoding="utf-8", errors="replace")


# ── format constants ──────────────────────────────────────────────────────────

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}

AUDIO_EXTS = {
    ".mp3", ".wav", ".m4a", ".flac", ".ogg", ".aac", ".opus", ".wma", ".amr",
    ".mp2", ".mp4a", ".m4b", ".m4r", ".3gp", ".3g2", ".oga", ".spx", ".caf",
    ".aiff", ".aif", ".aifc", ".au", ".ra", ".ram", ".rm", ".ac3", ".dts",
    ".ape", ".wv", ".mka", ".mid", ".midi", ".kar",
}
VIDEO_EXTS = {
    ".mp4", ".mov", ".avi", ".mkv", ".webm", ".m4v",
    ".flv", ".wmv", ".ts", ".mts", ".m2ts", ".vob", ".ogv",
    ".3gp", ".3g2", ".f4v", ".divx", ".mpg", ".mpeg", ".m2v",
    ".asf", ".rm", ".rmvb", ".dv", ".mxf",
}

# Approximate CPU realtime speed multiplier for each Whisper model size.
# e.g. "base" processes ~16 minutes of audio per minute of wall-clock time.
WHISPER_SPEED = {
    "tiny": 32, "base": 16, "small": 6, "medium": 2,
    "large": 1, "large-v2": 1, "large-v3": 1,
    "turbo": 8,   # large-v3-turbo: ~8x faster than large-v3
}


def _get_audio_duration(path: str) -> float:
    """Return audio/video duration in seconds, or 0.0 if unavailable.

    Uses ffprobe (same directory as ffmpeg) with an ffmpeg -i fallback.
    The old str.replace("ffmpeg","ffprobe") replaced every occurrence including
    the directory name, producing a non-existent path on Windows imageio bundles.
    """
    # ── Strategy 1: ffprobe in the same directory as the ffmpeg binary ────────
    try:
        _p = Path(FFMPEG_EXE)
        ffprobe = str(_p.with_name(_p.name.replace("ffmpeg", "ffprobe")))
        result = _sp.run(
            [ffprobe, "-v", "quiet", "-print_format", "json", "-show_format", path],
            capture_output=True, text=True,
        )
        data = json.loads(result.stdout)
        return float(data["format"]["duration"])
    except Exception:
        pass

    # ── Strategy 2: parse duration from ffmpeg -i stderr output ───────────────
    try:
        result = _sp.run(
            [FFMPEG_EXE, "-i", path],
            capture_output=True, text=True,
        )
        m = re.search(r"Duration:\s*(\d+):(\d+):(\d+\.?\d*)", result.stderr)
        if m:
            return int(m.group(1)) * 3600 + int(m.group(2)) * 60 + float(m.group(3))
    except Exception:
        pass

    return 0.0


def _fmt_duration(secs: float) -> str:
    secs = int(secs)
    h, rem = divmod(secs, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h}h {m:02d}m"
    if m:
        return f"{m}m {s:02d}s"
    return f"{s}s"


# ── dataclasses ───────────────────────────────────────────────────────────────

@dataclass
class SpeakerStats:
    name: str = ""
    words_per_minute: float = 0.0
    total_words: int = 0
    speaking_time_seconds: float = 0.0
    speaking_percentage: float = 0.0
    pace_label: str = ""          # Slow / Normal / Fast / Very Fast
    accent_indicators: str = ""   # Claude-inferred accent markers
    accent_confidence: str = ""   # low / medium / high


@dataclass
class ReportConfig:
    style: str = "formal"         # formal | casual | executive | bullet
    include_summary: bool = True
    include_key_points: bool = True
    include_action_items: bool = True
    include_transcript: bool = True
    include_speaker_profiles: bool = True
    include_speech_analytics: bool = True


@dataclass
class TranscriptResult:
    clean_transcript: str = ""
    speaker_dialogue: str = ""
    summary: str = ""
    key_points: list = field(default_factory=list)
    action_items: list = field(default_factory=list)
    speaker_map: dict = field(default_factory=dict)
    speaker_profiles: dict = field(default_factory=dict)
    speaker_stats: list = field(default_factory=list)
    detected_language: str = ""
    segments: list = field(default_factory=list)       # raw Whisper segments [{start,end,text}]
    stt_engine: str = "whisper_local"
    stt_seconds: float = 0.0
    interview_analysis: dict = field(default_factory=dict)  # filled when Interview Mode is on


# ── export generators ─────────────────────────────────────────────────────────

def _fmt_ts(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _fmt_srt_ts(seconds: float) -> str:
    ms = int((seconds % 1) * 1000)
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _fmt_vtt_ts(seconds: float) -> str:
    ms = int((seconds % 1) * 1000)
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def generate_srt(segments: list) -> str:
    if not segments:
        return ""
    lines = []
    for i, seg in enumerate(segments, 1):
        start = _fmt_srt_ts(float(seg.get("start", 0)))
        end   = _fmt_srt_ts(float(seg.get("end",   0)))
        text  = seg.get("text", "").strip()
        lines.append(f"{i}\n{start} --> {end}\n{text}\n")
    return "\n".join(lines)


def generate_vtt(segments: list) -> str:
    if not segments:
        return ""
    lines = ["WEBVTT", ""]
    for i, seg in enumerate(segments, 1):
        start = _fmt_vtt_ts(float(seg.get("start", 0)))
        end   = _fmt_vtt_ts(float(seg.get("end",   0)))
        text  = seg.get("text", "").strip()
        lines.append(f"{i}\n{start} --> {end}\n{text}\n")
    return "\n".join(lines)


def generate_docx(result: "TranscriptResult", stem: str, output_path: str, va_result=None) -> bool:
    if not DOCX_AVAILABLE:
        return False
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    # ── Colour helpers ────────────────────────────────────────────────────────
    _RGB = {
        "great":  RGBColor(22, 163, 74),
        "good":   RGBColor(37,  99,235),
        "ni":     RGBColor(217,119,  6),
        "missed": RGBColor(220, 38, 38),
        "header": RGBColor(30,  41, 59),
        "muted":  RGBColor(100,116,139),
        "accent": RGBColor(59, 130,246),
        "amber":  RGBColor(146, 64,  14),
        "blue":   RGBColor(30,  64,175),
    }
    _SCORE_RGB = {
        "Great": _RGB["great"], "Good": _RGB["good"],
        "Needs Improvement": _RGB["ni"], "Missed": _RGB["missed"],
    }

    def _shade_paragraph(para, hex_fill: str):
        """Apply a background shading to a paragraph."""
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        pPr.append(shd)

    def _set_para_spacing(para, before=0, after=0, line=None):
        pPr = para._p.get_or_add_pPr()
        pPr_spacing = OxmlElement("w:spacing")
        pPr_spacing.set(qn("w:before"), str(before))
        pPr_spacing.set(qn("w:after"),  str(after))
        if line:
            pPr_spacing.set(qn("w:line"), str(line))
            pPr_spacing.set(qn("w:lineRule"), "auto")
        pPr.append(pPr_spacing)

    def _add_section_heading(doc, title: str, level=1):
        h = doc.add_heading(title, level)
        h.runs[0].font.color.rgb = _RGB["header"]
        _set_para_spacing(h, before=160, after=80)
        return h

    def _add_labelled_block(doc, label: str, text: str, fill_hex: str, label_rgb: RGBColor):
        """Shaded labelled block (What was said / ideal / tip)."""
        lp = doc.add_paragraph()
        _shade_paragraph(lp, fill_hex)
        lp.paragraph_format.left_indent = Cm(0.5)
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(1)
        r = lp.add_run(label)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = label_rgb

        bp = doc.add_paragraph(text)
        _shade_paragraph(bp, fill_hex)
        bp.paragraph_format.left_indent = Cm(0.5)
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(6)
        bp.runs[0].font.size = Pt(9)
        bp.runs[0].font.color.rgb = _RGB["header"]
        return bp

    # ── Document setup ────────────────────────────────────────────────────────
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade_paragraph(title_p, "1E293B")
    _set_para_spacing(title_p, before=0, after=80)
    tr = title_p.add_run(stem)
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(255, 255, 255)

    # Meta line
    meta_parts = []
    if result.detected_language: meta_parts.append(f"Language: {result.detected_language}")
    if getattr(result, "stt_engine", ""): meta_parts.append(f"STT: {result.stt_engine}")
    if meta_parts:
        mp = doc.add_paragraph("   ".join(meta_parts))
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.runs[0].font.size = Pt(9)
        mp.runs[0].font.color.rgb = _RGB["muted"]
        _set_para_spacing(mp, before=40, after=120)

    # ── Summary ───────────────────────────────────────────────────────────────
    if result.summary:
        _add_section_heading(doc, "Summary")
        p = doc.add_paragraph(result.summary)
        p.runs[0].font.size = Pt(10)
        _set_para_spacing(p, after=100)

    # ── Key Points ────────────────────────────────────────────────────────────
    if result.key_points:
        _add_section_heading(doc, "Key Points")
        for kp in result.key_points:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(kp).font.size = Pt(10)

    # ── Action Items ──────────────────────────────────────────────────────────
    if result.action_items:
        _add_section_heading(doc, "Action Items")
        for ai in result.action_items:
            if isinstance(ai, dict):
                action   = ai.get("action", ai.get("item", str(ai)))
                owner    = ai.get("owner", "")
                tl       = ai.get("timeline", "")
                text     = action
                if owner: text += f"  (Owner: {owner})"
                if tl:    text += f"  [{tl}]"
            else:
                text = str(ai)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text).font.size = Pt(10)

    # ── Speaker Profiles ──────────────────────────────────────────────────────
    if result.speaker_profiles:
        _add_section_heading(doc, "Speaker Profiles")
        for name, profile in result.speaker_profiles.items():
            np = doc.add_paragraph()
            np.add_run(name + ": ").bold = True
            np.add_run(profile)

    # ── Transcript ────────────────────────────────────────────────────────────
    if result.speaker_dialogue:
        _add_section_heading(doc, "Speaker Dialogue")
        p = doc.add_paragraph(result.speaker_dialogue)
        p.runs[0].font.size = Pt(9)
    elif result.clean_transcript:
        _add_section_heading(doc, "Transcript")
        p = doc.add_paragraph(result.clean_transcript)
        p.runs[0].font.size = Pt(9)

    # ── Interview Coaching Analysis ───────────────────────────────────────────
    ia = result.interview_analysis or {}
    if ia and not ia.get("parse_error"):
        _add_section_heading(doc, "Interview Coaching Analysis")

        # Score banner table (3 cells: Score | Advance Likelihood | Deflection Rate)
        score   = ia.get("overall_score", "—")
        verdict = ia.get("overall_verdict", "")
        adv     = ia.get("advance_likelihood", "")
        defl_rt = ia.get("deflection_rate", "")

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        cells = tbl.rows[0].cells
        for i, (lbl, val, col_hex) in enumerate([
            (f"{score}/10", "Overall Score",  "166534"),
            (f"{adv}%"     if adv   else "—", "Advance Likelihood", "1D4ED8"),
            (f"{defl_rt}%" if defl_rt else "—", "Deflection Rate", "92400E"),
        ]):
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), col_hex)
            cells[i]._tc.get_or_add_tcPr().append(shd)
            p1 = cells[i].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(lbl)
            r1.bold = True; r1.font.size = Pt(16)
            r1.font.color.rgb = RGBColor(255,255,255)
            p2 = cells[i].add_paragraph(val)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.runs[0].font.size = Pt(8)
            p2.runs[0].font.color.rgb = RGBColor(200,220,255)

        if verdict:
            vp = doc.add_paragraph(verdict)
            vp.runs[0].font.size = Pt(10)
            vp.runs[0].italic = True
            vp.runs[0].font.color.rgb = _RGB["muted"]
            _set_para_spacing(vp, before=80, after=120)

        # Per-question breakdown
        qs = ia.get("questions", [])
        if qs:
            _add_section_heading(doc, "Question Breakdown", level=2)
            _DEFL_LABEL = {"partial": "! PARTIALLY DEFLECTED", "full": "X DID NOT ANSWER"}
            _DEFL_RGB   = {"partial": _RGB["ni"], "full": _RGB["missed"]}

            for q in qs:
                qid       = q.get("id", "")
                question  = q.get("question", "")
                sc        = q.get("score", "")
                reason    = q.get("score_reason", "")
                dfl       = (q.get("deflection") or "none").lower().strip()
                dfl_note  = q.get("deflection_note", "")
                said      = q.get("answer_said") or q.get("answer_summary", "")
                ideal     = q.get("model_answer") or q.get("ideal_answer", "")
                tip       = q.get("coaching_tip", "")
                sc_rgb    = _SCORE_RGB.get(sc, _RGB["muted"])

                # Q header row (shaded)
                qp = doc.add_paragraph()
                _shade_paragraph(qp, "F1F5F9")
                _set_para_spacing(qp, before=120, after=20)
                r_id = qp.add_run(f"Q{qid}  ")
                r_id.bold = True; r_id.font.size = Pt(9)
                r_id.font.color.rgb = _RGB["muted"]
                r_q = qp.add_run(question)
                r_q.bold = True; r_q.font.size = Pt(10)
                r_q.font.color.rgb = _RGB["header"]

                # Score badge line
                sp = doc.add_paragraph()
                _set_para_spacing(sp, before=0, after=20)
                r_sc = sp.add_run(f"  {sc.upper()}  ")
                r_sc.bold = True; r_sc.font.size = Pt(9)
                r_sc.font.color.rgb = sc_rgb
                if reason:
                    r_rs = sp.add_run(f"  —  {reason}")
                    r_rs.font.size = Pt(9)
                    r_rs.font.color.rgb = _RGB["muted"]
                    r_rs.italic = True

                # Deflection
                if dfl in _DEFL_LABEL:
                    dp = doc.add_paragraph()
                    _set_para_spacing(dp, before=0, after=20)
                    dr = dp.add_run(f"  {_DEFL_LABEL[dfl]}")
                    dr.bold = True; dr.font.size = Pt(9)
                    dr.font.color.rgb = _DEFL_RGB.get(dfl, _RGB["ni"])
                    if dfl_note:
                        dnp = doc.add_paragraph(dfl_note)
                        dnp.runs[0].font.size = Pt(9)
                        dnp.runs[0].italic = True
                        dnp.paragraph_format.left_indent = Cm(0.5)

                if said:  _add_labelled_block(doc, "WHAT WAS SAID",           said,  "F8FAFC", _RGB["muted"])
                if ideal: _add_labelled_block(doc, "WHAT YOU COULD HAVE SAID",ideal, "EFF6FF", _RGB["blue"])
                if tip:   _add_labelled_block(doc, "COACHING TIP",             tip,  "FFFBEB", _RGB["amber"])

        # Deep analysis
        adv_reason = ia.get("advance_reasoning", "")
        if adv_reason:
            _add_section_heading(doc, "Deep Analysis", level=2)
            dp = doc.add_paragraph(adv_reason)
            dp.runs[0].font.size = Pt(10)

    # ── Video Delivery Analysis ───────────────────────────────────────────────
    if va_result and not getattr(va_result, "error", None) and getattr(va_result, "persons", None):
        _add_section_heading(doc, "Video Delivery Analysis")

        # Overall summary row
        ovr = doc.add_paragraph()
        ovr.add_run(f"Overall Score: {va_result.overall_score:.0f}/100").bold = True
        ovr.runs[0].font.color.rgb = _RGB["header"]
        ovr.add_run(f"   Duration: {int(va_result.duration_seconds//60)}m {int(va_result.duration_seconds%60)}s"
                    f"   Participants: {va_result.person_count}")
        _set_para_spacing(ovr, after=80)

        for _pid, _p in va_result.persons.items():
            sc_rgb = _SCORE_RGB.get(
                "Great" if _p.overall >= 80 else "Good" if _p.overall >= 65 else
                "Needs Improvement" if _p.overall >= 50 else "Missed", _RGB["muted"])

            # Person header
            ph = doc.add_paragraph()
            _shade_paragraph(ph, "F1F5F9")
            _set_para_spacing(ph, before=120, after=20)
            r_role = ph.add_run(_p.role)
            r_role.bold = True; r_role.font.size = Pt(11)
            r_role.font.color.rgb = _RGB["header"]
            r_sc = ph.add_run(f"   {_p.overall:.0f}/100")
            r_sc.bold = True; r_sc.font.size = Pt(11)
            r_sc.font.color.rgb = sc_rgb

            # Metrics
            mp = doc.add_paragraph()
            mp.add_run("Confidence: ").bold = True
            mp.add_run(f"{_p.confidence:.0f}   ")
            mp.add_run("Composure: ").bold = True
            mp.add_run(f"{_p.composure:.0f}   ")
            mp.add_run("Eye Contact: ").bold = True
            mp.add_run(f"{_p.eye_contact:.0f}   ")
            mp.add_run("Engagement: ").bold = True
            mp.add_run(f"{_p.engagement:.0f}   ")
            mp.add_run("Energy: ").bold = True
            mp.add_run(f"{_p.energy:.0f}")
            for r in mp.runs: r.font.size = Pt(9)
            _set_para_spacing(mp, after=20)

            # Body language
            bl = doc.add_paragraph()
            bl.add_run("Open Posture: ").bold = True
            bl.add_run(f"{_p.open_body_pct:.0f}%   ")
            bl.add_run("Arms Crossed: ").bold = True
            bl.add_run(f"{_p.arm_crossed_pct:.0f}%   ")
            bl.add_run("Forward Lean: ").bold = True
            bl.add_run(f"{_p.forward_lean_pct:.0f}%   ")
            bl.add_run("Dominant Mood: ").bold = True
            bl.add_run(_p.dominant_emotion)
            for r in bl.runs: r.font.size = Pt(9)
            _set_para_spacing(bl, after=20)

            # Cultural scores
            if _p.cultural:
                cp = doc.add_paragraph()
                cp.add_run("American Interview Standard: ").bold = True
                cp.add_run(f"{_p.cultural.american_score:.0f}/100   ")
                cp.add_run("Indian→American Adaptation: ").bold = True
                cp.add_run(f"{_p.cultural.adaptation_score:.0f}/100")
                for r in cp.runs: r.font.size = Pt(9)
                _set_para_spacing(cp, after=10)
                for tip in (_p.cultural.american_tips or [])[:3]:
                    tp = doc.add_paragraph(style="List Bullet")
                    tp.add_run(tip).font.size = Pt(9)
                for tip in (_p.cultural.adaptation_tips or [])[:3]:
                    tp = doc.add_paragraph(style="List Bullet")
                    tp.add_run(f"→ {tip}").font.size = Pt(9)

        if getattr(va_result, "observations", None):
            _add_section_heading(doc, "Key Observations", level=2)
            for obs in va_result.observations:
                op = doc.add_paragraph(style="List Bullet")
                op.add_run(obs).font.size = Pt(9)

    doc.save(output_path)
    return True


# ── file loaders ──────────────────────────────────────────────────────────────


def load_audio_video(path: str, model_size: str = "base", on_progress=None,
                     on_stage_change=None, language: str = None, on_log=None,
                     use_gpu: bool = True) -> str:
    """Transcribe audio/video using OpenAI Whisper with timestamps.
    on_progress(pct: float) — live 0.0-1.0 progress updates.
    on_log(msg: str)        — human-readable step-by-step log messages.
    language: ISO-639-1 code (e.g. "es", "en") or None for auto-detect.
    """
    def _log(m):
        _safe_print(f"  {m}")
        if on_log: on_log(m)

    global _progress_cb
    if not WHISPER_AVAILABLE:
        raise ImportError("Run: pip install openai-whisper")

    # ── detect file duration for ETA estimate ─────────────────────────────────
    dur_secs = _get_audio_duration(path)
    dur_note = f" ({_fmt_duration(dur_secs)} audio)" if dur_secs > 0 else ""
    speed    = WHISPER_SPEED.get(model_size, 8)
    if dur_secs > 0:
        est_secs = dur_secs / speed
        _log(f"File duration: {_fmt_duration(dur_secs)}  |  Est. transcription time: {_fmt_duration(est_secs)} (Whisper '{model_size}')")
    else:
        _log(f"File duration: unknown  |  Loading Whisper '{model_size}' model…")

    device = "cpu"
    if use_gpu:
        try:
            import torch as _torch
            if _torch.cuda.is_available():
                device = "cuda"                          # NVIDIA
            elif hasattr(_torch.backends, "mps") and _torch.backends.mps.is_available():
                device = "mps"                           # Apple Silicon
            else:
                # Check DirectML (AMD / Intel on Windows)
                try:
                    import torch_directml as _dml
                    device = _dml.device()               # DirectML device handle
                except ImportError:
                    pass
        except Exception:
            device = "cpu"
    _device_label = device if isinstance(device, str) else "DirectML"
    if dur_secs > 0:
        _log(f"Loading Whisper '{model_size}' model… (device: {_device_label})")
    model = openai_whisper.load_model(model_size, device=device)
    if device == "cuda":
        try:
            import torch as _t
            _t.backends.cudnn.benchmark = True       # faster convolutions for fixed input sizes
            _t.backends.cuda.matmul.allow_tf32 = True  # ~10% faster matmul on Ampere+
        except Exception:
            pass
    _log(f"Model loaded on {_device_label}.")

    with _progress_lock:
        _progress_cb = on_progress

    audio_path = path
    tmp_path = None
    if Path(path).suffix.lower() in VIDEO_EXTS:
        _log("Extracting audio track from video…")
        tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        tmp_path = tmp.name
        tmp.close()
        # -vn: ignore video, -sn: ignore subtitles, force pcm_s16le mono 16k for Whisper
        try:
            proc = _sp.run(
                [FFMPEG_EXE, "-y", "-i", path,
                 "-vn", "-sn", "-acodec", "pcm_s16le",
                 "-ar", "16000", "-ac", "1",
                 tmp_path],
                capture_output=True,
            )
        except FileNotFoundError:
            raise RuntimeError(
                "ffmpeg not found. Re-run setup_windows.bat (or setup_mac.sh on Mac) "
                "to reinstall dependencies, then try again."
            )
        if proc.returncode != 0:
            err = (proc.stderr or b"").decode("utf-8", errors="replace").strip()
            tail = "\n".join(err.splitlines()[-8:]) if err else "(no stderr output)"
            raise RuntimeError(
                "ffmpeg failed extracting audio from video.\n"
                f"ffmpeg path: {FFMPEG_EXE}\n"
                f"input: {path}\n"
                f"stderr (tail):\n{tail}"
            )
        _log("Audio extraction complete.")
        audio_path = tmp_path

    if on_stage_change: on_stage_change("whisper")
    lang_note = f" (language: {language})" if language else " (language: auto-detect)"
    _log(f"Starting transcription{lang_note}…  This is the longest step.")
    transcribe_kwargs = {"verbose": False}
    if language:
        transcribe_kwargs["language"] = language
    try:
        result = model.transcribe(audio_path, **transcribe_kwargs)
    finally:
        with _progress_lock:
            _progress_cb = None

    total_words = sum(len(s["text"].split()) for s in result.get("segments", []))
    detected = result.get("language", "")
    _log(f"Transcription complete! ~{total_words:,} words detected{dur_note}."
         + (f"  Language: {detected}" if detected and not language else ""))

    if tmp_path:
        os.unlink(tmp_path)

    segs  = result.get("segments", [])
    lines = []
    for seg in segs:
        lines.append(f"[{_fmt_ts(seg['start'])}] {seg['text'].strip()}")
    return "\n".join(lines), detected, segs


def load_audio_video_panel(
    path: str, model_size: str = "base", num_speakers: int = None,
    language: str = None, on_log=None, use_gpu: bool = True,
) -> tuple:
    """
    Transcribe + diarize with WhisperX.
    Returns (timestamped_transcript, raw_whisperx_result).
    language: ISO-639-1 code or None for auto-detect.
    """
    def _log(m):
        _safe_print(f"  {m}")
        if on_log: on_log(m)

    try:
        import whisperx
        import torch
    except ImportError:
        _log("WhisperX not installed — falling back to standard Whisper (no diarization)")
        text, _ = load_audio_video(path, model_size, language=language, on_log=on_log)
        return text, {}

    hf_token = os.environ.get("HF_TOKEN", "")
    if not hf_token:
        _log("HF_TOKEN not set — falling back to standard Whisper (no diarization)")
        text, _ = load_audio_video(path, model_size, language=language, on_log=on_log)
        return text, {}

    audio_path = path
    tmp_path = None
    if Path(path).suffix.lower() in VIDEO_EXTS:
        _safe_print("  Extracting audio from video...")
        tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        tmp_path = tmp.name
        tmp.close()
        _sp.run([FFMPEG_EXE, "-i", path, "-ar", "16000", "-ac", "1", "-y", tmp_path, "-loglevel", "error"],
                capture_output=True)
        audio_path = tmp_path

    device = "cpu"
    if use_gpu:
        try:
            import torch as _t2
            if _t2.cuda.is_available():
                device = "cuda"
            elif hasattr(_t2.backends, "mps") and _t2.backends.mps.is_available():
                device = "mps"
        except Exception:
            device = "cpu"
    compute_type = "float16" if device in ("cuda", "mps") else "int8"

    lang_note = f" (language: {language})" if language else " (language: auto-detect)"
    dur_secs = _get_audio_duration(audio_path)
    if dur_secs > 0:
        speed = WHISPER_SPEED.get(model_size, 8)
        _log(f"File duration: {_fmt_duration(dur_secs)}  |  Est. WhisperX time: {_fmt_duration(dur_secs / speed)}")
    _log(f"Transcribing with WhisperX ({model_size}){lang_note}…")
    model = whisperx.load_model(model_size, device, compute_type=compute_type, language=language)
    audio = whisperx.load_audio(audio_path)
    transcribe_kwargs = {"batch_size": 16}
    if language:
        transcribe_kwargs["language"] = language
    result = model.transcribe(audio, **transcribe_kwargs)

    _log("Aligning word timestamps…")
    model_a, metadata = whisperx.load_align_model(
        language_code=result["language"], device=device
    )
    result = whisperx.align(result["segments"], model_a, metadata, audio, device)

    _log(f"Running speaker diarization{(' (' + str(num_speakers) + ' speakers)') if num_speakers else ''}…")
    diarize_model = whisperx.DiarizationPipeline(use_auth_token=hf_token, device=device)
    diarize_kwargs = {"num_speakers": num_speakers} if num_speakers else {}
    diarize_segments = diarize_model(audio_path, **diarize_kwargs)
    result = whisperx.assign_word_speakers(diarize_segments, result)
    _log("Diarization complete.")

    if tmp_path:
        os.unlink(tmp_path)

    lines = []
    for seg in result["segments"]:
        ts = _fmt_ts(seg["start"])
        speaker = seg.get("speaker", "UNKNOWN")
        lines.append(f"[{ts}] {speaker}: {seg['text'].strip()}")

    return "\n".join(lines), result


def load_srt(path: str) -> str:
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    lines, ts, buf = [], "", []
    for line in text.splitlines():
        line = line.strip()
        if re.match(r"^\d+$", line):
            if buf:
                lines.append(f"[{ts}] {' '.join(buf)}")
            buf = []
        elif "-->" in line:
            ts = line.split("-->")[0].strip()[:8]
        elif line:
            buf.append(line)
    if buf:
        lines.append(f"[{ts}] {' '.join(buf)}")
    return "\n".join(lines)


def load_vtt(path: str) -> str:
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    lines, ts, buf = [], "", []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith(("WEBVTT", "NOTE", "STYLE", "REGION")):
            continue
        if "-->" in line:
            if buf:
                lines.append(f"[{ts}] {' '.join(buf)}")
            ts = line.split("-->")[0].strip()[:8]
            buf = []
        elif line:
            buf.append(line)
    if buf:
        lines.append(f"[{ts}] {' '.join(buf)}")
    return "\n".join(lines)


def load_docx(path: str) -> str:
    if not DOCX_AVAILABLE:
        raise ImportError("Run: pip install python-docx")
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def load_pdf(path: str) -> str:
    if not PDF_AVAILABLE:
        raise ImportError("Run: pip install pdfplumber")
    parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"[Page {i}]\n{text}")
    return "\n\n".join(parts)


def load_file(path: str, whisper_model: str = "base") -> tuple:
    """Returns (raw_text, format_label)."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = p.suffix.lower()
    _safe_print(f"  Format detected: {ext or 'text'}")

    if ext in AUDIO_EXTS:
        return load_audio_video(path, whisper_model), "audio recording"
    if ext in VIDEO_EXTS:
        return load_audio_video(path, whisper_model), "video recording"
    if ext == ".srt":
        return load_srt(path), "SRT subtitle file"
    if ext == ".vtt":
        return load_vtt(path), "WebVTT subtitle file"
    if ext == ".docx":
        return load_docx(path), "Word document"
    if ext == ".pdf":
        return load_pdf(path), "PDF document"
    return p.read_text(encoding="utf-8", errors="replace"), "plain text"


# ── STT engine dispatcher ─────────────────────────────────────────────────────
# Each engine returns (timestamped_text: str, detected_lang: str, segments: list)
# segments is a list of {start, end, text} dicts (empty list if engine doesn't provide them)

STT_ENGINES = {
    "whisper_local":   "Whisper (Local / Offline)",
    "openai_whisper":  "OpenAI Whisper API",
    "groq_whisper":    "Groq Whisper",
    "deepgram":        "Deepgram",
    "assemblyai":      "AssemblyAI",
    "google_stt":      "Google Cloud STT",
    "azure_speech":    "Azure Speech",
    "elevenlabs":      "ElevenLabs Scribe",
    "revai":           "Rev.ai",
}


def _stt_openai_api(path: str, api_key: str, language: str = None, on_log=None,
                    model: str = "whisper-1") -> tuple:
    try:
        import openai as _oai
    except ImportError:
        raise ImportError("openai package required: pip install openai")
    client = _oai.OpenAI(api_key=api_key)
    _model = model or "whisper-1"

    def _log(m):
        if on_log: on_log(m)

    def _transcribe_one(chunk_path):
        if _model in ("gpt-4o-transcribe", "gpt-4o-mini-transcribe"):
            with open(chunk_path, "rb") as f:
                return client.audio.transcriptions.create(
                    model=_model, file=f, response_format="verbose_json",
                    **( {"language": language} if language else {} )
                )
        with open(chunk_path, "rb") as f:
            kw = {"model": _model, "file": f, "response_format": "verbose_json"}
            if language:
                kw["language"] = language
            return client.audio.transcriptions.create(**kw)

    chunks = _split_audio_for_api(path, chunk_secs=1200)
    is_chunked = len(chunks) > 1
    if is_chunked:
        _log(f"Audio split into {len(chunks)} x 20-min chunks for OpenAI API (25 MB limit)")

    all_segs, all_lines, detected_lang = [], [], ""
    for idx, (chunk_path, offset) in enumerate(chunks):
        if is_chunked:
            _log(f"Transcribing chunk {idx + 1}/{len(chunks)}...")
        try:
            resp = _transcribe_one(chunk_path)
            detected_lang = detected_lang or getattr(resp, "language", "")
            segs = [{"start": s.start + offset, "end": s.end + offset, "text": s.text}
                    for s in (resp.segments or [])]
            all_segs.extend(segs)
            all_lines.extend(f"[{_fmt_ts(s['start'])}] {s['text'].strip()}" for s in segs)
        finally:
            if is_chunked and chunk_path != path:
                try: os.unlink(chunk_path)
                except Exception: pass

    return "\n".join(all_lines) or "", detected_lang, all_segs


def _stt_groq(path: str, api_key: str, language: str = None, on_log=None,
              model: str = "whisper-large-v3-turbo") -> tuple:
    try:
        from groq import Groq
    except ImportError:
        raise ImportError("groq package required: pip install groq")
    client = Groq(api_key=api_key)
    _model = model or "whisper-large-v3-turbo"

    def _log(m):
        if on_log: on_log(m)

    chunks = _split_audio_for_api(path, chunk_secs=1200)
    is_chunked = len(chunks) > 1
    if is_chunked:
        _log(f"Audio split into {len(chunks)} x 20-min chunks for Groq API (25 MB limit)")

    all_segs, all_lines, detected_lang = [], [], ""
    for idx, (chunk_path, offset) in enumerate(chunks):
        if is_chunked:
            _log(f"Transcribing chunk {idx + 1}/{len(chunks)}...")
        try:
            with open(chunk_path, "rb") as f:
                kw = {"model": _model, "file": (Path(chunk_path).name, f),
                      "response_format": "verbose_json", "timestamp_granularities": ["segment"]}
                if language:
                    kw["language"] = language
                resp = client.audio.transcriptions.create(**kw)
            detected_lang = detected_lang or getattr(resp, "language", "")
            segs = [{"start": s.start + offset, "end": s.end + offset, "text": s.text}
                    for s in (getattr(resp, "segments", None) or [])]
            all_segs.extend(segs)
            all_lines.extend(f"[{_fmt_ts(s['start'])}] {s['text'].strip()}" for s in segs)
        finally:
            if is_chunked and chunk_path != path:
                try: os.unlink(chunk_path)
                except Exception: pass

    return "\n".join(all_lines) or "", detected_lang, all_segs


def _stt_deepgram(path: str, api_key: str, language: str = None, on_log=None,
                  model: str = "nova-2", on_stage_change=None) -> tuple:
    try:
        import httpx as _httpx
        from deepgram import DeepgramClient, PrerecordedOptions
    except ImportError:
        raise ImportError("deepgram-sdk required: pip install deepgram-sdk")
    dg = DeepgramClient(api_key)
    effective_model = model or "nova-2"
    if on_log:
        on_log(f"Deepgram model: {effective_model}", "info")
    lang_kw = {"language": language} if language else {"detect_language": True}
    opts = PrerecordedOptions(
        model=effective_model,
        punctuate=True, diarize=True, utterances=True,
        smart_format=True, numerals=True,
        **lang_kw,
    )

    if on_stage_change:
        on_stage_change("extracting")

    # Extract audio from video files before uploading — a 3-hour MP4 can be
    # several GB; the same content as 64 kbps mono MP3 is ~90 MB.  This also
    # removes the timeout issue: Deepgram processes at ~5× real-time so a
    # 3-hour file takes ~36 min server-side; we use no read timeout so the
    # connection never gets cut while waiting for the response.
    upload_path = path
    _tmp_audio  = None
    ext = Path(path).suffix.lower()
    if ext in VIDEO_EXTS or ext in AUDIO_EXTS:
        try:
            _tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
            _tmp.close()
            _sp.run(
                [FFMPEG_EXE, "-y", "-i", path,
                 "-acodec", "libmp3lame", "-ar", "16000", "-ac", "1", "-q:a", "5",
                 _tmp.name],
                capture_output=True, check=True,
            )
            upload_path = _tmp.name
            _tmp_audio  = _tmp.name
            if on_log:
                sz_mb = Path(upload_path).stat().st_size / 1_048_576
                on_log(f"Audio extracted for upload ({sz_mb:.1f} MB)", "info")
        except FileNotFoundError:
            if on_log:
                on_log(
                    "ffmpeg not found — audio extraction skipped, uploading original. "
                    "Re-run setup_windows.bat (or setup_mac.sh) to fix this.",
                    "warn",
                )
        except Exception as _e:
            if on_log:
                on_log(f"Audio extraction failed ({_e}), uploading original", "warn")

    if on_stage_change:
        on_stage_change("stt_cloud")

    try:
        with open(upload_path, "rb") as f:
            resp = dg.listen.rest.v("1").transcribe_file(
                {"buffer": f, "mimetype": "audio/mpeg"},
                opts,
                timeout=_httpx.Timeout(None, connect=15.0),  # no read timeout
            )
    finally:
        if _tmp_audio:
            try: os.unlink(_tmp_audio)
            except Exception: pass

    result = resp.results.channels[0].alternatives[0]
    detected_lang = getattr(resp.results.channels[0], "detected_language", None) or language or "en"
    segs, lines = [], []
    for utt in (resp.results.utterances or []):
        spk = f"Speaker {utt.speaker}: " if utt.speaker is not None else ""
        segs.append({"start": utt.start, "end": utt.end, "text": utt.transcript, "speaker": utt.speaker})
        lines.append(f"[{_fmt_ts(utt.start)}] {spk}{utt.transcript}")
    return "\n".join(lines) or result.transcript, detected_lang, segs


def _stt_assemblyai(path: str, api_key: str, language: str = None, on_log=None,
                    model: str = "best") -> tuple:
    try:
        import assemblyai as aai
    except ImportError:
        raise ImportError("assemblyai required: pip install assemblyai")
    aai.settings.api_key = api_key
    _model = getattr(aai.SpeechModel, model or "best", aai.SpeechModel.best)
    config = aai.TranscriptionConfig(
        speech_model=_model,
        language_code=language or None,
        language_detection=not bool(language),
        speaker_labels=True,
    )
    transcript = aai.Transcriber().transcribe(path, config=config)
    if transcript.status == aai.TranscriptStatus.error:
        raise RuntimeError(f"AssemblyAI error: {transcript.error}")
    segs, lines = [], []
    for utt in (transcript.utterances or []):
        start = utt.start / 1000.0
        end   = utt.end / 1000.0
        spk = f"Speaker {utt.speaker}: " if getattr(utt, "speaker", None) else ""
        segs.append({"start": start, "end": end, "text": utt.text, "speaker": getattr(utt, "speaker", None)})
        lines.append(f"[{_fmt_ts(start)}] {spk}{utt.text}")
    detected = getattr(transcript, "language_code", language or "en")
    return "\n".join(lines) or transcript.text, detected, segs


def _stt_google(path: str, api_key: str, language: str = None, on_log=None,
                model: str = "latest_long") -> tuple:
    try:
        from google.cloud import speech as _gspeech
        import google.auth as _gauth
    except ImportError:
        raise ImportError("google-cloud-speech required: pip install google-cloud-speech")
    import os as _os
    if api_key and not _os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
        raise ValueError("Google Cloud STT requires a service-account JSON file path as the API key. "
                         "Set it in the API key field.")
    _os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = api_key
    client = _gspeech.SpeechClient()
    with open(path, "rb") as f:
        audio = _gspeech.RecognitionAudio(content=f.read())
    lang = (language or "en") + "-US" if len(language or "en") == 2 else (language or "en-US")
    cfg = _gspeech.RecognitionConfig(
        encoding=_gspeech.RecognitionConfig.AudioEncoding.LINEAR16,
        language_code=lang,
        model=model or "latest_long",
        enable_automatic_punctuation=True,
        enable_word_time_offsets=True,
    )
    resp = client.recognize(config=cfg, audio=audio)
    lines, segs = [], []
    t = 0.0
    for r in resp.results:
        alt = r.alternatives[0]
        words = alt.words
        start = words[0].start_time.total_seconds() if words else t
        end   = words[-1].end_time.total_seconds()  if words else t + 3
        segs.append({"start": start, "end": end, "text": alt.transcript})
        lines.append(f"[{_fmt_ts(start)}] {alt.transcript}")
        t = end
    return "\n".join(lines), lang, segs


def _split_audio_for_api(path: str, chunk_secs: int = 1200) -> list:
    """Split audio into (chunk_path, start_offset_secs) tuples for APIs with a 25 MB cap.
    Encodes as 80 kbps mono MP3 — 20-min chunk ≈ 12 MB, well under the limit."""
    import math as _math
    dur = _get_audio_duration(path)
    if dur <= 0 or dur <= chunk_secs:
        return [(path, 0.0)]
    n = _math.ceil(dur / chunk_secs)
    chunks = []
    for i in range(n):
        tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
        tmp.close()
        offset = i * chunk_secs
        _sp.run(
            [FFMPEG_EXE, "-y", "-i", path,
             "-ss", str(offset), "-t", str(chunk_secs),
             "-acodec", "libmp3lame", "-ar", "16000", "-ac", "1", "-q:a", "5", tmp.name],
            capture_output=True,
        )
        chunks.append((tmp.name, float(offset)))
    return chunks


def _azure_split_audio(path: str, chunk_secs: int = 55) -> list:
    """Split audio into chunks for Azure (which has a 60s REST limit)."""
    import math as _math
    dur = _get_audio_duration(path)
    if dur <= 0 or dur <= chunk_secs:
        return [path]
    n = _math.ceil(dur / chunk_secs)
    chunks = []
    for i in range(n):
        tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
        tmp.close()
        _sp.run(
            [FFMPEG_EXE, "-y", "-i", path,
             "-ss", str(i * chunk_secs), "-t", str(chunk_secs),
             "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", tmp.name],
            capture_output=True,
        )
        chunks.append(tmp.name)
    return chunks


def _stt_azure(path: str, api_key: str, language: str = None, on_log=None,
               model: str = "conversation") -> tuple:
    try:
        import azure.cognitiveservices.speech as _az
    except ImportError:
        raise ImportError("azure-cognitiveservices-speech required: pip install azure-cognitiveservices-speech")
    # api_key format: "KEY|REGION" e.g. "abc123|eastus"
    sep = "|" if "|" in (api_key or "") else ":"
    parts = (api_key or "").split(sep, 1)
    if len(parts) != 2:
        raise ValueError("Azure key must be in format KEY|REGION (e.g. abc123|eastus)")
    key, region = parts

    def _log(m):
        if on_log: on_log(m)

    lang_code = ((language or "en") + "-US" if len(language or "en") == 2
                 else (language or "en-US"))

    def _transcribe_chunk(chunk_path: str) -> tuple:
        cfg = _az.SpeechConfig(subscription=key, region=region)
        cfg.speech_recognition_language = lang_code
        # Apply recognition mode
        if model == "dictation":
            cfg.set_property(_az.PropertyId.SpeechServiceConnection_RecoMode, "DICTATION")
        elif model == "command_and_search":
            cfg.set_property(_az.PropertyId.SpeechServiceConnection_RecoMode, "INTERACTIVE")
        audio_cfg = _az.audio.AudioConfig(filename=chunk_path)
        recognizer = _az.SpeechRecognizer(speech_config=cfg, audio_config=audio_cfg)
        results_chunk = []
        done_ev = [False]
        recognizer.session_stopped.connect(lambda e: done_ev.__setitem__(0, True))
        recognizer.canceled.connect(lambda e: done_ev.__setitem__(0, True))
        recognizer.recognized.connect(lambda e: results_chunk.append(e.result))
        recognizer.start_continuous_recognition()
        import time as _t
        while not done_ev[0]:
            _t.sleep(0.3)
        recognizer.stop_continuous_recognition()
        return results_chunk

    # Auto-chunk for long recordings (Azure REST limit is 60s)
    chunks = _azure_split_audio(path, chunk_secs=55)
    is_chunked = len(chunks) > 1
    if is_chunked:
        _log(f"Audio split into {len(chunks)} chunks (55s each) for Azure")

    all_results = []
    for i, chunk_path in enumerate(chunks):
        if is_chunked:
            _log(f"Transcribing chunk {i+1}/{len(chunks)}…")
        all_results.extend(_transcribe_chunk(chunk_path))
        if is_chunked and chunk_path != path:
            try: os.unlink(chunk_path)
            except Exception: pass

    lines, segs = [], []
    for r in all_results:
        if r.reason.name == "RecognizedSpeech":
            lines.append(r.text)
            segs.append({"start": 0, "end": 0, "text": r.text})
    return "\n".join(lines), language or "en", segs


def _stt_elevenlabs(path: str, api_key: str, language: str = None, on_log=None,
                    model: str = "scribe_v1") -> tuple:
    try:
        from elevenlabs.client import ElevenLabs
    except ImportError:
        raise ImportError("elevenlabs required: pip install elevenlabs")
    client = ElevenLabs(api_key=api_key)
    with open(path, "rb") as f:
        resp = client.speech_to_text.convert(
            file=f,
            model_id=model or "scribe_v1",
            language_code=language or None,
        )
    segs, lines = [], []
    for w in (getattr(resp, "words", None) or []):
        if getattr(w, "type", "") == "word":
            segs.append({"start": w.start or 0, "end": w.end or 0, "text": w.text})
    text = getattr(resp, "text", "") or " ".join(s["text"] for s in segs)
    return text, getattr(resp, "language_code", language or "en"), segs


def _stt_revai(path: str, api_key: str, language: str = None, on_log=None,
               model: str = "machine") -> tuple:
    try:
        from rev_ai import apiclient, jobstatus
        import time as _time
    except ImportError:
        raise ImportError("rev_ai required: pip install rev_ai")
    client = apiclient.RevAiAPIClient(api_key)
    # "fusion" uses Rev.ai's highest-accuracy pipeline (premium tier)
    submit_kw = {"language": language or "en"}
    if model == "fusion":
        submit_kw["metadata"] = "fusion"   # signals premium pipeline
    job = client.submit_job_local_file(path, **submit_kw)
    while True:
        details = client.get_job_details(job.id)
        if details.status == jobstatus.JobStatus.TRANSCRIBED:
            break
        if details.status == jobstatus.JobStatus.FAILED:
            raise RuntimeError(f"Rev.ai job failed: {details.failure_detail}")
        _time.sleep(3)
    transcript = client.get_transcript_object(job.id)
    segs, lines = [], []
    for mono in transcript.monologues:
        words = mono.elements
        text  = "".join(w.value for w in words if w.type == "text")
        if words:
            start = words[0].timestamp
            end   = words[-1].end_timestamp
        else:
            start = end = 0
        segs.append({"start": start, "end": end, "text": text})
        lines.append(f"[{_fmt_ts(start)}] {text}")
    return "\n".join(lines), language or "en", segs


def stt_transcribe(
    path: str, engine: str, api_key: str = None,
    whisper_model: str = "base", language: str = None,
    stt_model: str = None,   # model selection for cloud engines
    on_progress=None, on_stage_change=None, on_log=None,
    use_gpu: bool = True,
) -> tuple:
    """Unified STT dispatcher. Returns (text, detected_lang, segments, stt_secs)."""
    import time as _t
    t0 = _t.time()
    def _log(m):
        _safe_print(f"  [STT] {m}")
        if on_log: on_log(m)

    model_note = f" ({stt_model})" if stt_model else ""
    _log(f"STT engine: {STT_ENGINES.get(engine, engine)}{model_note}")

    if engine == "whisper_local":
        if on_stage_change: on_stage_change("extracting")
        text, lang, segs = load_audio_video(
            path, whisper_model,
            on_progress=on_progress,
            on_stage_change=on_stage_change,
            language=language,
            on_log=on_log,
            use_gpu=use_gpu,
        )
    elif engine == "openai_whisper":
        text, lang, segs = _stt_openai_api(path, api_key, language, on_log, model=stt_model)
    elif engine == "groq_whisper":
        text, lang, segs = _stt_groq(path, api_key, language, on_log, model=stt_model)
    elif engine == "deepgram":
        text, lang, segs = _stt_deepgram(path, api_key, language, on_log, model=stt_model, on_stage_change=on_stage_change)
    elif engine == "assemblyai":
        text, lang, segs = _stt_assemblyai(path, api_key, language, on_log, model=stt_model)
    elif engine == "google_stt":
        text, lang, segs = _stt_google(path, api_key, language, on_log, model=stt_model)
    elif engine == "azure_speech":
        text, lang, segs = _stt_azure(path, api_key, language, on_log, model=stt_model)
    elif engine == "elevenlabs":
        text, lang, segs = _stt_elevenlabs(path, api_key, language, on_log, model=stt_model)
    elif engine == "revai":
        text, lang, segs = _stt_revai(path, api_key, language, on_log, model=stt_model)
    else:
        raise ValueError(f"Unknown STT engine: {engine}")

    stt_secs = _t.time() - t0
    _log(f"Transcription done in {stt_secs:.1f}s")
    return text, lang, segs, stt_secs


# ── speech analytics ──────────────────────────────────────────────────────────

def calculate_speaker_stats_from_segments(segments: list) -> dict:
    """Compute per-speaker WPM from WhisperX diarized segments."""
    raw = {}
    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        words = len(seg["text"].split())
        duration = seg["end"] - seg["start"]
        if speaker not in raw:
            raw[speaker] = {"words": 0, "duration": 0.0}
        raw[speaker]["words"] += words
        raw[speaker]["duration"] += duration

    total_duration = sum(v["duration"] for v in raw.values()) or 1.0
    result = {}
    for speaker, data in raw.items():
        wpm = (data["words"] / data["duration"] * 60) if data["duration"] > 0 else 0.0
        result[speaker] = {
            "wpm": round(wpm, 1),
            "total_words": data["words"],
            "speaking_time": round(data["duration"], 1),
            "speaking_pct": round(data["duration"] / total_duration * 100, 1),
            "pace": (
                "Slow"      if wpm < 120 else
                "Normal"    if wpm < 150 else
                "Fast"      if wpm < 180 else
                "Very Fast"
            ),
        }
    return result


def calculate_overall_stats_from_text(text: str) -> dict:
    """Estimate overall WPM from timestamped lines like [00:01:23] text."""
    matches = list(re.finditer(
        r'\[(\d{1,2}):(\d{2})(?::(\d{2}))?\]\s*(?:\S+:\s*)?(.*)', text
    ))
    if len(matches) < 2:
        return {}

    def to_secs(m):
        h, mi, s = m.group(1), m.group(2), m.group(3) or "0"
        return int(h) * 3600 + int(mi) * 60 + int(s)

    t0 = to_secs(matches[0])
    t1 = to_secs(matches[-1])
    total_words = sum(len(m.group(4).split()) for m in matches)
    duration = t1 - t0
    if duration <= 0:
        return {}
    return {
        "overall_wpm": round(total_words / (duration / 60), 1),
        "total_words": total_words,
        "duration_secs": duration,
    }


# ── Claude prompts ────────────────────────────────────────────────────────────

STYLE_INSTRUCTIONS = {
    "formal":    "Use professional, formal language with clear section headings.",
    "casual":    "Use clear, conversational language. Keep it friendly and approachable.",
    "executive": "Be extremely concise. Summary max 3 sentences. Key points max 5 bullets. Skip minor details.",
    "bullet":    "Use bullet points for nearly everything. Minimize prose paragraphs.",
}

PANEL_SYSTEM_PROMPT = """\
You are an expert at processing panel discussion transcripts with multiple speakers.
Analyze the transcript and return ONLY a valid JSON object — no markdown, no extra text."""

STANDARD_SYSTEM_PROMPT = """\
You are an expert transcript processor.
Analyze the transcript and return ONLY a valid JSON object — no markdown, no extra text."""


def build_panel_prompt(content, fmt, num_speakers, style, speech_data, language=None, language_variant=None):
    style_note = STYLE_INSTRUCTIONS.get(style, STYLE_INSTRUCTIONS["formal"])
    speakers_hint = (
        f"Expected number of speakers: {num_speakers}"
        if num_speakers else "Number of speakers: detect automatically"
    )
    speech_section = ""
    if speech_data:
        stats_lines = "\n".join(
            f"  {spk}: {d['wpm']} WPM ({d['pace']}), {d['speaking_pct']}% of conversation"
            for spk, d in speech_data.items()
        )
        speech_section = f"\nPre-calculated speech rates:\n{stats_lines}\nUse these exact WPM values.\n"

    lang_section = ""
    if language_variant:
        lang_section = f"\nLanguage/dialect: {language_variant}. Focus accent analysis on markers specific to this regional variety (vocabulary, phonology hints, idiomatic phrases, grammar patterns).\n"
    elif language and language != "auto":
        lang_section = f"\nTranscript language: {language}. Tailor accent analysis accordingly.\n"

    return f"""\
Format: {fmt}
{speakers_hint}
Style: {style_note}
{speech_section}{lang_section}
<transcript>
{content}
</transcript>

Return JSON with exactly these keys (analysis fields first — ALWAYS include these):
{{
  "speaker_map": {{"SPEAKER_00": "name or role", ...}},
  "summary": "Executive summary",
  "key_points": ["point 1", ...],
  "action_items": [{"action": "what needs to be done", "owner": "who", "timeline": "when"}, ...],
  "speaker_profiles": {{"Name": "2-3 sentence profile of contributions"}},
  "speaker_stats": [
    {{
      "name": "resolved name",
      "words_per_minute": 142.5,
      "pace_label": "Normal",
      "speaking_percentage": 35.2,
      "accent_indicators": "Likely British English — uses 'whilst', 'cheers'. Confidence: medium.",
      "accent_confidence": "medium"
    }}
  ],
  "clean_transcript": "Full cleaned transcript with resolved speaker names and timestamps — OMIT THIS FIELD if the transcript is very long (>2000 words) to stay within token limits",
  "speaker_dialogue": "Readable dialogue with speaker labels — OMIT THIS FIELD if the transcript is very long (>2000 words) to stay within token limits"
}}

IMPORTANT: Always write the analysis fields (summary, key_points, action_items, speaker_profiles, speaker_stats) first and completely. The clean_transcript and speaker_dialogue fields are optional — skip them if including them would exceed your output limit.

For accent_indicators: analyze vocabulary, syntax, idiomatic expressions, and regional phrases.
Always state confidence level (low/medium/high)."""


def build_standard_prompt(content, fmt, style, overall_stats, language=None, language_variant=None):
    style_note = STYLE_INSTRUCTIONS.get(style, STYLE_INSTRUCTIONS["formal"])
    stats_note = ""
    if overall_stats and overall_stats.get("overall_wpm"):
        stats_note = (
            f"\nOverall speech rate: {overall_stats['overall_wpm']} WPM "
            f"({overall_stats['total_words']} words over "
            f"{overall_stats['duration_secs']//60} min)"
        )

    lang_note = ""
    if language_variant:
        lang_note = f"\nLanguage/dialect: {language_variant}. Focus accent analysis on markers specific to this regional variety (vocabulary, phonology hints, idiomatic phrases, grammar patterns)."
    elif language and language != "auto":
        lang_note = f"\nTranscript language: {language}. Tailor accent analysis accordingly."

    has_speakers = bool(re.search(r"Speaker \d+:", content))
    speaker_fields = """
  "speaker_map": {"Speaker 0": "resolved name or role", "Speaker 1": "resolved name or role"},
  "speaker_profiles": {"Resolved Name": "2-3 sentence profile of their contributions and communication style"},""" if has_speakers else ""

    return f"""\
Format: {fmt}
Style: {style_note}
{stats_note}{lang_note}

<transcript>
{content}
</transcript>

Return JSON with exactly these keys (in this order — analysis fields first, transcript last):
{{{speaker_fields}
  "summary": "Executive summary",
  "key_points": ["point 1", ...],
  "action_items": [{{"action": "what", "owner": "who", "timeline": "when"}}],
  "speaker_stats": [
    {{
      "name": "resolved speaker name or role",
      "words_per_minute": 0,
      "pace_label": "Normal",
      "speaking_percentage": 100,
      "accent_indicators": "Inferred accent from vocabulary and expressions. Confidence: low/medium/high.",
      "accent_confidence": "medium"
    }}
  ],
  "clean_transcript": "Full cleaned transcript — OMIT if transcript is very long (>2000 words)",
  "speaker_dialogue": "Readable dialogue with speaker labels — OMIT if transcript is very long (>2000 words)"
}}

IMPORTANT: Always complete the analysis fields first. clean_transcript and speaker_dialogue are optional — skip them for long transcripts to avoid hitting output limits."""


# ── processing ────────────────────────────────────────────────────────────────

def _parse_json(raw: str) -> dict:
    """Parse JSON from an AI response, recovering from truncated/malformed output."""
    if not raw or not raw.strip():
        raise json.JSONDecodeError("Empty response from AI — model returned no content", "", 0)

    s = raw.strip()

    # Gemini and some models wrap JSON in text + fences: "Here is the result:\n```json\n{...}\n```"
    # Extract the first fenced block if one exists anywhere in the response
    fence = re.search(r"```(?:json)?\s*\n?([\s\S]*?)```", s)
    if fence:
        s = fence.group(1).strip()
    else:
        # No fenced block — strip leading/trailing fences only
        s = re.sub(r"^```(?:json)?\s*", "", s)
        s = re.sub(r"\s*```$", "", s)

    # Fast path: well-formed JSON
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        pass

    # Find the outermost { ... } block (handles text before the JSON)
    start = s.find('{')
    if start == -1:
        raise json.JSONDecodeError("No JSON object found in AI response", s, 0)

    # Walk forward tracking brace/bracket depth.
    # If the response was truncated, depth will be > 0 at end-of-string.
    depth = arr = 0
    in_str = esc = False
    i = start
    while i < len(s):
        c = s[i]
        if esc:                      esc = False; i += 1; continue
        if c == '\\' and in_str:     esc = True;  i += 1; continue
        if c == '"':                 in_str = not in_str; i += 1; continue
        if not in_str:
            if   c == '{': depth += 1
            elif c == '}':
                depth -= 1
                if depth == 0 and arr == 0:
                    # Found end of outermost object
                    try:
                        return json.loads(s[start:i + 1])
                    except json.JSONDecodeError:
                        pass  # keep walking
            elif c == '[': arr += 1
            elif c == ']' and arr > 0: arr -= 1
        i += 1

    # Truncated response — append the missing closing chars and retry
    if depth > 0 or arr > 0:
        closing = ']' * arr + '}' * depth
        try:
            return json.loads(s[start:] + closing)
        except json.JSONDecodeError:
            pass

    raise json.JSONDecodeError("Cannot parse/repair JSON from AI response", s, 0)


def _chunk(text: str, max_chars: int = 600_000) -> list:
    if len(text) <= max_chars:
        return [text]
    chunks, buf, buf_len = [], [], 0
    for line in text.split("\n"):
        line_len = len(line) + 1
        if buf_len + line_len > max_chars and buf:
            chunks.append("\n".join(buf))
            buf, buf_len = [line], line_len
        else:
            buf.append(line)
            buf_len += line_len
    if buf:
        chunks.append("\n".join(buf))
    return chunks


def _merge_summaries(client, results: list) -> str:
    _safe_print("  Merging summaries...")
    combined = "\n".join(f"- {r.get('summary', '')}" for r in results)
    return client.chat(
        system="You are a helpful assistant.",
        user=f"Combine into one 3-5 sentence executive summary:\n{combined}",
        max_tokens=512,
    ) or combined


def process_transcript(
    client,
    raw_text: str,
    fmt: str,
    panel_mode: bool = False,
    num_speakers: int = None,
    config: ReportConfig = None,
    raw_whisperx: dict = None,
    language: str = None,
    language_variant: str = None,
    speaker_names: str = None,
    on_log=None,
    on_token_usage=None,   # callable(total_input, total_output)
    image_context: str = "",  # visual descriptions from uploaded images
) -> TranscriptResult:
    def _log(m):
        _safe_print(f"  {m}")
        if on_log: on_log(m)

    _tok_in  = [0]
    _tok_out = [0]
    def _on_usage(inp, out):
        _tok_in[0]  += inp
        _tok_out[0] += out
        if on_token_usage:
            on_token_usage(_tok_in[0], _tok_out[0])

    config = config or ReportConfig()
    chunks = _chunk(raw_text)
    n = len(chunks)

    audio_speech_data = None
    if raw_whisperx and raw_whisperx.get("segments"):
        audio_speech_data = calculate_speaker_stats_from_segments(raw_whisperx["segments"])

    overall_text_stats = calculate_overall_stats_from_text(raw_text) if not audio_speech_data else None

    if n > 1:
        _log(f"Transcript is large — splitting into {n} chunks for Claude…")

    results = []
    for i, chunk in enumerate(chunks, 1):
        label = f"chunk {i}/{n}" if n > 1 else "transcript"
        _log(f"Sending {label} to AI for analysis… (summary, key points, accent detection)")

        if panel_mode:
            sys_prompt = PANEL_SYSTEM_PROMPT
            prompt = build_panel_prompt(chunk, fmt, num_speakers, config.style, audio_speech_data, language, language_variant)
        else:
            sys_prompt = STANDARD_SYSTEM_PROMPT
            prompt = build_standard_prompt(chunk, fmt, config.style, overall_text_stats, language, language_variant)

        if image_context:
            prompt = (
                f"=== Visual Context (from uploaded images) ===\n{image_context}\n\n"
                f"=== Audio/Text Content ===\n\n"
            ) + prompt

        if speaker_names:
            prompt = (
                f"There are {speaker_names} in this recording.\n"
                f"Label each speaker distinctly as Speaker 1, Speaker 2, etc. when identifying who said what.\n\n"
            ) + prompt

        if n > 1:
            prompt = f"[Part {i} of {n}]\n\n" + prompt

        _parse_ok = False
        # Attempt 0: full prompt, generous tokens
        # Attempt 1: strict JSON-only system prompt
        # Attempt 2: analysis-only (no transcript echo) — for very long recordings
        _ANALYSIS_ONLY_SCHEMA = """{
  "speaker_map": {},
  "summary": "...",
  "key_points": [],
  "action_items": [],
  "speaker_profiles": {},
  "speaker_stats": [],
  "clean_transcript": "",
  "speaker_dialogue": ""
}"""
        for _parse_attempt in range(3):
            if _parse_attempt == 0:
                _sys = sys_prompt
                _max_tok = 16000
            elif _parse_attempt == 1:
                _sys = ("IMPORTANT: Respond with ONLY a raw JSON object. "
                        "No markdown, no code fences, no explanation — just the JSON.\n\n"
                        + sys_prompt)
                _max_tok = 8000
            else:
                # Final fallback: analysis fields only, skip transcript echo
                _sys = ("Return ONLY a JSON object with these keys: "
                        "speaker_map, summary, key_points, action_items, "
                        "speaker_profiles, speaker_stats. "
                        "Leave clean_transcript and speaker_dialogue as empty strings. "
                        "No markdown, no fences.")
                _max_tok = 4000
            raw = client.chat(
                system=_sys,
                user=prompt,
                max_tokens=_max_tok,
                thinking=False,
                on_usage=_on_usage,
            )
            try:
                results.append(_parse_json(raw))
                _parse_ok = True
                break
            except json.JSONDecodeError:
                if _parse_attempt == 0:
                    _log("⚠️ JSON parse failed — retrying with stricter prompt…")
                elif _parse_attempt == 1:
                    _log("⚠️ Still failed — retrying analysis-only (no transcript echo)…")
                else:
                    _log("⚠️ All 3 attempts failed — using raw response as fallback")
                    results.append({"raw": raw, "parse_error": True,
                                    "summary": raw[:2000] if raw else "AI returned no content.",
                                    "clean_transcript": "", "speaker_dialogue": ""})
        _log(f"Claude analysis complete{f' ({i}/{n})' if n > 1 else ''}.")

    r = results[0] if n == 1 else None

    # Build merged result; fall back to raw_text when the AI omitted transcript
    # fields (happens when token limit hit before reaching those keys — now placed
    # last in the JSON schema so analysis fields are always written first).
    _ct = (r.get("clean_transcript", "") if r else
           "\n\n".join(x.get("clean_transcript", "") for x in results))
    _sd = (r.get("speaker_dialogue", "") if r else
           "\n\n".join(x.get("speaker_dialogue", "") for x in results))

    merged = TranscriptResult(
        clean_transcript=_ct or raw_text,
        speaker_dialogue=_sd or raw_text,
        summary=r.get("summary", "") if r else _merge_summaries(client, results),
        key_points=r.get("key_points", []) if r else [p for x in results for p in x.get("key_points", [])],
        action_items=r.get("action_items", []) if r else [a for x in results for a in x.get("action_items", [])],
        speaker_map=(r or results[0]).get("speaker_map", {}),
        speaker_profiles=r.get("speaker_profiles", {}) if r else {k: v for x in results for k, v in x.get("speaker_profiles", {}).items()},
    )

    raw_stats = (r or results[0]).get("speaker_stats", [])
    for s in raw_stats:
        spk_key = s.get("name", "")
        if audio_speech_data:
            matched = next(
                (v for k, v in audio_speech_data.items() if k in spk_key or spk_key in k), None
            )
            if matched:
                s["words_per_minute"] = matched["wpm"]
                s["pace_label"] = matched["pace"]
                s["speaking_percentage"] = matched["speaking_pct"]

        merged.speaker_stats.append(SpeakerStats(
            name=s.get("name", ""),
            words_per_minute=s.get("words_per_minute", 0.0),
            pace_label=s.get("pace_label", ""),
            speaking_percentage=s.get("speaking_percentage", 0.0),
            accent_indicators=s.get("accent_indicators", ""),
            accent_confidence=s.get("accent_confidence", ""),
        ))

    return merged


# ── Interview Mode ────────────────────────────────────────────────────────────

_INTERVIEW_SYSTEM = """\
You are an expert interview coach and communication analyst.
Analyse the provided interview transcript and return a structured JSON object.
Be specific, honest, and actionable. Use the exact keys shown below.
When a candidate profile/resume is provided, every model_answer MUST be personalised to that
specific person — draw on their real job titles, projects, companies, skills, and experiences.
The suggested answer should sound exactly like THAT candidate speaking naturally, not a generic template.
"""

_INTERVIEW_PROMPT = """\
{profile_section}Analyse this interview transcript carefully. Return ONLY valid JSON — no markdown fences.

SCOPE — INTERVIEW ONLY:
Only analyse the live interview between the interviewer and the candidate.
Exclude entirely any post-interview debrief, internal team discussion, coaching call, or conversation that happens after the candidate leaves. These sections typically involve only the interviewers/recruiters talking among themselves (e.g. "great candidate", "let's move forward"). Do NOT include questions or statements from those sections, and do NOT let them influence the overall_score or advance_likelihood.

CRITICAL: Include EVERY question asked by the interviewer in the "questions" array — do not skip, merge, or omit any. If a follow-up or clarifying question was asked, include it as its own entry.

EXCLUDE from scoring: Do NOT include any question where the interviewer is inviting the candidate to ask their own questions — e.g. "Do you have any questions for me?", "Is there anything you'd like to ask?", "Any questions from your side?", or similar closing invitations. These are not interview questions and must not appear in the "questions" array or affect the score.

Rules for answer_said:
- Quote or closely paraphrase what the candidate ACTUALLY said — 3 to 5 sentences.
- Include the specific points, examples, numbers, stories, or projects they mentioned.
- Do NOT generalise or summarise vaguely. Capture the real substance of their words.
- If they gave no answer or deflected, say so plainly.

Rules for deflection:
- "none" = candidate answered directly and on-topic.
- "partial" = candidate answered but avoided the core of the question, gave a vague or generic response, or pivoted to a different topic without fully addressing what was asked.
- "full" = candidate completely skipped, refused, or gave a non-answer (e.g. "I'd rather not say", silence, or a wholly unrelated response).

Rules for model_answer:
- Write as if YOU are the candidate speaking right now — first-person, present tense.
- If a candidate profile was provided above, use their REAL background: name actual companies, projects, technologies, and experiences from their profile. Make it sound like them specifically.
- If no profile was provided, write a confident, believable answer that feels natural for this role.
- Natural, conversational voice — no bullets, no headers, no "I would say…". Just speak the answer.
- 3–5 sentences. Confident but not robotic. Sound like a real person, not an AI template.

{{
  "questions": [
    {{
      "id": 1,
      "question": "<exact question text from the transcript>",
      "speaker": "<interviewer name or 'Interviewer'>",
      "answer_said": "<3-5 sentences of exactly what the candidate said — specific points, examples, stories>",
      "deflection": "<none|partial|full>",
      "deflection_note": "<one sentence explaining HOW they deflected, or empty string if none>",
      "score": "<Great|Good|Needs Improvement|Missed>",
      "score_reason": "<one sentence why>",
      "model_answer": "<first-person natural answer as if you are the candidate speaking — confident, no bullets>",
      "coaching_tip": "<one specific, actionable piece of advice for this answer>"
    }}
  ],
  "overall_score": "<0-10>",
  "overall_verdict": "<Great|Good|Needs Improvement>",
  "strengths": ["<strength 1>", "<strength 2>"],
  "weaknesses": ["<weakness 1>", "<weakness 2>"],
  "prep_guide": ["<topic to study 1>", "<topic to study 2>"]
}}

--- DEEP MODE (only fill if requested) ---
  "deflection_rate": "<0-100 % of questions deflected>",
  "advance_likelihood": "<0-100 % chance of advancing>",
  "advance_reasoning": "<why this likelihood>",
  "weak_question_prep": [
    {{"question": "...", "study_topics": ["..."], "sample_answer": "..."}}
  ]
--- END DEEP MODE ---

TRANSCRIPT:
{transcript}

DEEP MODE REQUESTED: {deep_mode}
"""


def run_interview_analysis(
    transcript: str,
    client: "LLMClient",
    deep_mode: bool = False,
    on_log=None,
    candidate_profile: str = "",
) -> dict:
    def _log(m):
        _safe_print(f"  [Interview] {m}")
        if on_log: on_log(m)

    _log("Running interview analysis…")
    profile_section = (
        "CANDIDATE PROFILE — use this to personalise every model_answer to this specific person:\n"
        "---\n" + candidate_profile.strip()[:8000] + "\n---\n\n"
    ) if candidate_profile and candidate_profile.strip() else ""
    # Character budget: Claude 200K context = ~800K chars; GPT-4o 128K = ~512K chars.
    # 400K chars (~100K tokens) covers a 4-hour interview on any major provider safely
    # and leaves plenty of room for the 16K output + prompt overhead.
    _TRANSCRIPT_CHAR_LIMIT = 400_000
    prompt = _INTERVIEW_PROMPT.format(
        profile_section=profile_section,
        transcript=transcript[:_TRANSCRIPT_CHAR_LIMIT],
        deep_mode="YES" if deep_mode else "NO",
    )
    raw = client.chat(system=_INTERVIEW_SYSTEM, user=prompt, max_tokens=16000)
    _log("Interview analysis complete.")
    try:
        # Strip markdown fences if model ignores instructions
        clean = re.sub(r"```(?:json)?|```", "", raw).strip()
        return json.loads(clean)
    except Exception:
        return {"raw": raw, "parse_error": True}


# ── History ───────────────────────────────────────────────────────────────────

import time as _time_mod
import uuid as _uuid_mod


def save_history_entry(entry: dict, history_path: "Path") -> None:
    history_path.parent.mkdir(parents=True, exist_ok=True)
    with open(history_path, "a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False, default=str) + "\n")


def load_history(history_path: "Path") -> list:
    if not history_path.exists():
        return []
    entries = []
    with open(history_path, encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.strip()
            if line:
                try:
                    entries.append(json.loads(line))
                except Exception:
                    pass
    return list(reversed(entries))  # newest first


# ── combined report builder ───────────────────────────────────────────────────

def build_combined_report(result: TranscriptResult, config: ReportConfig) -> str:
    divider = "=" * 60
    thin = "-" * 40
    sections = []

    if result.detected_language:
        sections += ["DOCUMENT INFO", thin,
                     f"  Language : {result.detected_language}", ""]

    if config.include_summary:
        sections += [divider, "SUMMARY", divider, result.summary, ""]

    if config.include_key_points and result.key_points:
        sections += ["KEY POINTS", thin]
        sections += [f"  • {p}" for p in result.key_points]
        sections.append("")

    if config.include_action_items and result.action_items:
        sections += ["ACTION ITEMS", thin]
        sections += [f"  ☐  {a}" for a in result.action_items]
        sections.append("")

    if config.include_speech_analytics and result.speaker_stats:
        sections += ["SPEECH ANALYTICS", thin]
        for s in result.speaker_stats:
            wpm_str = f"{s.words_per_minute} WPM" if s.words_per_minute else "N/A"
            pct_str = f"{s.speaking_percentage}% of conversation" if s.speaking_percentage else ""
            sections.append(f"  {s.name}")
            sections.append(f"    Speech rate : {wpm_str} — {s.pace_label}  {pct_str}")
            if s.accent_indicators:
                sections.append(f"    Accent      : {s.accent_indicators}")
        sections.append("")

    if config.include_speaker_profiles and result.speaker_profiles:
        sections += ["SPEAKER PROFILES", thin]
        for name, profile in result.speaker_profiles.items():
            sections += [f"  {name}", f"  {profile}", ""]

    if config.include_transcript:
        sections += [divider, "FULL TRANSCRIPT", divider, result.clean_transcript, ""]

    # ── Interview Coaching Analysis ───────────────────────────────────────────
    ia = result.interview_analysis
    if ia and not ia.get("parse_error"):
        sections += [divider, "INTERVIEW COACHING ANALYSIS", divider]
        score   = ia.get("overall_score", "—")
        verdict = ia.get("overall_verdict", "")
        adv     = ia.get("advance_likelihood", "")
        defl    = ia.get("deflection_rate", "")
        sections.append(f"  Overall Score     : {score} / 10  —  {verdict}")
        if adv:
            sections.append(f"  Advance Likelihood: {adv}%")
        if defl:
            sections.append(f"  Deflection Rate   : {defl}%")
        sections.append("")

        qs = ia.get("questions", [])
        if qs:
            _SCORE_ICON = {
                "Great": "★", "Good": "◑",
                "Needs Improvement": "△", "Missed": "✗",
            }
            _DEFL_LABEL = {
                "partial": "⚠️  Partially deflected",
                "full":    "🚫  Did not answer",
            }
            sections.append("  QUESTION BREAKDOWN")
            sections.append("  " + thin)
            for q in qs:
                qid        = q.get("id", "")
                question   = q.get("question", "")
                sc         = q.get("score", "")
                reason     = q.get("score_reason", "")
                defl       = (q.get("deflection") or "none").lower().strip()
                defl_note  = q.get("deflection_note", "")
                said       = q.get("answer_said") or q.get("answer_summary", "")
                ideal      = q.get("model_answer") or q.get("ideal_answer", "")
                tip        = q.get("coaching_tip", "")
                icon       = _SCORE_ICON.get(sc, "·")

                sections.append(f"  Q{qid}: {question}")
                sections.append(f"    Score  : {icon} {sc}" + (f"  — {reason}" if reason else ""))
                if defl in _DEFL_LABEL:
                    sections.append(f"    {_DEFL_LABEL[defl]}")
                    if defl_note:
                        sections.append(f"    {defl_note}")
                if said:
                    sections.append(f"    What was said:")
                    for _line in said.splitlines():
                        sections.append(f"      {_line}")
                if ideal:
                    sections.append(f"    What you could have said:")
                    for _line in ideal.splitlines():
                        sections.append(f"      {_line}")
                if tip:
                    sections.append(f"    Coaching Tip: {tip}")
                sections.append("")

        adv_reason = ia.get("advance_reasoning", "")
        if adv_reason:
            sections += ["  " + divider, "  DEEP ANALYSIS", "  " + thin,
                         f"  Deflection Rate   : {defl_pct}%" if (defl_pct := ia.get("deflection_rate","")) else "",
                         f"  Advance Likelihood: {adv_pct}%" if (adv_pct := ia.get("advance_likelihood","")) else "",
                         "", f"  {adv_reason}", ""]

    return "\n".join(sections)


# ── save outputs ──────────────────────────────────────────────────────────────

def save_results(result: TranscriptResult, config: ReportConfig, output_dir: str, stem: str) -> dict:
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    paths = {}
    paths["transcript"] = str(out / f"{stem}_transcript.txt")
    paths["speakers"]   = str(out / f"{stem}_speakers.txt")
    paths["combined"]   = str(out / f"{stem}_combined.txt")
    paths["report"]     = str(out / f"{stem}_report.md")
    paths["json"]       = str(out / f"{stem}_full.json")
    paths["srt"]        = str(out / f"{stem}.srt")
    paths["vtt"]        = str(out / f"{stem}.vtt")
    paths["docx"]       = str(out / f"{stem}_report.docx")

    Path(paths["transcript"]).write_text(result.clean_transcript, encoding="utf-8")
    Path(paths["speakers"]).write_text(result.speaker_dialogue, encoding="utf-8")
    Path(paths["combined"]).write_text(build_combined_report(result, config), encoding="utf-8")

    report_lines = [f"# {stem}", ""]
    if result.detected_language:
        report_lines += [f"**Language:** {result.detected_language}", ""]
    report_lines += ["## Summary", result.summary, ""]
    if result.key_points:
        report_lines += ["## Key Points", *[f"- {p}" for p in result.key_points], ""]
    if result.action_items:
        def _fmt_ai(a):
            if isinstance(a, dict):
                action = a.get("action", a.get("item", str(a)))
                owner  = a.get("owner", "")
                tl     = a.get("timeline", "")
                line   = f"- [ ] {action}"
                if owner: line += f"  *(Owner: {owner})*"
                if tl:    line += f"  *[{tl}]*"
                return line
            return f"- [ ] {a}"
        report_lines += ["## Action Items", *[_fmt_ai(a) for a in result.action_items], ""]
    if result.speaker_stats:
        report_lines += ["## Speech Analytics", ""]
        for s in result.speaker_stats:
            report_lines += [
                f"### {s.name}",
                f"- **Speech rate:** {s.words_per_minute} WPM — {s.pace_label}",
                f"- **Speaking time:** {s.speaking_percentage}% of conversation",
                f"- **Accent:** {s.accent_indicators or 'N/A'}",
                "",
            ]
    if result.speaker_profiles:
        report_lines += ["## Speaker Profiles", ""]
        for name, profile in result.speaker_profiles.items():
            report_lines += [f"### {name}", profile, ""]

    Path(paths["report"]).write_text("\n".join(report_lines), encoding="utf-8")

    with open(paths["json"], "w", encoding="utf-8") as f:
        json.dump(result.__dict__, f, indent=2, ensure_ascii=False, default=str)

    # SRT / VTT subtitles (only if segments available)
    if result.segments:
        Path(paths["srt"]).write_text(generate_srt(result.segments), encoding="utf-8")
        Path(paths["vtt"]).write_text(generate_vtt(result.segments), encoding="utf-8")

    # DOCX report
    generate_docx(result, stem, paths["docx"])

    return paths


# ── main entry point ──────────────────────────────────────────────────────────

def run(
    file_path: str,
    file_path_2: str = None,             # optional second file — merged into one transcript
    output_dir: str = "transcript_output",
    whisper_model: str = "base",
    stt_engine: str = "whisper_local",   # see STT_ENGINES keys
    stt_api_key: str = None,             # API key for cloud STT engines
    stt_model: str = None,               # model selection for cloud STT engines
    panel_mode: bool = False,
    num_speakers: int = None,
    config: ReportConfig = None,
    api_key: str = None,
    provider: str = "anthropic",
    model: str = None,
    base_url: str = None,
    language: str = None,
    language_variant: str = None,
    speaker_names: str = None,
    interview_mode: bool = False,
    interview_deep: bool = False,
    candidate_profile: str = "",
    history_path: "Path | None" = None,
    on_whisper_progress=None,
    on_raw_transcript=None,
    on_stage_change=None,
    on_log=None,
    on_stt_done=None,               # callable(stt_secs: float) — fired when STT completes
    on_token_usage=None,            # callable(total_input, total_output) — live token counts
    cancel_event=None,              # threading.Event — set to abort before LLM analysis starts
    pre_transcribed=None,           # (raw_text, lang, segments) — skip STT when provided
    transcription_only: bool = False,  # skip AI analysis — return raw transcript immediately
    image_paths: list = None,       # optional images (slides, whiteboard) for visual context
    use_gpu: bool = True,           # use CUDA if available for Whisper
) -> TranscriptResult:
    def _log(m):
        _safe_print(f"  {m}")
        if on_log: on_log(m)

    config = config or ReportConfig()
    _safe_print(f"\nTranscript Agent {'(Panel)' if panel_mode else ''}")
    _safe_print("=" * 50)

    fname = Path(file_path).name
    ext   = Path(file_path).suffix.lower()
    _log(f"File: {fname}")
    if language:
        _log(f"Language: {language_variant or language}")

    raw_whisperx  = {}
    _detected_lang = ""
    _segments      = []
    _stt_secs      = 0.0

    if pre_transcribed is not None:
        # Resume from a cached STT result — skip transcription entirely
        raw_text, _detected_lang, _segments = pre_transcribed
        _log(f"Resuming from saved transcript (~{len(raw_text.split()):,} words) — skipping transcription")
        if on_stage_change: on_stage_change("whisper")
        if on_stt_done: on_stt_done(0.0)
        fmt = "audio/video (cached)"
    elif ext in (AUDIO_EXTS | VIDEO_EXTS):
        if panel_mode:
            _log("Mode: Panel (multi-speaker diarization)")
            if on_stage_change: on_stage_change("extracting")
            raw_text, raw_whisperx = load_audio_video_panel(
                file_path, whisper_model, num_speakers, language=language, on_log=on_log,
                use_gpu=use_gpu,
            )
            _detected_lang = raw_whisperx.get("language", "")
            fmt = "panel audio/video (diarized)"
        else:
            _model_note = f"  |  Whisper model: {whisper_model}" if stt_engine == "whisper_local" else (f"  |  Model: {stt_model}" if stt_model else "")
            _log(f"STT engine: {STT_ENGINES.get(stt_engine, stt_engine)}{_model_note}")
            raw_text, _detected_lang, _segments, _stt_secs = stt_transcribe(
                file_path, stt_engine,
                api_key=stt_api_key,
                whisper_model=whisper_model,
                language=language,
                stt_model=stt_model,
                on_progress=on_whisper_progress,
                on_stage_change=on_stage_change,
                on_log=on_log,
                use_gpu=use_gpu,
            )
            fmt = "audio/video"
        if on_stt_done:
            on_stt_done(_stt_secs)
    else:
        _log(f"Mode: Document  ({ext or 'text'})")
        raw_text, fmt = load_file(file_path, whisper_model)
        _log(f"Document loaded: ~{len(raw_text.split()):,} words")

    # ── Merge second file if provided ────────────────────────────────────────────
    if file_path_2 and Path(file_path_2).exists() and ext in (AUDIO_EXTS | VIDEO_EXTS):
        _log(f"▶ File 2: {Path(file_path_2).name} — transcribing…")
        _dur1 = _get_audio_duration(file_path) or (
            max((s.get("end", 0) for s in _segments), default=0) if _segments else 0
        )
        _raw2, _lang2, _segs2, _secs2 = stt_transcribe(
            file_path_2, stt_engine,
            api_key=stt_api_key,
            whisper_model=whisper_model,
            language=language,
            stt_model=stt_model,
            on_progress=on_whisper_progress,
            on_stage_change=on_stage_change,
            on_log=on_log,
            use_gpu=use_gpu,
        )
        # Offset file-2 segment timestamps so they follow file 1 continuously
        for seg in _segs2:
            if "start" in seg: seg["start"] = round(seg["start"] + _dur1, 3)
            if "end"   in seg: seg["end"]   = round(seg["end"]   + _dur1, 3)
        raw_text  = raw_text.rstrip() + "\n\n" + _raw2.lstrip()
        _segments = _segments + _segs2
        _stt_secs += _secs2
        _log(f"✅ Both files transcribed — {len(raw_text.split()):,} words total")

    if on_raw_transcript:
        on_raw_transcript(raw_text)

    if transcription_only:
        return TranscriptResult(
            clean_transcript=raw_text,
            speaker_dialogue=raw_text,
            detected_language=_detected_lang,
            segments=_segments,
            stt_engine=stt_engine,
            stt_seconds=_stt_secs,
        )

    if cancel_event and cancel_event.is_set():
        raise RuntimeError("Job cancelled before AI analysis")

    _log(f"Text ready: ~{len(raw_text.split()):,} words  |  Passing to AI…")

    if on_stage_change: on_stage_change("claude")
    _model = model or ("claude-opus-4-8" if provider == "anthropic" else "gpt-4o")
    client = LLMClient(provider=provider, api_key=api_key, model=_model, base_url=base_url,
                       use_gpu=use_gpu)

    image_context = ""
    if image_paths:
        valid_images = [p for p in image_paths if p and Path(p).exists()]
        if valid_images:
            _log(f"Analyzing {len(valid_images)} image(s) for visual context…")
            if on_stage_change: on_stage_change("images")
            image_context = client.describe_images(valid_images, on_log=on_log)
            _log(f"✅ Visual context ready ({len(valid_images)} image(s))")

    if speaker_names:
        _log(f"Speaker count provided: {speaker_names}")

    # Track final cumulative token counts for history
    _final_tok = [0, 0]   # [in, out]
    def _token_usage_wrapper(inp, out):
        _final_tok[0] = inp
        _final_tok[1] = out
        if on_token_usage:
            on_token_usage(inp, out)

    result = process_transcript(
        client, raw_text, fmt,
        panel_mode=panel_mode,
        num_speakers=num_speakers,
        config=config,
        raw_whisperx=raw_whisperx,
        language=language,
        language_variant=language_variant,
        speaker_names=speaker_names,
        on_log=on_log,
        on_token_usage=_token_usage_wrapper,
        image_context=image_context,
    )

    result.detected_language = language_variant or language or _detected_lang or "Auto-detected"
    result.segments    = _segments
    result.stt_engine  = stt_engine
    result.stt_seconds = _stt_secs

    # ── Interview Mode ────────────────────────────────────────────────────────
    if interview_mode:
        _log("Running Interview Mode analysis…")
        if on_stage_change: on_stage_change("interview")
        result.interview_analysis = run_interview_analysis(
            raw_text, client, deep_mode=interview_deep, on_log=on_log,
            candidate_profile=candidate_profile,
        )

    paths = save_results(result, config, output_dir, Path(file_path).stem)

    # ── Save to history ───────────────────────────────────────────────────────
    if history_path:
        ia = result.interview_analysis
        entry = {
            "id": _uuid_mod.uuid4().hex[:12],
            "timestamp": _time_mod.strftime("%Y-%m-%d %H:%M"),
            "filename": fname,
            "stt_engine": STT_ENGINES.get(stt_engine, stt_engine),
            "stt_secs": round(_stt_secs, 1),
            "ai_provider": provider,
            "ai_model": _model,
            "language": result.detected_language,
            "word_count": len(raw_text.split()),
            "overall_score": ia.get("overall_score", "") if ia else "",
            "overall_verdict": ia.get("overall_verdict", "") if ia else "",
            "tok_in":  _final_tok[0],
            "tok_out": _final_tok[1],
            "paths": paths,
            "summary": result.summary[:300],
            "interview_questions": [
                {
                    "question":    q.get("question", ""),
                    "answer_said": q.get("answer_said") or q.get("answer_summary", ""),
                    "score":       q.get("score", ""),
                    "score_reason": q.get("score_reason", ""),
                    "deflection":  q.get("deflection", "none"),
                }
                for q in ia.get("questions", [])
            ] if ia else [],
        }
        save_history_entry(entry, Path(history_path))

    _log(f"✅ Analysis complete — outputs ready for download")
    _safe_print(f"\n✓ Done! Outputs: {output_dir}/")
    return result


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Transcript Agent")
    parser.add_argument("file", help="Input file path")
    parser.add_argument("--output", default="transcript_output")
    parser.add_argument("--whisper", default="base",
                        choices=["tiny", "base", "small", "medium", "large"])
    parser.add_argument("--panel", action="store_true")
    parser.add_argument("--speakers", type=int, default=None)
    parser.add_argument("--style", default="formal",
                        choices=["formal", "casual", "executive", "bullet"])
    parser.add_argument("--images", nargs="+", metavar="IMAGE",
                        help="Supporting images (slides, whiteboard, docs) for visual context")
    args = parser.parse_args()

    run(args.file, args.output, args.whisper, args.panel, args.speakers,
        ReportConfig(style=args.style),
        image_paths=args.images or [])
