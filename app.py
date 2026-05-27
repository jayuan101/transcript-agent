#!/usr/bin/env python3
"""Transcript Agent — Gradio UI with drag-and-drop | v2.6"""

import gradio as gr
import os
import sys
import uuid
import threading
import queue as Q
import time
import re
import urllib.parse
import mimetypes
import tempfile
import zoneinfo
from pathlib import Path

try:
    from dotenv import load_dotenv
    # When frozen (.exe), load .env from the .exe's directory, not the temp extraction dir
    _env_dir = Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).parent
    load_dotenv(_env_dir / ".env")
except ImportError:
    pass

try:
    import job_db as _jdb
    _JDB_OK = True
except ImportError:
    _jdb = None
    _JDB_OK = False

try:
    import psutil as _psutil
    _PSUTIL_OK = True
except ImportError:
    _psutil = None
    _PSUTIL_OK = False

# ── version & auto-update ─────────────────────────────────────────────────────
APP_VERSION = "3.16"
_RELEASES_API = "https://api.github.com/repos/jayuan101/transcript-agent-releases/releases/latest"
_update_info: dict = {}
_update_downloaded = threading.Event()
_update_new_path: list = [None]   # [Path | None]


def _silent_download():
    """Download the update exe in the background as soon as an update is detected."""
    try:
        import urllib.request
        exe_path = Path(sys.executable).resolve()
        suffix = ".exe" if sys.platform == "win32" else ""
        new_path = exe_path.parent / f"TranscriptAgent_update{suffix}"
        urllib.request.urlretrieve(_update_info.get("url", ""), str(new_path))
        _update_new_path[0] = new_path
        _update_downloaded.set()
    except Exception:
        pass


def _is_installed_app() -> bool:
    """True only when running as the installed .exe launched by launcher.py."""
    return getattr(sys, "frozen", False) and bool(os.environ.get("TRANSCRIPT_AGENT_WINDOWED"))


def _check_for_update():
    if not _is_installed_app():
        return
    try:
        import requests as _r
        r = _r.get(_RELEASES_API, timeout=10, headers={"User-Agent": "TranscriptAgent"})
        if r.status_code != 200:
            return
        data = r.json()
        latest = data.get("tag_name", "").lstrip("v")
        if not latest or latest == APP_VERSION:
            return
        exe_url = next(
            (a["browser_download_url"] for a in data.get("assets", []) if a["name"].endswith(".exe")),
            None,
        )
        _update_info.update({
            "version": latest,
            "changelog": data.get("body", "").strip() or "See release page for details.",
            "url": exe_url,
            "assets": data.get("assets", []),
        })
        if exe_url:
            threading.Thread(target=_silent_download, daemon=True).start()
    except Exception:
        pass


threading.Thread(target=_check_for_update, daemon=True).start()


def _get_update_banner():
    current_badge = (
        f'<span style="display:inline-block;background:#0f2a1a;border:1px solid #166534;'
        f'border-radius:6px;padding:3px 10px;font-size:0.8em;color:#86efac;margin-bottom:6px;">'
        f'v{APP_VERSION}</span>'
    )
    if not _is_installed_app():
        return (
            f'<div style="padding:6px 0;">{current_badge}'
            f'<span style="font-size:0.8em;color:#64748b;margin-left:8px;">Docker / browser mode</span></div>'
        )
    if not _update_info:
        return (
            f'<div style="padding:6px 0;">{current_badge}'
            f'<span style="font-size:0.8em;color:#64748b;margin-left:8px;">Up to date</span></div>'
        )
    v = _update_info["version"]
    changelog_lines = _update_info.get("changelog", "").strip().splitlines()
    changelog_html  = "".join(
        f'<li style="margin:3px 0;">{ln.lstrip("- ").strip()}</li>'
        for ln in changelog_lines if ln.strip()
    )
    return (
        f'<div id="update-banner" style="background:#1e3a5f;border:1px solid #3b82f6;'
        f'border-radius:8px;padding:12px 16px;margin:4px 0 8px;color:#e2e8f0;">'
        f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap;margin-bottom:6px;">'
        f'{current_badge}'
        f'<span style="font-size:0.8em;color:#94a3b8;">&#8594;</span>'
        f'<span style="display:inline-block;background:#1e3a5f;border:1px solid #3b82f6;'
        f'border-radius:6px;padding:3px 10px;font-size:0.8em;color:#93c5fd;font-weight:600;">'
        f'v{v} available</span>'
        f'</div>'
        f'<details style="font-size:0.85em;color:#cbd5e1;">'
        f'<summary style="cursor:pointer;color:#93c5fd;user-select:none;">What\'s new in v{v}</summary>'
        f'<ul style="margin:8px 0 0 16px;padding:0;line-height:1.6;">{changelog_html}</ul>'
        f'</details>'
        f'</div>'
    )


def _do_update():
    if not _is_installed_app():
        yield "⚠️ Auto-update only works in the installed .exe — not the source version."
        return

    if not _update_info:
        yield "No update available."
        return

    import urllib.request, subprocess
    exe_path = Path(sys.executable).resolve()
    is_win   = sys.platform == "win32"

    # ── Acquire the downloaded file (or download now with progress) ───────────
    if _update_downloaded.is_set() and _update_new_path[0] and _update_new_path[0].exists():
        new_path = _update_new_path[0]
        yield "📦 Update already downloaded — restarting in 3 seconds..."
        time.sleep(3)
    else:
        is_mac = sys.platform == "darwin"
        def _asset_match(name):
            if is_mac:
                return "mac" in name.lower() or (not name.endswith(".exe") and "." not in name)
            return name.endswith(".exe")
        url = next(
            (a["browser_download_url"] for a in _update_info.get("assets", [])
             if _asset_match(a["name"])), None
        )
        if not url:
            yield "❌ No download URL found. Visit the releases page to update manually."
            return

        new_path = exe_path.parent / ("TranscriptAgent_update.exe" if is_win else "TranscriptAgent_update")
        _progress = {"pct": 0}

        def _hook(count, block, total):
            if total > 0:
                _progress["pct"] = min(100, int(count * block * 100 / total))

        _dl_done = threading.Event()

        def _dl():
            try:
                urllib.request.urlretrieve(url, str(new_path), reporthook=_hook)
            finally:
                _dl_done.set()

        yield "⬇️ Downloading update — 0%..."
        threading.Thread(target=_dl, daemon=True).start()
        while not _dl_done.is_set():
            time.sleep(0.5)
            yield f"⬇️ Downloading — {_progress['pct']}%..."

        if not new_path.exists():
            yield "❌ Download failed. Check your connection and try again."
            return

        yield "📦 Download complete — restarting in 3 seconds..."
        time.sleep(3)

    # ── Apply update ──────────────────────────────────────────────────────────
    try:
        if is_win:
            # Rename the running exe (allowed on Windows), move new one in, relaunch.
            old_name = exe_path.stem + "_old" + exe_path.suffix
            script = exe_path.parent / "_ta_update.bat"
            script.write_text(
                "@echo off\n"
                "timeout /t 2 /nobreak >nul\n"
                f'ren "{exe_path}" "{old_name}"\n'
                f'move /y "{new_path}" "{exe_path}"\n'
                f'start "" "{exe_path}"\n'
                "ping -n 6 127.0.0.1 >nul\n"
                f'del /f "{exe_path.parent}\\{old_name}" 2>nul\n'
                'del /f "%~f0"\n',
                encoding="utf-8",
            )
            subprocess.Popen(
                str(script),
                shell=True,
                creationflags=subprocess.CREATE_NO_WINDOW,
                cwd=str(exe_path.parent),
            )
        else:
            script = exe_path.parent / "_ta_update.sh"
            script.write_text(
                "#!/bin/bash\nsleep 2\n"
                f'mv "{exe_path}" "{exe_path}_old"\n'
                f'mv "{new_path}" "{exe_path}"\n'
                f'chmod +x "{exe_path}"\n'
                f'"{exe_path}" &\n'
                "sleep 5\n"
                f'rm -f "{exe_path}_old"\n'
                'rm -- "$0"\n'
            )
            script.chmod(0o755)
            subprocess.Popen(["bash", str(script)])
    except Exception as e:
        yield f"❌ Install failed: {e}"
        return

    sys.exit(0)


def _get_proxyscrape_proxies() -> list:
    """Return proxy URLs to try when a direct download is blocked (403)."""
    import requests as _r
    api_key = os.environ.get("PROXYSCRAPE_API_KEY", "").strip()
    params = {
        "request": "displayproxies",
        "proxy_format": "protocolipport",
        "format": "text",
        "anonymity": "elite,anonymous",
        "timeout": "8000",
        "country": "US,GB,DE,CA,AU,NL",
    }
    if api_key:
        params["apiKey"] = api_key
    try:
        r = _r.get(
            "https://api.proxyscrape.com/v3/free-proxy-list/get",
            params=params,
            timeout=10,
        )
        if r.ok:
            lines = [l.strip() for l in r.text.splitlines() if l.strip()]
            return ["http://" + l if "://" not in l else l for l in lines[:20]]
    except Exception:
        pass
    return []


def _resolve_nextcloud_token_url(url: str):
    """
    Decode a Nextcloud share token URL and return a direct S3 object URL.
    Supports nextim.itcapp.ai and path-style s3.amazonaws.com share URLs.
    """
    import base64, json
    try:
        parsed = urllib.parse.urlparse(url)
        if "/api/nextcloud/share" not in parsed.path:
            return None
        qs = urllib.parse.parse_qs(parsed.query)
        token_b64 = qs.get("token", [""])[0]
        if not token_b64:
            return None
        pad = (4 - len(token_b64) % 4) % 4
        token_data = json.loads(base64.b64decode(token_b64 + "=" * pad).decode())
        file_key = token_data.get("key", "")
        if not file_key:
            return None
        host = parsed.hostname or ""
        if host == "s3.amazonaws.com":
            # path-style: /s3.amazonaws.com/{bucket}/api/nextcloud/share
            bucket = parsed.path.lstrip("/").split("/")[0]
        elif "itcapp.ai" in host or "nextim" in host:
            # ITC App Nextcloud server — recordings live in this S3 bucket
            bucket = "nextcloud-talk-recordings-itc-eit"
        else:
            return None
        if not bucket:
            return None
        encoded_key = urllib.parse.quote(file_key, safe="")
        return f"https://{bucket}.s3.amazonaws.com/{encoded_key}"
    except Exception:
        return None


def _download_url(url: str, dest_dir: Path, on_progress=None) -> Path:
    """Download a URL to dest_dir; uses Content-Disposition to pick filename.

    Handles S3 application-level redirects (PermanentRedirect XML response with
    HTTP 200/301) by retrying against the correct endpoint.
    On 403 Forbidden, retries via ProxyScrape proxy (uses PROXYSCRAPE_API_KEY env var).
    on_progress(bytes_received, total_bytes_or_0): called periodically during download.
    """
    import requests

    def _do_get(u: str, proxies=None, timeout=300):
        return requests.get(
            u,
            stream=True,
            timeout=timeout,
            headers={"User-Agent": "TranscriptAgent/1.0"},
            allow_redirects=True,
            proxies=proxies,
        )

    resp = _do_get(url)

    if resp.status_code == 401:
        raise ValueError(
            "The URL requires a login (401 Unauthorized). "
            "Download the file manually and paste its local path instead."
        )

    if resp.status_code == 403:
        resp.close()
        resp = None

        # Strategy 1: decode Nextcloud share token → direct S3 object URL
        direct_url = _resolve_nextcloud_token_url(url)
        if direct_url:
            try:
                _r = _do_get(direct_url)
                if _r.ok:
                    resp = _r
                else:
                    _r.close()
            except Exception:
                pass

        # Strategy 2: retry via ProxyScrape proxy
        if resp is None:
            for proxy_url in _get_proxyscrape_proxies():
                try:
                    pr = {"http": proxy_url, "https": proxy_url}
                    _r = _do_get(url, proxies=pr, timeout=90)
                    if _r.ok:
                        resp = _r
                        break
                    _r.close()
                except Exception:
                    continue

        if resp is None:
            raise ValueError(
                "403 Forbidden — the server denied access to this URL.\n"
                "The share link may require a company VPN or internal network access.\n"
                "Download the file manually and paste its local path instead."
            )

    resp.raise_for_status()

    # Peek at the first chunk so we can detect S3 XML-style errors that arrive
    # as a normal 200 response.
    chunks = resp.iter_content(chunk_size=65536)
    try:
        first = next(chunks)
    except StopIteration:
        first = b""

    ct = (resp.headers.get("Content-Type", "") or "").split(";")[0].strip().lower()
    head = first[:512].lstrip()

    looks_xml = (
        ct in ("application/xml", "text/xml")
        or head.startswith(b"<?xml")
        or head.startswith(b"<Error")
    )

    if looks_xml:
        body = head.decode("utf-8", errors="replace")
        # Detect S3 PermanentRedirect and retry against the suggested endpoint.
        m_code   = re.search(r"<Code>([^<]+)</Code>", body)
        m_ep     = re.search(r"<Endpoint>([^<]+)</Endpoint>", body)
        m_bucket = re.search(r"<Bucket>([^<]+)</Bucket>", body)
        if m_code and m_code.group(1) == "PermanentRedirect" and m_ep:
            # Use resp.url (the actual S3 URL after following HTTP redirects) not
            # the original URL, so we get the correct object path.
            parsed = urllib.parse.urlparse(resp.url)
            new_host = m_ep.group(1).strip()
            bucket = m_bucket.group(1).strip() if m_bucket else ""
            _qs = parsed.query
            _is_presigned = any(k in _qs for k in ("X-Amz-Algorithm", "X-Amz-Credential", "AWSAccessKeyId"))
            if bucket and not new_host.startswith(bucket + "."):
                new_path = "/" + bucket.strip("/") + "/" + parsed.path.lstrip("/")
            else:
                new_path = parsed.path
            # Try with presigned params preserved first — the credential region
            # often already matches the redirect target (global → regional endpoint).
            new_url = urllib.parse.urlunparse(
                (parsed.scheme or "https", new_host, new_path,
                 parsed.params, _qs, parsed.fragment)
            )
            resp.close()
            resp = _do_get(new_url)
            if resp.status_code in (400, 403) and _is_presigned:
                # Params didn't work at the new endpoint — truly wrong region.
                # Try without presigned params as a last resort (public bucket).
                new_url_no_auth = urllib.parse.urlunparse(
                    (parsed.scheme or "https", new_host, new_path,
                     parsed.params, "", parsed.fragment)
                )
                resp.close()
                resp = _do_get(new_url_no_auth)
            if resp.status_code in (400, 403):
                _orig_host = urllib.parse.urlparse(url).hostname or ""
                _hint = (
                    f"Open {_orig_host} in your browser, download the recording, "
                    "then drag the file into the app or paste its local path."
                ) if _orig_host else "Download the file manually and paste its local path."
                raise ValueError(
                    f"The recording server ({_orig_host}) generated a presigned URL "
                    "for the wrong AWS region — this is a server configuration issue "
                    "that cannot be worked around automatically.\n\n" + _hint
                )
            resp.raise_for_status()
            chunks = resp.iter_content(chunk_size=65536)
            try:
                first = next(chunks)
            except StopIteration:
                first = b""
            ct = (resp.headers.get("Content-Type", "") or "").split(";")[0].strip().lower()
            head = first[:512].lstrip()
            if head.startswith(b"<?xml") or head.startswith(b"<Error"):
                snippet = head.decode("utf-8", errors="replace")[:400]
                raise ValueError(
                    "Download failed: server returned an XML error after "
                    "following S3 redirect.\n" + snippet
                )
        else:
            snippet = body[:400]
            msg = "Download failed: server returned an XML error instead of a media file."
            if m_code:
                msg += f"\nS3 error: {m_code.group(1)}"
            raise ValueError(msg + "\n" + snippet)

    # Prefer Content-Disposition filename; fall back to final (post-redirect) URL path
    cd = resp.headers.get("Content-Disposition", "")
    filename = None
    if cd:
        m = re.search(r"filename\*=(?:UTF-8'')?([^\s;]+)", cd, re.I)
        if m:
            filename = urllib.parse.unquote(m.group(1).strip('"'))
        else:
            m = re.search(r'filename=["\']?([^"\';]+)', cd, re.I)
            if m:
                filename = m.group(1).strip()
    if not filename:
        url_path = urllib.parse.urlparse(resp.url).path   # use final URL after redirects
        filename = Path(urllib.parse.unquote(url_path)).name or "download"
    if not Path(filename).suffix:
        filename += mimetypes.guess_extension(ct) or ""

    dest = dest_dir / filename
    total = 0
    last_progress = time.time()
    last_report   = time.time()
    total_size    = int(resp.headers.get("Content-Length", 0) or 0)
    with open(dest, "wb") as f:
        if first:
            f.write(first)
            total += len(first)
            last_progress = time.time()
        for chunk in chunks:
            if not chunk:
                if time.time() - last_progress > 60:
                    raise ValueError(
                        "Download stalled — no data received for 60 seconds. "
                        "The server dropped the connection. Download the file via "
                        "your browser and drag it into the app instead."
                    )
                continue
            f.write(chunk)
            total += len(chunk)
            last_progress = time.time()
            if on_progress and time.time() - last_report >= 2:
                on_progress(total, total_size)
                last_report = time.time()

    # Sanity check: a real audio/video file is never just a few hundred bytes.
    if total < 1024:
        try:
            preview = dest.read_bytes()[:400].decode("utf-8", errors="replace")
        except Exception:
            preview = ""
        try:
            dest.unlink()
        except Exception:
            pass
        raise ValueError(
            f"Download failed: only {total} bytes received — the URL did not "
            f"return a valid media file.\n{preview}"
        )

    return dest

# ── Windows sleep prevention ──────────────────────────────────────────────────
_ES_CONTINUOUS      = 0x80000000
_ES_SYSTEM_REQUIRED = 0x00000001   # blocks idle auto-sleep; screen can still turn off
_sleep_active       = False
_sleep_thread       = None

def _set_lid_action(action: int):
    """Set lid-close power action: 0=do nothing, 1=sleep. Requires no elevation on most systems."""
    if sys.platform != "win32":
        return
    import subprocess
    for idx in ("setacvalueindex", "setdcvalueindex"):
        subprocess.run(
            ["powercfg", f"/{idx}", "SCHEME_CURRENT", "SUB_BUTTONS", "LIDACTION", str(action)],
            capture_output=True, check=False,
        )
    subprocess.run(["powercfg", "/setactive", "SCHEME_CURRENT"], capture_output=True, check=False)

def _prevent_sleep():
    """Block idle sleep on Windows (Win32 API) and macOS (caffeinate)."""
    global _sleep_active, _sleep_thread
    _sleep_active = True
    if sys.platform == "win32":
        import ctypes
        ctypes.windll.kernel32.SetThreadExecutionState(_ES_CONTINUOUS | _ES_SYSTEM_REQUIRED)
        _set_lid_action(0)
        if _sleep_thread is None or not _sleep_thread.is_alive():
            def _refresh():
                import ctypes as _ct
                while _sleep_active:
                    _ct.windll.kernel32.SetThreadExecutionState(_ES_CONTINUOUS | _ES_SYSTEM_REQUIRED)
                    time.sleep(60)
            _sleep_thread = threading.Thread(target=_refresh, daemon=True)
            _sleep_thread.start()
    elif sys.platform == "darwin":
        if _sleep_thread is None or not _sleep_thread.is_alive():
            def _caffeinate():
                import subprocess as _sp
                while _sleep_active:
                    proc = _sp.Popen(["caffeinate", "-i", "-t", "3600"])
                    while _sleep_active and proc.poll() is None:
                        time.sleep(5)
                    proc.kill()
            _sleep_thread = threading.Thread(target=_caffeinate, daemon=True)
            _sleep_thread.start()

    elif sys.platform.startswith("linux"):
        if _sleep_thread is None or not _sleep_thread.is_alive():
            def _inhibit():
                import subprocess as _sp
                while _sleep_active:
                    proc = _sp.Popen(
                        ["systemd-inhibit", "--what=sleep:idle",
                         "--who=TranscriptAgent", "--why=Transcription in progress",
                         "--mode=block", "sleep", "3600"],
                        stdout=_sp.DEVNULL, stderr=_sp.DEVNULL,
                    )
                    while _sleep_active and proc.poll() is None:
                        time.sleep(5)
                    proc.kill()
            _sleep_thread = threading.Thread(target=_inhibit, daemon=True)
            _sleep_thread.start()


def _allow_sleep():
    global _sleep_active
    _sleep_active = False
    if sys.platform == "win32":
        import ctypes
        ctypes.windll.kernel32.SetThreadExecutionState(_ES_CONTINUOUS)
        _set_lid_action(1)
    # macOS + Linux: background thread exits naturally when _sleep_active turns False

from transcript_agent import (
    run, ReportConfig, build_combined_report, LLMClient,
    AUDIO_EXTS, VIDEO_EXTS,
)

# ── AI provider configuration ─────────────────────────────────────────────────
_PROVIDERS = {
    "Claude (Anthropic)": {
        "type": "anthropic",
        "placeholder": "sk-ant-api03-…",
        "info": "console.anthropic.com → API keys → Create key",
        "models": [
            "claude-opus-4-7",
            "claude-sonnet-4-6",
            "claude-haiku-4-5-20251001",
            "claude-3-5-sonnet-20241022",
            "claude-3-5-haiku-20241022",
            "claude-3-opus-20240229",
        ],
        "base_url": None,
    },
    "OpenAI": {
        "type": "openai",
        "placeholder": "sk-…",
        "info": "platform.openai.com → API keys",
        "models": [
            "gpt-4.1",
            "gpt-4.1-mini",
            "gpt-4o",
            "gpt-4o-mini",
            "o3",
            "o3-mini",
            "o1",
            "gpt-4-turbo",
            "gpt-3.5-turbo",
        ],
        "base_url": None,
    },
    "Google Gemini": {
        "type": "openai_compat",
        "placeholder": "AIzaSy…",
        "info": "aistudio.google.com → Get API key",
        "models": [
            "gemini-2.5-pro",
            "gemini-2.5-flash",
            "gemini-2.5-pro-preview-05-06",
            "gemini-2.0-flash-exp",
            "gemini-2.0-flash-thinking-exp",
            "gemini-1.5-pro",
            "gemini-1.5-flash",
            "gemini-1.5-flash-8b",
        ],
        "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
    },
    "Groq": {
        "type": "openai_compat",
        "placeholder": "gsk_…",
        "info": "console.groq.com → API keys",
        "models": [
            "llama-3.3-70b-versatile",
            "llama-3.1-70b-versatile",
            "llama3-70b-8192",
            "llama3-8b-8192",
            "deepseek-r1-distill-llama-70b",
            "mixtral-8x7b-32768",
            "gemma2-9b-it",
        ],
        "base_url": "https://api.groq.com/openai/v1",
    },
    "Mistral": {
        "type": "openai_compat",
        "placeholder": "…",
        "info": "console.mistral.ai → API keys",
        "models": [
            "mistral-large-latest",
            "mistral-medium-latest",
            "mistral-small-latest",
            "open-mixtral-8x22b",
            "open-mistral-nemo",
        ],
        "base_url": "https://api.mistral.ai/v1",
    },
    "Together AI": {
        "type": "openai_compat",
        "placeholder": "…",
        "info": "api.together.ai → Settings → API keys",
        "models": [
            "meta-llama/Llama-3-70b-chat-hf",
            "meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
            "mistralai/Mixtral-8x22B-Instruct-v0.1",
            "Qwen/Qwen2-72B-Instruct",
            "google/gemma-2-27b-it",
        ],
        "base_url": "https://api.together.xyz/v1",
    },
    "Perplexity": {
        "type": "openai_compat",
        "placeholder": "pplx-…",
        "info": "perplexity.ai → Settings → API",
        "models": [
            "llama-3.1-sonar-large-128k-online",
            "llama-3.1-sonar-huge-128k-online",
            "llama-3.1-sonar-small-128k-online",
            "llama-3.1-70b-instruct",
        ],
        "base_url": "https://api.perplexity.ai",
    },
    "Ollama (Local)": {
        "type": "openai_compat",
        "placeholder": "none required",
        "info": "ollama.ai — run models locally, no API key needed",
        "models": [
            "llama3.2",
            "llama3.1:70b",
            "mistral",
            "gemma2:27b",
            "qwen2.5:72b",
            "phi3.5",
        ],
        "base_url": "http://localhost:11434/v1",
    },
    "Custom (OpenAI-compatible)": {
        "type": "openai_compat",
        "placeholder": "sk-… or leave blank",
        "info": "Enter your API base URL in the field below",
        "models": ["custom-model"],
        "base_url": "",
    },
}

# ── Speech-to-text (transcription) providers ─────────────────────────────────
_STT_PROVIDERS = {
    "Whisper (Local)": {
        "id": "whisper",
        "key_placeholder": None,
        "key_info": None,
        "models": ["tiny", "base", "small", "medium", "large"],
        "default_model": "base",
        "model_info": "tiny = fastest   |   large = most accurate",
    },
    "Deepgram (Cloud)": {
        "id": "deepgram",
        "key_placeholder": "dg-…",
        "key_info": "Get key: console.deepgram.com → API Keys",
        "models": [
            "nova-2", "nova-2-general", "nova-2-meeting", "nova-2-phonecall",
            "nova-2-voicemail", "nova", "enhanced", "base",
            "whisper-large", "whisper-medium", "whisper-small",
        ],
        "default_model": "nova-2",
        "model_info": "nova-2 = best accuracy   |   nova = fast & accurate   |   enhanced/base = cheaper",
    },
    "AssemblyAI (Cloud)": {
        "id": "assemblyai",
        "key_placeholder": "your_assemblyai_key",
        "key_info": "Get key: app.assemblyai.com → API Keys",
        "models": ["best", "nano"],
        "default_model": "best",
        "model_info": "best = highest accuracy with speaker labels   |   nano = faster & cheaper",
    },
    "Groq Whisper (Cloud)": {
        "id": "groq_whisper",
        "key_placeholder": "gsk_…",
        "key_info": "Get key: console.groq.com → API Keys",
        "models": ["whisper-large-v3-turbo", "whisper-large-v3", "distil-whisper-large-v3-en"],
        "default_model": "whisper-large-v3-turbo",
        "model_info": "whisper-large-v3-turbo = fastest   |   whisper-large-v3 = most accurate",
    },
    "OpenAI Whisper API (Cloud)": {
        "id": "openai_whisper",
        "key_placeholder": "sk-…",
        "key_info": "Get key: platform.openai.com → API Keys",
        "models": ["whisper-1", "gpt-4o-transcribe", "gpt-4o-mini-transcribe"],
        "default_model": "whisper-1",
        "model_info": "whisper-1 = standard   |   gpt-4o-transcribe = highest accuracy (newer)",
    },
    "Google Cloud STT (Cloud)": {
        "id": "google_stt",
        "key_placeholder": "AIza…",
        "key_info": "Get key: console.cloud.google.com → APIs & Services → Credentials",
        "models": ["latest_long", "latest_short", "command_and_search", "phone_call"],
        "default_model": "latest_long",
        "model_info": "latest_long = best for recordings   |   latest_short = best for short clips",
    },
    "ElevenLabs (Cloud)": {
        "id": "elevenlabs",
        "key_placeholder": "sk_…",
        "key_info": "Get key: elevenlabs.io → Profile → API Keys",
        "models": ["scribe_v1"],
        "default_model": "scribe_v1",
        "model_info": "scribe_v1 = highest accuracy, 32 languages, speaker diarization",
    },
    "Rev.ai (Cloud)": {
        "id": "rev_ai",
        "key_placeholder": "your_rev_ai_token",
        "key_info": "Get key: rev.ai → Dashboard → Access Tokens",
        "models": ["machine", "fusion"],
        "default_model": "machine",
        "model_info": "machine = fastest   |   fusion = highest accuracy (slower, premium)",
    },
}

# When frozen (.exe), respect TRANSCRIPT_OUTPUT_DIR set by launcher.py so outputs
# land in ~/TranscriptAgent/outputs rather than the temp extraction directory
_out_override = os.environ.get("TRANSCRIPT_OUTPUT_DIR", "")
OUT_DIR = Path(_out_override) if _out_override else Path(__file__).parent / "outputs"
OUT_DIR.mkdir(exist_ok=True)

JOB_STATUS_FILE = OUT_DIR / ".job_status.json"

if _JDB_OK:
    try:
        _jdb.init_db(OUT_DIR)
    except Exception:
        _JDB_OK = False

def _write_job_status(status: str, **kwargs):
    import json, datetime
    data = {"status": status, "updated": datetime.datetime.now().isoformat(), **kwargs}
    try:
        JOB_STATUS_FILE.write_text(json.dumps(data, indent=2))
    except Exception:
        pass

def _read_job_status():
    import json
    try:
        return json.loads(JOB_STATUS_FILE.read_text())
    except Exception:
        return None

def load_last_result():
    """Load the most recently completed job and return its results for all output tabs."""
    import json
    status = _read_job_status()
    if not status:
        return [gr.update()] * 17 + [gr.update(value="No completed job found. Run a transcription first.", visible=True)]
    s = status.get("status")
    if s == "running":
        name = status.get("stem", "Unknown file")
        return [gr.update()] * 17 + [gr.update(
            value=f"⏳ **Transcription still in progress** — {name}. Keep this tab open and results will load automatically when done.",
            visible=True,
        )]
    if s not in ("done", "error"):
        return [gr.update()] * 17 + [gr.update(value="No completed job found. Run a transcription first.", visible=True)]
    try:
        job_dir = Path(status["job_dir"])
        stem    = status["stem"]
        data    = status.get("result", {})   # formatted strings saved at completion
        f_t    = str(job_dir / f"{stem}_transcript.txt")
        f_s    = str(job_dir / f"{stem}_speakers.txt")
        f_r    = str(job_dir / f"{stem}_report.md")
        f_c    = str(job_dir / f"{stem}_combined.txt")
        f_j    = str(job_dir / f"{stem}_full.json")
        f_p    = str(job_dir / f"{stem}_report.pdf") if (job_dir / f"{stem}_report.pdf").exists() else None
        f_docx = str(job_dir / f"{stem}_report.docx") if (job_dir / f"{stem}_report.docx").exists() else None
        f_srt  = str(job_dir / f"{stem}_transcript.srt") if (job_dir / f"{stem}_transcript.srt").exists() else None
        f_vtt  = str(job_dir / f"{stem}_transcript.vtt") if (job_dir / f"{stem}_transcript.vtt").exists() else None
        if s == "done":
            completed = status.get("completed", "")[:16].replace("T", " ")
            msg = f"✅ Loaded: **{stem}** (completed {completed})"
        else:
            err = status.get("error", "unknown error")
            msg = f"⚠️ Last job failed — partial results for **{stem}**: {err}"
        return [
            data.get("summary", ""),
            data.get("transcript", ""),
            data.get("dialogue", ""),
            data.get("profiles", ""),
            data.get("interview", ""),
            data.get("analytics", ""),
            data.get("combined", ""),
            gr.update(value=f_t if Path(f_t).exists() else None, visible=Path(f_t).exists()),
            gr.update(value=f_s if Path(f_s).exists() else None, visible=Path(f_s).exists()),
            gr.update(value=f_r if Path(f_r).exists() else None, visible=Path(f_r).exists()),
            gr.update(value=f_c if Path(f_c).exists() else None, visible=Path(f_c).exists()),
            gr.update(value=f_j if Path(f_j).exists() else None, visible=Path(f_j).exists()),
            gr.update(value=f_p, visible=f_p is not None),
            gr.update(value=f_docx, visible=f_docx is not None),
            gr.update(value=f_srt, visible=f_srt is not None),
            gr.update(value=f_vtt, visible=f_vtt is not None),
            gr.update(open=True),
            gr.update(value=msg, visible=True),
        ]
    except Exception as e:
        return [gr.update()] * 17 + [gr.update(value=f"Error loading results: {e}", visible=True)]

def get_job_banner():
    """Return HTML banner showing current job status — called on page load."""
    status = _read_job_status()
    if not status:
        return ""
    s = status.get("status")
    name = status.get("stem", "Unknown file")
    updated = status.get("updated", "")[:16].replace("T", " ")
    if s == "running":
        import datetime as _dtt
        try:
            started = _dtt.datetime.fromisoformat(status.get("updated", ""))
            age_hours = ((_dtt.datetime.now() - started).total_seconds()) / 3600
        except Exception:
            age_hours = 0
        if age_hours > 4:
            return (
                '<div id="job-banner" style="background:#422006;border:1px solid #f97316;border-radius:8px;'
                'padding:12px 16px;margin-bottom:12px;display:flex;align-items:center;gap:10px;color:#fed7aa!important;">'
                '<span style="font-size:1.3em">⚠️</span>'
                '<div><strong style="color:#fdba74!important;">Previous job did not complete</strong>'
                f'<div style="color:#fed7aa!important;font-size:0.85em">{name} — started {updated} but never finished '
                '(browser disconnected or computer slept). Click <strong>Load Last Result</strong> to check '
                'if any results were saved, or run a new job.</div>'
                '</div></div>'
            )
        return (
            '<div id="job-banner" style="background:#1e3a5f;border:1px solid #3b82f6;border-radius:8px;'
            'padding:12px 16px;margin-bottom:12px;display:flex;align-items:center;gap:10px;color:#cbd5e1!important;">'
            '<span style="font-size:1.3em">⏳</span>'
            '<div><strong style="color:#93c5fd!important;">Transcription in progress</strong>'
            f'<div style="color:#cbd5e1!important;font-size:0.85em">{name} — started {updated}. '
            'Keep this tab open to see results, or come back later and click <strong>Load Last Result</strong>.</div>'
            '</div></div>'
        )
    elif s == "done":
        completed = status.get("completed", "")[:16].replace("T", " ")
        return (
            '<div id="job-banner" style="background:#14532d;border:1px solid #4ade80;border-radius:8px;'
            'padding:12px 16px;margin-bottom:12px;display:flex;align-items:center;gap:10px;color:#bbf7d0!important;">'
            '<span style="font-size:1.3em">✅</span>'
            '<div><strong style="color:#4ade80!important;">Last transcription completed</strong>'
            f'<div style="color:#bbf7d0!important;font-size:0.85em">{name} — finished {completed}. '
            'Click <strong>Load Last Result</strong> to view it.</div>'
            '</div></div>'
        )
    elif s == "error":
        return (
            '<div id="job-banner" style="background:#450a0a;border:1px solid #f87171;border-radius:8px;'
            'padding:12px 16px;margin-bottom:12px;display:flex;align-items:center;gap:10px;color:#fca5a5!important;">'
            '<span style="font-size:1.3em">🚨</span>'
            '<div><strong style="color:#f87171!important;">Last transcription failed</strong>'
            f'<div style="color:#fca5a5!important;font-size:0.85em">{name} — {updated}. '
            f'Error: {status.get("error","unknown")}</div>'
            '</div></div>'
        )
    return ""



SUPPORTED = list(AUDIO_EXTS | VIDEO_EXTS | {".srt", ".vtt", ".txt", ".md", ".docx", ".pdf"})


# ── History helpers ───────────────────────────────────────────────────────────

_STATUS_PILL = {
    "done":         ('<span style="background:#14532d;color:#4ade80;border-radius:5px;'
                     'padding:2px 8px;font-size:0.78em;font-weight:700;">✓ Done</span>'),
    "error":        ('<span style="background:#450a0a;color:#f87171;border-radius:5px;'
                     'padding:2px 8px;font-size:0.78em;font-weight:700;">✗ Error</span>'),
    "running":      ('<span style="background:#1e3a5f;color:#60a5fa;border-radius:5px;'
                     'padding:2px 8px;font-size:0.78em;font-weight:700;">⏳ Running</span>'),
    "whisper_done": ('<span style="background:#1e3a5f;color:#93c5fd;border-radius:5px;'
                     'padding:2px 8px;font-size:0.78em;font-weight:700;">🎤 Whisper done</span>'),
    "pending":      ('<span style="background:#374151;color:#9ca3af;border-radius:5px;'
                     'padding:2px 8px;font-size:0.78em;font-weight:700;">• Pending</span>'),
}


def _build_history_html() -> str:
    import html as _html
    if not _JDB_OK:
        return "<p style='color:#9ca3af;font-size:0.85em;'>Job history unavailable (job_db not loaded).</p>"
    try:
        jobs = _jdb.list_jobs(limit=50)
    except Exception as e:
        return f"<p style='color:#f87171;font-size:0.85em;'>Error reading history: {_html.escape(str(e))}</p>"
    if not jobs:
        return "<p style='color:#9ca3af;font-size:0.85em;'>No jobs yet — run a transcription to see history here.</p>"
    rows = []
    for j in jobs:
        ts   = (j.get("created_at") or "")[:16].replace("T", " ")
        pill = _STATUS_PILL.get(j.get("status", ""), _STATUS_PILL["pending"])
        name = _html.escape(j.get("stem") or j.get("original_filename") or "—")
        jid  = _html.escape(j.get("job_id", ""))
        err  = (f'<br><span style="color:#f87171;font-size:0.76em;">'
                f'{_html.escape((j["error"] or "")[:80])}</span>') if j.get("error") else ""
        load_btn = (
            f'<button onclick="taLoadJob(\'{jid}\')" '
            f'style="background:#312e81;color:#a5b4fc;border:1px solid #4338ca;border-radius:6px;'
            f'padding:3px 10px;font-size:0.78em;cursor:pointer;white-space:nowrap;">'
            f'📂 Load</button>'
        )
        rows.append(
            f'<tr style="border-bottom:1px solid #1e293b;">'
            f'<td style="padding:6px 10px;color:#94a3b8;font-size:0.82em;">{ts}</td>'
            f'<td style="padding:6px 10px;color:#e2e8f0;max-width:200px;overflow:hidden;'
            f'text-overflow:ellipsis;white-space:nowrap;" title="{name}">{name}{err}</td>'
            f'<td style="padding:6px 10px;">{pill}</td>'
            f'<td style="padding:6px 10px;">{load_btn}</td>'
            f'</tr>'
        )
    body = "".join(rows)
    return (
        '<div style="overflow-x:auto;">'
        '<table style="width:100%;border-collapse:collapse;font-size:0.85em;">'
        '<thead><tr style="border-bottom:1px solid #334155;">'
        '<th style="padding:6px 10px;text-align:left;color:#94a3b8;font-weight:600;">Started</th>'
        '<th style="padding:6px 10px;text-align:left;color:#94a3b8;font-weight:600;">File</th>'
        '<th style="padding:6px 10px;text-align:left;color:#94a3b8;font-weight:600;">Status</th>'
        '<th style="padding:6px 10px;text-align:left;color:#94a3b8;font-weight:600;"></th>'
        '</tr></thead>'
        f'<tbody>{body}</tbody>'
        '</table></div>'
    )


def load_job_from_history(job_id_input: str = ""):
    """Load a specific job from the DB by job_id."""
    no_change = [gr.update()] * 17 + [gr.update(visible=False)]
    if not _JDB_OK:
        return no_change[:-1] + [gr.update(value="Job history not available.", visible=True)]
    jid = (job_id_input or "").strip()
    if not jid:
        return no_change[:-1] + [gr.update(value="Enter a Job ID from the table above.", visible=True)]
    try:
        job = _jdb.get_job(jid)
    except Exception as e:
        return no_change[:-1] + [gr.update(value=f"Error: {e}", visible=True)]
    if not job:
        return no_change[:-1] + [gr.update(value=f"Job `{jid}` not found.", visible=True)]
    if job.get("status") not in ("done", "error"):
        s = job.get("status", "unknown")
        return no_change[:-1] + [gr.update(value=f"Job `{jid}` status is **{s}** — only completed jobs can be loaded.", visible=True)]
    job_dir = Path(job.get("job_dir", ""))
    stem    = job.get("stem", jid)
    f_t    = str(job_dir / f"{stem}_transcript.txt")
    f_s    = str(job_dir / f"{stem}_speakers.txt")
    f_r    = str(job_dir / f"{stem}_report.md")
    f_c    = str(job_dir / f"{stem}_combined.txt")
    f_j    = str(job_dir / f"{stem}_full.json")
    f_p    = str(job_dir / f"{stem}_report.pdf") if (job_dir / f"{stem}_report.pdf").exists() else None
    f_docx = str(job_dir / f"{stem}_report.docx") if (job_dir / f"{stem}_report.docx").exists() else None
    f_srt  = str(job_dir / f"{stem}_transcript.srt") if (job_dir / f"{stem}_transcript.srt").exists() else None
    f_vtt  = str(job_dir / f"{stem}_transcript.vtt") if (job_dir / f"{stem}_transcript.vtt").exists() else None
    ts  = (job.get("created_at") or "")[:16].replace("T", " ")
    msg = f"✅ Loaded job `{jid}` — **{stem}** ({ts})"
    if job.get("status") == "error":
        msg = f"⚠️ Loaded failed job `{jid}` — partial results for **{stem}**: {job.get('error','')}"
    return [
        job.get("result_summary", ""),
        job.get("result_transcript", ""),
        job.get("result_dialogue", ""),
        job.get("result_profiles", ""),
        job.get("result_interview", ""),
        job.get("result_analytics", ""),
        job.get("result_combined", ""),
        gr.update(value=f_t if Path(f_t).exists() else None, visible=Path(f_t).exists()),
        gr.update(value=f_s if Path(f_s).exists() else None, visible=Path(f_s).exists()),
        gr.update(value=f_r if Path(f_r).exists() else None, visible=Path(f_r).exists()),
        gr.update(value=f_c if Path(f_c).exists() else None, visible=Path(f_c).exists()),
        gr.update(value=f_j if Path(f_j).exists() else None, visible=Path(f_j).exists()),
        gr.update(value=f_p, visible=f_p is not None),
        gr.update(value=f_docx, visible=f_docx is not None),
        gr.update(value=f_srt, visible=f_srt is not None),
        gr.update(value=f_vtt, visible=f_vtt is not None),
        gr.update(open=True),
        gr.update(value=msg, visible=True),
    ]

FORMATS_MD = """
**Accepted formats**
🎵 `.mp3` `.wav` `.m4a` `.flac` `.ogg` `.aac`
🎬 `.mp4` `.mov` `.avi` `.mkv` `.webm`
📝 `.srt` `.vtt`   📄 `.pdf` `.docx` `.txt` `.md`
"""

CSS = """
footer { display: none !important; }

/* Hide Gradio HTML blocks that have no content (prevents ghost oval shapes) */
[data-testid="html"]:not(:has(*)),
[data-testid="html"] > .prose:empty {
    display: none !important;
}

/* ── Design tokens ───────────────────────────────────────────────────────── */
:root {
  --ta-primary:        #6366f1;
  --ta-primary-dk:     #4f46e5;
  --ta-primary-glow:   rgba(99,102,241,0.28);
  --ta-accent:         #8b5cf6;
  --ta-radius-xl:      24px;
  --ta-radius-lg:      18px;
  --ta-radius-md:      12px;
  --ta-radius-sm:      8px;
  --ta-shadow-sm:      0 1px 3px rgba(99,102,241,0.07), 0 1px 2px rgba(0,0,0,0.04);
  --ta-shadow-md:      0 4px 16px rgba(99,102,241,0.10), 0 1px 4px rgba(0,0,0,0.06);
  --ta-shadow-lg:      0 8px 32px rgba(99,102,241,0.14), 0 2px 8px rgba(0,0,0,0.08);
  --ta-glass-bg:       rgba(255,255,255,0.82);
  --ta-glass-border:   rgba(255,255,255,0.65);
}

/* ── Page — mesh gradient background ─────────────────────────────────────── */
body {
  background:
    radial-gradient(ellipse 90% 60% at 15% -5%,  rgba(99,102,241,0.10) 0%, transparent 55%),
    radial-gradient(ellipse 70% 50% at 85% 105%, rgba(139,92,246,0.09) 0%, transparent 55%),
    radial-gradient(ellipse 50% 40% at 50%  50%, rgba(99,102,241,0.04) 0%, transparent 70%),
    #f1f1f9 !important;
  background-attachment: fixed !important;
}

/* ── Checkbox — fully custom so both checked and unchecked are visible ── */
input[type="checkbox"] {
    -webkit-appearance: none !important;
    appearance: none !important;
    width: 18px !important;
    height: 18px !important;
    min-width: 18px !important;
    border: 2px solid var(--ta-primary) !important;
    border-radius: 5px !important;
    background: #ffffff !important;
    cursor: pointer !important;
    position: relative !important;
    vertical-align: middle !important;
    transition: background 0.15s, border-color 0.15s, box-shadow 0.15s !important;
    flex-shrink: 0 !important;
}
input[type="checkbox"]:checked {
    background: var(--ta-primary) !important;
    border-color: var(--ta-primary) !important;
}
input[type="checkbox"]:checked::after {
    content: '' !important;
    position: absolute !important;
    left: 4px !important;
    top: 1px !important;
    width: 6px !important;
    height: 10px !important;
    border: 2px solid #fff !important;
    border-top: none !important;
    border-left: none !important;
    transform: rotate(45deg) !important;
    display: block !important;
}
input[type="checkbox"]:hover { border-color: var(--ta-primary-dk) !important; box-shadow: 0 0 0 3px var(--ta-primary-glow) !important; }
input[type="checkbox"]:focus { outline: 2px solid var(--ta-primary) !important; outline-offset: 2px !important; }

/* Dark mode checkboxes */
html.dark input[type="checkbox"] {
    background: #1a1a2e !important;
    border-color: #818cf8 !important;
}
html.dark input[type="checkbox"]:checked {
    background: #6366f1 !important;
    border-color: #6366f1 !important;
}
html.dark input[type="checkbox"]:hover { border-color: #a5b4fc !important; }

.checkbox-wrap { align-items: center !important; gap: 8px !important; }
.checkbox-group label, .checkbox-wrap label {
    font-size: 0.9em !important;
    font-weight: 500 !important;
    cursor: pointer !important;
    display: flex !important;
    align-items: center !important;
    gap: 8px !important;
}
html.dark .checkbox-group label span,
html.dark .checkbox-wrap label span { color: #e2e8f0 !important; }

/* ── Process / CTA button ────────────────────────────────────────────────── */
@keyframes ta-btn-shimmer {
  0%   { background-position: 200% center; }
  100% { background-position: -200% center; }
}
.big-btn button {
    background: linear-gradient(135deg, #059669 0%, #10b981 40%, #34d399 70%, #10b981 100%) !important;
    background-size: 300% 100% !important;
    color: #fff !important;
    font-size: 1.06em !important;
    font-weight: 800 !important;
    letter-spacing: 0.05em !important;
    text-transform: uppercase !important;
    border: none !important;
    border-radius: var(--ta-radius-lg) !important;
    padding: 16px 24px !important;
    min-height: 58px !important;
    width: 100% !important;
    box-shadow: 0 4px 20px rgba(16,185,129,0.50), 0 1px 4px rgba(0,0,0,0.12), inset 0 1px 0 rgba(255,255,255,0.2) !important;
    transition: all 0.25s ease !important;
    position: relative !important;
    overflow: hidden !important;
}
.big-btn button:hover {
    background-position: right center !important;
    box-shadow: 0 10px 36px rgba(16,185,129,0.65), 0 3px 10px rgba(0,0,0,0.15) !important;
    transform: translateY(-3px) scale(1.01) !important;
}
.big-btn button:active {
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 14px rgba(16,185,129,0.45) !important;
}

/* ── Upload zone ─────────────────────────────────────────────────────────── */
@keyframes ta-upload-pulse {
  0%, 100% { border-color: rgba(99,102,241,0.35); box-shadow: 0 0 0 0 rgba(99,102,241,0); }
  50%       { border-color: rgba(99,102,241,0.75); box-shadow: 0 0 0 6px rgba(99,102,241,0.07); }
}
@keyframes ta-upload-float {
  0%, 100% { transform: translateY(0); }
  50%       { transform: translateY(-4px); }
}
.upload-container {
    border-radius: var(--ta-radius-xl) !important;
    border: 2px dashed rgba(99,102,241,0.4) !important;
    background: rgba(238,242,255,0.6) !important;
    transition: all 0.25s ease !important;
    min-height: 140px !important;
}
.upload-container:not(:hover) { animation: ta-upload-pulse 3s ease-in-out infinite !important; }
.upload-container:hover {
    border: 2px solid var(--ta-primary) !important;
    background: rgba(238,242,255,0.9) !important;
    box-shadow: 0 0 0 4px var(--ta-primary-glow), var(--ta-shadow-md) !important;
    transform: scale(1.01) !important;
}
.upload-container svg, .upload-container .upload-icon {
    animation: ta-upload-float 3s ease-in-out infinite !important;
}
/* Upload text */
.upload-container .file-name, .upload-container span {
    font-weight: 600 !important;
    color: #6366f1 !important;
}

/* ── Pill-style tabs ──────────────────────────────────────────────────────── */
.tabs > .tab-nav {
    background: rgba(99,102,241,0.07) !important;
    border-radius: 40px !important;
    padding: 4px !important;
    gap: 2px !important;
    border-bottom: none !important;
}
.tabs > .tab-nav button {
    border-radius: 30px !important;
    padding: 6px 16px !important;
    border: none !important;
    font-weight: 500 !important;
    font-size: 0.87em !important;
    transition: all 0.18s ease !important;
    color: #6b7280 !important;
    background: transparent !important;
}
.tabs > .tab-nav button.selected {
    background: var(--ta-primary) !important;
    color: #fff !important;
    box-shadow: 0 2px 8px rgba(99,102,241,0.35) !important;
    border-bottom: none !important;
}
.tabs > .tab-nav button:hover:not(.selected) {
    background: rgba(99,102,241,0.12) !important;
    color: var(--ta-primary) !important;
}

/* ── Accordion headers — clean modern style ──────────────────────────────── */
.accordion > .label-wrap, details > summary {
    border-left: 3px solid var(--ta-primary) !important;
    padding-left: 12px !important;
    border-radius: 0 var(--ta-radius-sm) var(--ta-radius-sm) 0 !important;
    font-weight: 700 !important;
    letter-spacing: 0.01em !important;
}
.accordion, details {
    border-radius: var(--ta-radius-lg) !important;
    overflow: hidden !important;
}

/* ── Secondary / ghost buttons ───────────────────────────────────────────── */
button:not(.big-btn button):not(#ta-btn-light):not(#ta-btn-dark) {
    border-radius: var(--ta-radius-md) !important;
    font-weight: 600 !important;
    transition: background 0.18s, box-shadow 0.18s, transform 0.15s !important;
}
button:not(.big-btn button):not(#ta-btn-light):not(#ta-btn-dark):hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 3px 10px rgba(99,102,241,0.18) !important;
}

/* ── Radio buttons (Transcription Engine toggle) ─────────────────────────── */
.radio-group label {
    border-radius: var(--ta-radius-md) !important;
    transition: all 0.18s ease !important;
}
.radio-group label:hover {
    background: rgba(99,102,241,0.06) !important;
}
input[type="radio"] {
    accent-color: var(--ta-primary) !important;
}

/* ── Blocks — elevated cards (no backdrop-filter: avoids stacking context
     that traps Gradio dropdown listboxes) ──────────────────────────────── */
.block, .form, .panel-full-width {
    background: rgba(255,255,255,0.94) !important;
    border: 1px solid rgba(225,225,240,0.80) !important;
    border-radius: var(--ta-radius-lg) !important;
    box-shadow: var(--ta-shadow-md) !important;
    transition: box-shadow 0.25s ease !important;
}
.block:hover, .form:hover {
    box-shadow: var(--ta-shadow-lg) !important;
}

/* ── Block labels — modern small-caps style ──────────────────────────────── */
.block > .label-wrap > span, .block-label {
    font-size: 0.72em !important;
    font-weight: 800 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    color: var(--ta-primary) !important;
}

/* ── Inputs — clean with indigo focus ───────────────────────────────────── */
input[type="text"], input[type="password"], textarea, select {
    border-radius: var(--ta-radius-md) !important;
    border: 1.5px solid #e4e4f0 !important;
    background: rgba(255,255,255,0.9) !important;
    transition: border-color 0.18s, box-shadow 0.18s, background 0.18s !important;
    font-size: 0.92em !important;
}
input[type="text"]:focus, input[type="password"]:focus, textarea:focus {
    border-color: var(--ta-primary) !important;
    background: #fff !important;
    box-shadow: 0 0 0 3px var(--ta-primary-glow), 0 1px 4px rgba(99,102,241,0.12) !important;
    outline: none !important;
}

/* Info text under inputs */
.info {
    font-size: 0.74em !important;
    color: #9090a8 !important;
    letter-spacing: 0.01em !important;
}

/* ── Scrollbar ───────────────────────────────────────────────────────────── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #c4c4d0; border-radius: 6px; }
::-webkit-scrollbar-thumb:hover { background: #9494a8; }

/* scrollable dropdowns — generic */
[role="listbox"] {
    max-height: 220px !important;
    overflow-y: auto !important;
}

/* provider + model dropdowns — taller scroll area, thin scrollbar */
#provider-sel [role="listbox"],
#model-sel [role="listbox"] {
    max-height: 280px !important;
    overflow-y: auto !important;
    scrollbar-width: thin !important;
}
#provider-sel [role="listbox"]::-webkit-scrollbar,
#model-sel [role="listbox"]::-webkit-scrollbar { width: 6px !important; }
#provider-sel [role="listbox"]::-webkit-scrollbar-thumb,
#model-sel [role="listbox"]::-webkit-scrollbar-thumb {
    background: #94a3b8 !important; border-radius: 4px !important;
}
html.dark #provider-sel [role="listbox"]::-webkit-scrollbar-thumb,
html.dark #model-sel [role="listbox"]::-webkit-scrollbar-thumb {
    background: #475569 !important;
}

/* Hero banner always uses white text — dark gradient background in both modes */
#ta-hero, #ta-hero * { color: #fff !important; }
#ta-hero .subtitle { color: #93c5fd !important; }

/* log panel CSS vars */
:root {
    --ta-log-bg: #f8f9fc;
    --ta-log-border: #e5e7eb;
    --ta-log-ts: #9ca3af;
    --ta-log-hdr: #374151;
}
html.dark {
    --ta-log-bg: #0d0d18;
    --ta-log-border: #1e1e30;
    --ta-log-ts: #4a4a60;
    --ta-log-hdr: #c4c4e0;
}

/* Fix banner text — Gradio overrides <strong> color to white */
#api-banner strong, #api-banner b { color: inherit !important; font-weight: 700; }
#api-banner-sub { color: #92400e !important; }

/* ── Dark mode static rules (JS-injected sheet wins by cascade order) ── */
html.dark { color-scheme: dark; color: #f0f0ff !important; }
html.dark body {
    background:
      radial-gradient(ellipse 90% 60% at 15% -5%,  rgba(99,102,241,0.14) 0%, transparent 55%),
      radial-gradient(ellipse 70% 50% at 85% 105%, rgba(139,92,246,0.12) 0%, transparent 55%),
      #0a0a12 !important;
    background-attachment: fixed !important;
}
html.dark .gradio-container, html.dark .main, html.dark .contain {
    background: transparent !important; color: #f0f0ff !important;
}
html.dark .block, html.dark .form, html.dark .panel-full-width, html.dark .compact,
html.dark .wrap {
    background: #1e1e2a !important;
    border-color: rgba(99,102,241,0.18) !important;
}
html.dark .upload-container {
    background: rgba(30,30,50,0.70) !important;
    border-color: rgba(99,102,241,0.35) !important;
}
html.dark .block > .label-wrap > span, html.dark .block-label {
    color: #a5b4fc !important;
}
html.dark input, html.dark textarea, html.dark select {
    background: #14141e !important; color: #f0f0ff !important; border-color: #2e2e42 !important;
}
html.dark span, html.dark p, html.dark div, html.dark h1, html.dark h2,
html.dark h3, html.dark h4, html.dark li, html.dark td { color: #f0f0ff !important; }
html.dark .label-wrap span, html.dark .block-label, html.dark label span,
html.dark .info, html.dark .file-name { color: #8888a8 !important; }
html.dark .tabs > .tab-nav {
    background: rgba(99,102,241,0.12) !important; border-bottom: none !important;
}
html.dark .tabs > .tab-nav button {
    color: #8888a8 !important; background: transparent !important; border-color: transparent !important;
}
html.dark .tabs > .tab-nav button.selected {
    color: #fff !important; background: #6366f1 !important; border-bottom: none !important;
    box-shadow: 0 2px 8px rgba(99,102,241,0.4) !important;
}
html.dark .tabitem { background: #0a0a12 !important; }
html.dark .prose, html.dark .markdown { color: #f0f0ff !important; }
html.dark .prose *, html.dark .markdown * { color: #f0f0ff !important; }
html.dark [role="listbox"] { background: #1e1e2a !important; border-color: #2e2e42 !important; }
html.dark [role="option"] { color: #f0f0ff !important; background: #1e1e2a !important; }
html.dark [role="option"]:hover, html.dark [role="option"][aria-selected="true"] {
    background: #2a2a3a !important; color: #fff !important;
}
html.dark .accordion, html.dark details { background: #1e1e2a !important; border-color: #2e2e42 !important; }
html.dark .accordion .label-wrap, html.dark details summary { color: #f0f0ff !important; }
html.dark .accordion > .label-wrap, html.dark details > summary {
    border-left-color: #818cf8 !important;
}
html.dark .checkbox-group label span, html.dark .radio-group label span { color: #f0f0ff !important; }
html.dark .file-preview { background: #1e1e2a !important; color: #f0f0ff !important; }
html.dark .dropdown-arrow svg { fill: #8888a8 !important; }
html.dark button { background: #1e1e2a !important; border-color: #2e2e42 !important; color: #f0f0ff !important; }
html.dark button.selected { background: #2a2a3a !important; }
html.dark .big-btn button { background: linear-gradient(135deg,#047857,#10b981,#34d399) !important; color: #fff !important; border: none !important; }
html.dark #ta-btn-light { background: transparent !important; color: #8888a8 !important; }

/* ── Download buttons ── */
.dl-btn button {
    border-radius: 20px !important;
    font-size: 0.8em !important;
    font-weight: 600 !important;
    padding: 6px 14px !important;
    border-width: 1.5px !important;
    white-space: nowrap !important;
    min-width: 80px !important;
}
.dl-pdf button   { border-color: #ef4444 !important; color: #ef4444 !important; }
.dl-docx button  { border-color: #3b82f6 !important; color: #3b82f6 !important; }
.dl-md button    { border-color: #8b5cf6 !important; color: #8b5cf6 !important; }
.dl-txt button   { border-color: #6b7280 !important; color: #6b7280 !important; }
.dl-srt button   { border-color: #f59e0b !important; color: #f59e0b !important; }
.dl-vtt button   { border-color: #10b981 !important; color: #10b981 !important; }
.dl-json button  { border-color: #06b6d4 !important; color: #06b6d4 !important; }
.dl-pdf button:hover   { background: #ef4444 !important; color: #fff !important; }
.dl-docx button:hover  { background: #3b82f6 !important; color: #fff !important; }
.dl-md button:hover    { background: #8b5cf6 !important; color: #fff !important; }
.dl-txt button:hover   { background: #6b7280 !important; color: #fff !important; }
.dl-srt button:hover   { background: #f59e0b !important; color: #fff !important; }
.dl-vtt button:hover   { background: #10b981 !important; color: #fff !important; }
.dl-json button:hover  { background: #06b6d4 !important; color: #fff !important; }
html.dark #ta-btn-dark  { background: #6366f1 !important; color: #fff !important; }
html.dark ::-webkit-scrollbar-track { background: #0a0a12 !important; }
html.dark ::-webkit-scrollbar-thumb { background: #2e2e42 !important; border-radius: 6px !important; }
html.dark ::-webkit-scrollbar-thumb:hover { background: #4a4a60 !important; }

/* ── Adaptive CSS variables used by step-tracker and ETA panel ── */
:root {
    --ta-card-bg:          #ffffff;
    --ta-card-border:      #e5e7eb;
    --ta-card-text:        #111827;
    --ta-card-sub:         #6b7280;
    --ta-card-val:         #111827;
    --ta-step-done-bg:     #ecfdf5;
    --ta-step-done-bdr:    #10b981;
    --ta-step-done-clr:    #065f46;
    --ta-step-act-bg:      #eef2ff;
    --ta-step-act-bdr:     #6366f1;
    --ta-step-act-clr:     #4338ca;
    --ta-step-wait-bg:     #f9fafb;
    --ta-step-wait-bdr:    #e5e7eb;
    --ta-step-wait-clr:    #9ca3af;
    --ta-conn-line-done:   #10b981;
    --ta-conn-line-wait:   #e5e7eb;
    --ta-stat-bg:          rgba(255,255,255,0.90);
    --ta-stat-label:       #4338ca;
    --ta-stat-val:         #6366f1;
}
html.dark {
    --ta-card-bg:          #1e1e2a;
    --ta-card-border:      #2e2e42;
    --ta-card-text:        #f0f0ff;
    --ta-card-sub:         #8888a8;
    --ta-card-val:         #f0f0ff;
    --ta-step-done-bg:     #052e16;
    --ta-step-done-bdr:    #10b981;
    --ta-step-done-clr:    #6ee7b7;
    --ta-step-act-bg:      #1e1b4b;
    --ta-step-act-bdr:     #818cf8;
    --ta-step-act-clr:     #c7d2fe;
    --ta-step-wait-bg:     #14141e;
    --ta-step-wait-bdr:    #2e2e42;
    --ta-step-wait-clr:    #6b6b88;
    --ta-conn-line-done:   #10b981;
    --ta-conn-line-wait:   #2e2e42;
    --ta-stat-bg:          rgba(20,20,30,0.90);
    --ta-stat-label:       #a5b4fc;
    --ta-stat-val:         #f0f0ff;
}

"""

_SB = (
    "background:#ffffff;border:2px solid #6366f1;border-radius:14px;"
    "padding:16px 20px;font-size:1.05em;font-family:sans-serif;"
    "min-height:60px;box-shadow:0 4px 16px rgba(99,102,241,0.18);"
)

_ANIM = (
    "<style>"
    "@keyframes pgslide{0%{left:-45%}100%{left:110%}}"
    "</style>"
)

def _fmt_eta(eta_secs: int) -> str:
    secs = max(0, int(eta_secs))
    if secs < 10:
        return "almost done!"
    if secs < 60:
        return f"{secs} seconds left"
    m, s = divmod(secs, 60)
    if m == 1 and s == 0:
        return "1 minute left"
    if s == 0:
        return f"{m} minutes left"
    if m == 1:
        return f"1 minute {s} sec left"
    return f"{m} minutes {s} sec left"


def _finish_time_str(eta_secs: int, tz_name: str = "") -> str:
    from datetime import datetime, timezone, timedelta
    finish_utc = datetime.now(timezone.utc) + timedelta(seconds=eta_secs)
    if tz_name:
        try:
            finish = finish_utc.astimezone(zoneinfo.ZoneInfo(tz_name))
        except Exception:
            finish = finish_utc
    else:
        finish = finish_utc
    hour = finish.hour % 12 or 12
    ampm = "AM" if finish.hour < 12 else "PM"
    return f"{hour}:{finish.minute:02d} {ampm}"


def _status_compact(icon: str, title: str, elapsed: str = "") -> str:
    """Minimal one-line status — used when eta_panel carries the detail."""
    elap = (f'<span style="color:var(--ta-card-sub);font-size:.85em;margin-left:10px;">'
            f'elapsed: {elapsed}</span>') if elapsed else ""
    return (f'<div style="background:var(--ta-card-bg);border:2px solid #6366f1;border-radius:14px;'
            f'padding:16px 20px;font-size:1.05em;font-family:sans-serif;min-height:60px;'
            f'box-shadow:0 4px 16px rgba(99,102,241,0.18);">'
            f'<div style="color:var(--ta-card-text);font-weight:700;font-size:1em;">'
            f'{icon} {title}{elap}</div></div>')


def _status_html(icon: str, title: str, subtitle: str = "", elapsed: str = "",
                 pct: float = None, eta_secs: int = None, tz_name: str = "") -> str:
    elap = (
        f'<span style="color:#6b7280;font-size:.88em;margin-left:12px;font-weight:400;">'
        f'elapsed: {elapsed}</span>'
    ) if elapsed else ""

    sub = (
        f'<div style="color:#374151;font-size:.93em;margin-top:5px;">{subtitle}</div>'
    ) if subtitle else ""

    eta_html = ""
    if eta_secs is not None and eta_secs > 0:
        finish_str = _finish_time_str(eta_secs, tz_name)
        eta_html = (
            f'<div style="margin-top:8px;display:flex;gap:8px;flex-wrap:wrap;">'
            f'<span style="background:#eef2ff;border-radius:20px;padding:4px 14px;'
            f'color:#4338ca;font-weight:700;font-size:0.85em;">⏱ {_fmt_eta(eta_secs)}</span>'
            f'<span style="background:#ecfdf5;border-radius:20px;padding:4px 14px;'
            f'color:#065f46;font-weight:700;font-size:0.85em;">🕐 Done by {finish_str}</span>'
            f'</div>'
        )

    if pct is not None:
        fill = f"{pct*100:.0f}%"
        bar_html = (
            f'<div style="margin-top:10px;background:#eef2ff;border-radius:8px;height:6px;overflow:hidden;">'
            f'<div style="width:{fill};height:100%;background:linear-gradient(90deg,#6366f1,#8b5cf6);border-radius:8px;'
            f'transition:width 0.6s ease;box-shadow:0 0 8px rgba(99,102,241,0.5);"></div></div>'
            f'<div style="color:#4338ca;font-weight:700;font-size:.88em;margin-top:5px;">{fill} complete</div>'
        )
    else:
        bar_html = (
            f'{_ANIM}'
            f'<div style="margin-top:10px;background:#eef2ff;border-radius:8px;height:6px;'
            f'overflow:hidden;position:relative;">'
            f'<div style="position:absolute;width:45%;height:100%;'
            f'background:linear-gradient(90deg,#6366f1,#8b5cf6);'
            f'border-radius:8px;animation:pgslide 1.4s ease-in-out infinite;'
            f'box-shadow:0 0 8px rgba(99,102,241,0.5);"></div></div>'
        )

    return (
        f'<div style="{_SB}">'
        f'<div style="color:#111827;font-weight:700;font-size:1.1em;">{icon} {title}{elap}</div>'
        f'{sub}{eta_html}{bar_html}'
        f'</div>'
    )

PACE_EMOJI = {"Slow": "🐢", "Normal": "🚶", "Fast": "🏃", "Very Fast": "⚡"}

# ── language options ──────────────────────────────────────────────────────────

LANGUAGES = [
    ("🔍 Auto-detect",                "auto"),
    ("🇺🇸 English",                   "en"),
    ("🇪🇸 Spanish",                   "es"),
    ("🇫🇷 French",                    "fr"),
    ("🇧🇷 Portuguese",                "pt"),
    ("🇩🇪 German",                    "de"),
    ("🇮🇹 Italian",                   "it"),
    ("🇯🇵 Japanese",                  "ja"),
    ("🇨🇳 Chinese (Mandarin)",         "zh"),
    ("🇰🇷 Korean",                    "ko"),
    ("🇸🇦 Arabic",                    "ar"),
    ("🇷🇺 Russian",                   "ru"),
    # ── Indian languages ─────────────────────────────────────────────────────
    ("🇮🇳 Hindi",                     "hi"),
    ("🇮🇳 Bengali",                   "bn"),
    ("🇮🇳 Tamil",                     "ta"),
    ("🇮🇳 Telugu",                    "te"),
    ("🇮🇳 Gujarati",                  "gu"),
    ("🇮🇳 Kannada",                   "kn"),
    ("🇮🇳 Malayalam",                 "ml"),
    ("🇮🇳 Marathi",                   "mr"),
    ("🇮🇳 Punjabi",                   "pa"),
    ("🇮🇳 Urdu",                      "ur"),
    # ── Other languages ───────────────────────────────────────────────────────
    ("🇳🇱 Dutch",                     "nl"),
    ("🇵🇱 Polish",                    "pl"),
    ("🇹🇷 Turkish",                   "tr"),
    ("🇻🇳 Vietnamese",                "vi"),
    ("🇹🇭 Thai",                      "th"),
    ("🇸🇪 Swedish",                   "sv"),
    ("🇳🇴 Norwegian",                 "no"),
    ("🇩🇰 Danish",                    "da"),
    ("🇬🇷 Greek",                     "el"),
    ("🇮🇱 Hebrew",                    "he"),
    ("🇷🇴 Romanian",                  "ro"),
    ("🇭🇺 Hungarian",                 "hu"),
    ("🇨🇿 Czech",                     "cs"),
    ("🇫🇮 Finnish",                   "fi"),
    ("🇺🇦 Ukrainian",                 "uk"),
    ("🇮🇩 Indonesian",                "id"),
]

# ── per-language regional variants ────────────────────────────────────────────
# First entry in each list is the "auto / no hint" value; these are filtered out
# before being passed to the backend so only meaningful hints are forwarded.

LANGUAGE_VARIANTS = {
    "en": [
        ("🌍 Auto (General English)",          "General English"),
        ("🇺🇸 American English",               "American English (en-US)"),
        ("🇬🇧 British English",                "British English (en-GB)"),
        ("🇦🇺 Australian English",             "Australian English (en-AU)"),
        ("🇨🇦 Canadian English",               "Canadian English (en-CA)"),
        ("🇮🇪 Irish English",                  "Irish English (en-IE)"),
        ("🏴 Scottish English",                "Scottish English"),
        ("🇿🇦 South African English",          "South African English (en-ZA)"),
        ("🇮🇳 Indian English",                 "Indian English (en-IN)"),
        ("🇸🇬 Singaporean English",            "Singaporean English (en-SG)"),
        ("🇵🇭 Filipino English",               "Filipino English (en-PH)"),
        ("🌍 Nigerian English",                "Nigerian English"),
        ("🇯🇲 Caribbean English",              "Caribbean English"),
    ],
    "es": [
        ("🌎 Auto (General Spanish)",          "General Spanish"),
        ("🇨🇴 Colombian Spanish",              "Colombian Spanish (es-CO)"),
        ("🇲🇽 Mexican Spanish",                "Mexican Spanish (es-MX)"),
        ("🇪🇸 Castilian Spanish (Spain)",      "Castilian Spanish (es-ES)"),
        ("🇦🇷 Argentinian Spanish",            "Argentinian Spanish (es-AR)"),
        ("🇨🇱 Chilean Spanish",                "Chilean Spanish (es-CL)"),
        ("🇻🇪 Venezuelan Spanish",             "Venezuelan Spanish (es-VE)"),
        ("🇵🇪 Peruvian Spanish",               "Peruvian Spanish (es-PE)"),
        ("🇨🇺 Cuban Spanish",                  "Cuban Spanish (es-CU)"),
        ("🇵🇷 Puerto Rican Spanish",           "Puerto Rican Spanish (es-PR)"),
        ("🇩🇴 Dominican Spanish",              "Dominican Spanish (es-DO)"),
        ("🇬🇹 Guatemalan Spanish",             "Guatemalan Spanish (es-GT)"),
        ("🇪🇨 Ecuadorian Spanish",             "Ecuadorian Spanish (es-EC)"),
        ("🇧🇴 Bolivian Spanish",               "Bolivian Spanish (es-BO)"),
        ("🇺🇾 Uruguayan Spanish",              "Uruguayan Spanish (es-UY)"),
        ("🇵🇾 Paraguayan Spanish",             "Paraguayan Spanish (es-PY)"),
        ("🇭🇳 Honduran Spanish",               "Honduran Spanish (es-HN)"),
        ("🇸🇻 Salvadoran Spanish",             "Salvadoran Spanish (es-SV)"),
        ("🇳🇮 Nicaraguan Spanish",             "Nicaraguan Spanish (es-NI)"),
        ("🇨🇷 Costa Rican Spanish",            "Costa Rican Spanish (es-CR)"),
        ("🇵🇦 Panamanian Spanish",             "Panamanian Spanish (es-PA)"),
        ("🇺🇸 US Latino Spanish",              "US Latino Spanish (es-US)"),
    ],
    "fr": [
        ("🌍 Auto (General French)",           "General French"),
        ("🇫🇷 France French",                  "France French (fr-FR)"),
        ("🇨🇦 Canadian French (Québécois)",    "Canadian French (fr-CA)"),
        ("🇧🇪 Belgian French",                 "Belgian French (fr-BE)"),
        ("🇨🇭 Swiss French",                   "Swiss French (fr-CH)"),
        ("🌍 West African French",             "West African French"),
        ("🌍 North African French",            "North African French"),
        ("🇲🇬 Malagasy French",               "Malagasy French"),
    ],
    "pt": [
        ("🌍 Auto (General Portuguese)",       "General Portuguese"),
        ("🇧🇷 Brazilian Portuguese",           "Brazilian Portuguese (pt-BR)"),
        ("🇵🇹 European Portuguese",            "European Portuguese (pt-PT)"),
        ("🇦🇴 Angolan Portuguese",             "Angolan Portuguese (pt-AO)"),
        ("🇲🇿 Mozambican Portuguese",          "Mozambican Portuguese (pt-MZ)"),
        ("🇨🇻 Cape Verdean Portuguese",        "Cape Verdean Portuguese"),
        ("🇸🇹 São Tomé Portuguese",            "São Tomé Portuguese"),
    ],
    "de": [
        ("🌍 Auto (General German)",           "General German"),
        ("🇩🇪 Standard German (Germany)",      "Standard German (de-DE)"),
        ("🇦🇹 Austrian German",                "Austrian German (de-AT)"),
        ("🇨🇭 Swiss German",                   "Swiss German (de-CH)"),
        ("🇱🇺 Luxembourg German",              "Luxembourg German"),
        ("🌍 Bavarian dialect",                "Bavarian German"),
        ("🌍 Low German (Plattdeutsch)",        "Low German"),
    ],
    "it": [
        ("🌍 Auto (General Italian)",          "General Italian"),
        ("🇮🇹 Standard Italian",               "Standard Italian (it-IT)"),
        ("🇨🇭 Swiss Italian",                  "Swiss Italian (it-CH)"),
        ("🌍 Sicilian",                        "Sicilian Italian"),
        ("🌍 Neapolitan",                      "Neapolitan Italian"),
        ("🌍 Venetian dialect",                "Venetian Italian"),
        ("🌍 Roman dialect",                   "Roman Italian"),
    ],
    "zh": [
        ("🌍 Auto (General Chinese)",          "General Chinese"),
        ("🇨🇳 Mandarin Simplified (Mainland)", "Mainland Mandarin (zh-CN)"),
        ("🇹🇼 Mandarin Traditional (Taiwan)",  "Taiwan Mandarin (zh-TW)"),
        ("🇭🇰 Cantonese (Hong Kong)",          "Cantonese (zh-HK)"),
        ("🇸🇬 Singaporean Mandarin",           "Singaporean Mandarin (zh-SG)"),
        ("🌍 Shanghainese / Wu",               "Shanghainese (Wu dialect)"),
    ],
    "ja": [
        ("🌍 Auto (General Japanese)",         "General Japanese"),
        ("🇯🇵 Standard Japanese (Tokyo)",      "Standard Japanese (Tokyo)"),
        ("🌍 Kansai / Osaka dialect",          "Kansai dialect"),
        ("🌍 Kyushu dialect",                  "Kyushu dialect"),
        ("🌍 Tohoku dialect",                  "Tohoku dialect"),
    ],
    "ko": [
        ("🌍 Auto (General Korean)",           "General Korean"),
        ("🇰🇷 Standard Korean (Seoul)",        "Standard Korean (Seoul)"),
        ("🌍 Gyeongsang dialect",              "Gyeongsang dialect"),
        ("🌍 Jeolla dialect",                  "Jeolla dialect"),
        ("🌍 Jeju dialect",                    "Jeju dialect"),
    ],
    "ar": [
        ("🌍 Auto / Modern Standard Arabic",   "Modern Standard Arabic"),
        ("🇪🇬 Egyptian Arabic",                "Egyptian Arabic"),
        ("🇸🇦 Saudi Arabic",                   "Saudi Arabic"),
        ("🇦🇪 Gulf Arabic",                    "Gulf Arabic"),
        ("🇱🇧 Levantine Arabic",               "Levantine Arabic"),
        ("🇲🇦 Moroccan Arabic (Darija)",       "Moroccan Arabic"),
        ("🇮🇶 Iraqi Arabic",                   "Iraqi Arabic"),
        ("🇹🇳 Tunisian Arabic",                "Tunisian Arabic"),
        ("🇩🇿 Algerian Arabic",                "Algerian Arabic"),
        ("🇾🇪 Yemeni Arabic",                  "Yemeni Arabic"),
        ("🇸🇩 Sudanese Arabic",                "Sudanese Arabic"),
    ],
    "ru": [
        ("🌍 Auto (General Russian)",          "General Russian"),
        ("🇷🇺 Standard Russian (Moscow)",      "Standard Russian"),
        ("🌍 St. Petersburg Russian",          "St. Petersburg Russian"),
        ("🌍 Siberian Russian",                "Siberian Russian"),
        ("🌍 Ural Russian",                    "Ural Russian"),
    ],
    # ── Indian languages ────────────────────────────────────────────────────
    "hi": [
        ("🔍 Auto (General Hindi)",            "General Hindi"),
        ("🇮🇳 Standard Hindi (Delhi)",         "Standard Hindi (Delhi)"),
        ("🇮🇳 Mumbai Hindi",                   "Mumbai Hindi"),
        ("🇮🇳 Bihari Hindi",                   "Bihari Hindi"),
        ("🇮🇳 Rajasthani Hindi",               "Rajasthani Hindi"),
        ("🇮🇳 Bhojpuri-accented Hindi",        "Bhojpuri-accented Hindi"),
    ],
    "bn": [
        ("🔍 Auto (General Bengali)",          "General Bengali"),
        ("🇮🇳 Indian Bengali (Kolkata)",       "Indian Bengali (Kolkata)"),
        ("🇧🇩 Bangladeshi Bengali (Dhaka)",    "Bangladeshi Bengali (Dhaka)"),
        ("🌍 Sylheti dialect",                 "Sylheti dialect"),
        ("🌍 Chittagonian dialect",            "Chittagonian dialect"),
    ],
    "ta": [
        ("🔍 Auto (General Tamil)",            "General Tamil"),
        ("🇮🇳 Indian Tamil (Chennai)",         "Indian Tamil (Chennai)"),
        ("🇱🇰 Sri Lankan Tamil",               "Sri Lankan Tamil"),
        ("🇸🇬 Singaporean Tamil",              "Singaporean Tamil"),
        ("🇲🇾 Malaysian Tamil",                "Malaysian Tamil"),
    ],
    "te": [
        ("🔍 Auto (General Telugu)",           "General Telugu"),
        ("🇮🇳 Standard Telugu (Hyderabad)",    "Standard Telugu (Hyderabad)"),
        ("🇮🇳 Coastal Andhra Telugu",          "Coastal Andhra Telugu"),
        ("🇮🇳 Rayalaseema Telugu",             "Rayalaseema Telugu"),
    ],
    "gu": [
        ("🔍 Auto (General Gujarati)",         "General Gujarati"),
        ("🇮🇳 Standard Gujarati (Ahmedabad)",  "Standard Gujarati"),
        ("🇮🇳 Saurashtra Gujarati",            "Saurashtra Gujarati"),
        ("🇬🇧 British Gujarati",               "British Gujarati"),
        ("🇺🇸 American Gujarati",              "American Gujarati"),
    ],
    "kn": [
        ("🔍 Auto (General Kannada)",          "General Kannada"),
        ("🇮🇳 Standard Kannada (Bangalore)",   "Standard Kannada (Bangalore)"),
        ("🇮🇳 Dharwad Kannada",                "Dharwad Kannada"),
        ("🇮🇳 Old Mysore Kannada",             "Old Mysore Kannada"),
    ],
    "ml": [
        ("🔍 Auto (General Malayalam)",        "General Malayalam"),
        ("🇮🇳 Standard Malayalam (Kerala)",    "Standard Malayalam"),
        ("🇮🇳 Central Kerala dialect",         "Central Kerala dialect"),
        ("🇮🇳 North Kerala (Malabar) dialect", "North Kerala dialect"),
        ("🌍 Gulf Malayalam",                  "Gulf Malayalam"),
    ],
    "mr": [
        ("🔍 Auto (General Marathi)",          "General Marathi"),
        ("🇮🇳 Standard Marathi (Pune/Mumbai)", "Standard Marathi"),
        ("🇮🇳 Nagpuri Marathi",               "Nagpuri Marathi"),
        ("🇮🇳 Konkani-accented Marathi",       "Konkani-accented Marathi"),
    ],
    "pa": [
        ("🔍 Auto (General Punjabi)",          "General Punjabi"),
        ("🇮🇳 Indian Punjabi (Amritsar)",      "Indian Punjabi"),
        ("🇵🇰 Pakistani Punjabi (Lahore)",     "Pakistani Punjabi"),
        ("🇬🇧 British Punjabi",               "British Punjabi"),
        ("🇨🇦 Canadian Punjabi",              "Canadian Punjabi"),
    ],
    "ur": [
        ("🌍 Auto (General Urdu)",             "General Urdu"),
        ("🇵🇰 Pakistani Urdu",                 "Pakistani Urdu"),
        ("🇮🇳 Indian Urdu",                    "Indian Urdu"),
        ("🌍 Deccani Urdu",                    "Deccani Urdu"),
    ],
    # ── Other languages ─────────────────────────────────────────────────────
    "nl": [
        ("🌍 Auto (General Dutch)",            "General Dutch"),
        ("🇳🇱 Netherlands Dutch",              "Netherlands Dutch (nl-NL)"),
        ("🇧🇪 Belgian Dutch / Flemish",        "Belgian Dutch / Flemish (nl-BE)"),
        ("🇸🇷 Surinamese Dutch",               "Surinamese Dutch"),
    ],
    "tr": [
        ("🌍 Auto (General Turkish)",          "General Turkish"),
        ("🇹🇷 Istanbul Turkish",               "Istanbul Turkish"),
        ("🌍 Anatolian Turkish",               "Anatolian Turkish"),
        ("🇨🇾 Cypriot Turkish",                "Cypriot Turkish"),
    ],
    "vi": [
        ("🌍 Auto (General Vietnamese)",       "General Vietnamese"),
        ("🇻🇳 Northern Vietnamese (Hanoi)",    "Northern Vietnamese"),
        ("🇻🇳 Southern Vietnamese (Ho Chi Minh City)", "Southern Vietnamese"),
        ("🇻🇳 Central Vietnamese",             "Central Vietnamese"),
    ],
    "sv": [
        ("🌍 Auto (General Swedish)",          "General Swedish"),
        ("🇸🇪 Sweden Swedish",                 "Sweden Swedish"),
        ("🇫🇮 Finland Swedish",                "Finland Swedish"),
    ],
    "no": [
        ("🌍 Auto (General Norwegian)",        "General Norwegian"),
        ("🌍 Bokmål",                          "Norwegian Bokmål"),
        ("🌍 Nynorsk",                         "Norwegian Nynorsk"),
    ],
    "pl": [
        ("🌍 Auto (General Polish)",           "General Polish"),
        ("🌍 Warsaw Polish",                   "Warsaw Polish"),
        ("🌍 Silesian-accented Polish",        "Silesian-accented Polish"),
        ("🌍 Kashubian-accented Polish",       "Kashubian-accented Polish"),
    ],
    "th": [
        ("🌍 Auto (General Thai)",             "General Thai"),
        ("🇹🇭 Central Thai (Bangkok)",         "Central Thai"),
        ("🌍 Northern Thai (Kham Mueang)",     "Northern Thai"),
        ("🌍 Northeastern Thai (Isan)",        "Northeastern Thai"),
        ("🌍 Southern Thai",                   "Southern Thai"),
    ],
    "el": [
        ("🌍 Auto (General Greek)",            "General Greek"),
        ("🇬🇷 Standard Modern Greek",          "Standard Modern Greek"),
        ("🇨🇾 Cypriot Greek",                  "Cypriot Greek"),
    ],
    "he": [
        ("🌍 Auto (General Hebrew)",           "General Hebrew"),
        ("🇮🇱 Modern Israeli Hebrew",          "Modern Israeli Hebrew"),
        ("🌍 Mizrahi-accented Hebrew",         "Mizrahi-accented Hebrew"),
        ("🌍 Ashkenazi-accented Hebrew",       "Ashkenazi-accented Hebrew"),
    ],
    "ro": [
        ("🌍 Auto (General Romanian)",         "General Romanian"),
        ("🇷🇴 Romanian (Romania)",             "Romanian (ro-RO)"),
        ("🇲🇩 Moldovan Romanian",              "Moldovan Romanian"),
    ],
    "hu": [
        ("🌍 Auto (General Hungarian)",        "General Hungarian"),
        ("🇭🇺 Standard Hungarian",             "Standard Hungarian"),
        ("🇷🇴 Transylvanian Hungarian",        "Transylvanian Hungarian"),
        ("🇸🇰 Slovak Hungarian",               "Slovak Hungarian"),
    ],
    "cs": [
        ("🌍 Auto (General Czech)",            "General Czech"),
        ("🇨🇿 Bohemian Czech",                 "Bohemian Czech"),
        ("🌍 Moravian Czech",                  "Moravian Czech"),
        ("🌍 Silesian Czech",                  "Silesian Czech"),
    ],
    "fi": [
        ("🌍 Auto (General Finnish)",          "General Finnish"),
        ("🇫🇮 Standard Finnish",               "Standard Finnish"),
        ("🇫🇮 Helsinki Finnish",               "Helsinki Finnish"),
        ("🌍 Finland Swedish-accented Finnish","Finland Swedish-accented Finnish"),
    ],
    "da": [
        ("🌍 Auto (General Danish)",           "General Danish"),
        ("🇩🇰 Standard Danish (Copenhagen)",   "Standard Danish"),
        ("🌍 Jutlandic Danish",                "Jutlandic Danish"),
    ],
    "uk": [
        ("🌍 Auto (General Ukrainian)",        "General Ukrainian"),
        ("🇺🇦 Standard Ukrainian (Kyiv)",      "Standard Ukrainian"),
        ("🌍 Western Ukrainian",               "Western Ukrainian"),
        ("🌍 Eastern Ukrainian",               "Eastern Ukrainian"),
    ],
    "id": [
        ("🌍 Auto (General Indonesian)",       "General Indonesian"),
        ("🇮🇩 Standard Bahasa Indonesia",      "Standard Bahasa Indonesia"),
        ("🌍 Javanese-accented",               "Javanese-accented Indonesian"),
        ("🌍 Sundanese-accented",              "Sundanese-accented Indonesian"),
        ("🌍 Balinese-accented",               "Balinese-accented Indonesian"),
        ("🌍 Batak-accented",                  "Batak-accented Indonesian"),
    ],
}

# Values that represent "no specific variant" — filtered out before passing to backend
_VARIANT_AUTO_VALUES = {v[0][1] for v in LANGUAGE_VARIANTS.values()}


# ── helpers ───────────────────────────────────────────────────────────────────

def stats_to_markdown(speaker_stats) -> str:
    if not speaker_stats:
        return "_No speech analytics available. Upload an audio or video file to see speaker stats._"
    lines = []
    for s in speaker_stats:
        emoji = PACE_EMOJI.get(s.pace_label, "")
        wpm   = f"**{s.words_per_minute} WPM**" if s.words_per_minute else "N/A"
        pct   = f"{s.speaking_percentage}% of conversation" if s.speaking_percentage else ""

        lines += [
            f"### {emoji} {s.name}",
            "",
            "| Metric | Value |",
            "|--------|-------|",
            f"| 🎙️ Speech rate | {wpm} — {s.pace_label} |",
        ]
        if pct:
            lines.append(f"| 🕐 Speaking time | {pct} |")
        if s.accent_indicators:
            lines.append(f"| 🌍 Accent analysis | {s.accent_indicators} |")
        if s.accent_confidence:
            lines.append(f"| 🎯 Confidence | {s.accent_confidence.capitalize()} |")
        lines.append("")

    lines += [
        "---",
        "",
        "**Pace reference:** 🐢 Slow < 120 wpm · 🚶 Normal 120–150 · 🏃 Fast 150–180 · ⚡ Very Fast > 180",
    ]
    return "\n".join(lines)


# ── processing — generator streams every update live to the UI ────────────────
#
# Output order (15 items, must match outputs= list in .click()):
def _generate_pdf(stem: str, combined_text: str, path: Path) -> str:
    """Render the combined report as a formatted PDF using fpdf2."""  # noqa
    from fpdf import FPDF
    import datetime

    # Indigo palette
    _C_PRIMARY   = (99,  102, 241)   # #6366f1
    _C_PRIMARY_DK= (79,  70,  229)   # #4f46e5
    _C_ACCENT    = (139, 92,  246)   # #8b5cf6
    _C_BG_HDR    = (238, 242, 255)   # #eef2ff  — section header fill
    _C_TEXT      = (30,  30,  42)    # #1e1e2a
    _C_SUB       = (107, 107, 128)   # #6b6b80
    _C_RULE      = (199, 210, 254)   # #c7d2fe  — light indigo divider
    _C_WHITE     = (255, 255, 255)
    _C_FOOTER    = (156, 163, 175)   # gray-400

    class _PDF(FPDF):
        def header(self):
            # Indigo top bar
            self.set_fill_color(*_C_PRIMARY_DK)
            self.rect(0, 0, 210, 10, "F")
            # Accent stripe
            self.set_fill_color(*_C_ACCENT)
            self.rect(0, 10, 210, 2, "F")

        def footer(self):
            self.set_y(-13)
            # Bottom rule
            self.set_draw_color(*_C_RULE)
            self.set_line_width(0.4)
            self.line(15, self.get_y(), 195, self.get_y())
            self.ln(1)
            self.set_font("Helvetica", "I", 7.5)
            self.set_text_color(*_C_FOOTER)
            self.cell(0, 6, f"Transcript Agent  |  Page {self.page_no()}  |  {datetime.date.today()}", align="C")

    def _safe(text: str) -> str:
        import unicodedata
        # Replace common unicode → ASCII equivalents
        text = (text
                .replace("☐", "[ ]").replace("☑", "[x]")
                .replace("•", "-").replace("•", "-")
                .replace("’", "'").replace("‘", "'")
                .replace("“", '"').replace("”", '"')
                .replace("–", "-").replace("—", "--")
                .replace("…", "...")
                # Verdict / status emoji → text
                .replace("✅", "[OK]").replace("🟡", "[~]")
                .replace("⚠️", "[!]").replace("❌", "[X]")
                .replace("🟢", "[+]").replace("🟠", "[~]").replace("🔴", "[-]")
                .replace("⚡", "[!]").replace("📝", "").replace("💬", "")
                .replace("🎯", "").replace("📋", "").replace("🎤", "")
                .replace("📁", "").replace("🤖", ""))
        # Normalize and drop anything still outside latin-1
        text = unicodedata.normalize("NFKD", text)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(18, 18, 18)
    pdf.add_page()

    # Push content below the header bars
    pdf.set_y(18)

    # ── Title block ──────────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(*_C_PRIMARY_DK)
    pdf.multi_cell(0, 10, _safe(stem), align="C")
    pdf.ln(1)

    # Indigo rule under title
    pdf.set_draw_color(*_C_PRIMARY)
    pdf.set_line_width(0.8)
    pdf.line(18, pdf.get_y(), 192, pdf.get_y())
    pdf.ln(5)

    # ── Body ─────────────────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*_C_TEXT)

    for line in combined_text.splitlines():
        stripped = line.rstrip()
        # Skip pure divider lines (===... or ---...)
        if stripped and set(stripped) <= {"=", "-", " "} and len(stripped) > 4:
            continue
        if not stripped:
            pdf.ln(3)
            continue
        inner = stripped.strip()

        # Markdown-style headers: ## or ###
        if inner.startswith("### "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*_C_PRIMARY)
            pdf.multi_cell(0, 6, _safe(inner[4:]))
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*_C_TEXT)
        elif inner.startswith("## "):
            pdf.ln(5)
            pdf.set_fill_color(*_C_BG_HDR)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*_C_PRIMARY_DK)
            pdf.cell(0, 8, _safe(inner[3:]), new_x="LMARGIN", new_y="NEXT", fill=True)
            # Thin indigo underline
            pdf.set_draw_color(*_C_PRIMARY)
            pdf.set_line_width(0.3)
            pdf.line(18, pdf.get_y(), 192, pdf.get_y())
            pdf.ln(3)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*_C_TEXT)
        # ALL-CAPS section headers
        elif (inner.isupper() and 2 < len(inner) < 80
                and not set(inner) <= {"=", "-", " "}):
            pdf.ln(5)
            pdf.set_fill_color(*_C_BG_HDR)
            pdf.set_draw_color(*_C_PRIMARY)
            pdf.set_line_width(0.3)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*_C_PRIMARY_DK)
            pdf.cell(0, 8, _safe(inner), new_x="LMARGIN", new_y="NEXT", fill=True, border="B")
            pdf.ln(3)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*_C_TEXT)
        # Bullet lines
        elif inner.startswith("- ") or inner.startswith("* "):
            pdf.set_x(22)
            pdf.set_fill_color(*_C_PRIMARY)
            pdf.ellipse(18.5, pdf.get_y() + 2.2, 1.6, 1.6, "F")
            pdf.multi_cell(0, 5, _safe(inner[2:]))
        # Timestamp lines e.g. [00:01:23]
        elif inner.startswith("[") and "]" in inner[:12]:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(*_C_SUB)
            pdf.multi_cell(0, 5, _safe(stripped))
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*_C_TEXT)
        else:
            pdf.multi_cell(0, 5, _safe(stripped))

    pdf.output(str(path))
    return str(path)


def _generate_docx(stem: str, combined_text: str, path) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.core_properties.title = f"Transcript Report — {stem}"
        for line in combined_text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=2)
            elif set(stripped) <= {"=", "-", " "} and len(stripped) > 4:
                pass  # skip divider lines
            else:
                doc.add_paragraph(stripped)
        doc.save(str(path))
        return str(path)
    except Exception:
        return None


def _ts_add_secs(ts: str, extra: int, sep: str) -> str:
    parts = ts.split(":")
    h = int(parts[0]) if len(parts) == 3 else 0
    m = int(parts[-2])
    s = int(parts[-1]) + extra
    if s >= 60:
        m += s // 60; s = s % 60
    if m >= 60:
        h += m // 60; m = m % 60
    return f"{h:02d}:{m:02d}:{s:02d}{sep}000"


def _transcript_to_srt(transcript: str) -> str:
    import re
    ts_re = re.compile(r'^\[(\d{1,2}:\d{2}(?::\d{2})?)\]\s*(.*)')
    parsed = []
    for line in transcript.splitlines():
        m = ts_re.match(line.strip())
        if m:
            parsed.append((m.group(1), m.group(2)))
    if not parsed:
        return ""
    parts = []
    for i, (ts, text) in enumerate(parsed):
        parts_ts = ts.split(":")
        h = int(parts_ts[0]) if len(parts_ts) == 3 else 0
        m2 = int(parts_ts[-2])
        s2 = int(parts_ts[-1])
        start = f"{h:02d}:{m2:02d}:{s2:02d},000"
        end = _ts_add_secs(parsed[i + 1][0], 0, ",") if i + 1 < len(parsed) else _ts_add_secs(ts, 3, ",")
        parts.append(f"{i + 1}\n{start} --> {end}\n{text}\n")
    return "\n".join(parts)


def _transcript_to_vtt(transcript: str) -> str:
    srt = _transcript_to_srt(transcript)
    if not srt:
        return ""
    vtt = "WEBVTT\n\n"
    for block in srt.split("\n\n"):
        lines = block.strip().splitlines()
        if len(lines) >= 3:
            vtt += lines[1].replace(",", ".") + "\n" + "\n".join(lines[2:]) + "\n\n"
    return vtt


#   0  status_bar       1  summary_out      2  transcript_out   3  dialogue_out
#   4  profiles_out     5  interview_out    6  analytics_out    7  combined_out
#   8  dl_transcript    9  dl_speakers      10 dl_report
#   11 dl_combined      12 dl_json          13 dl_pdf
#   14 dl_docx          15 dl_srt           16 dl_vtt
#   17 download_accordion  18 log_out       19 eta_panel  20 result_state
# ---------------------------------------------------------------------------

_NOCHANGE = (gr.update(),) * 21   # yield this to keep connection alive without changes

def _out(status=gr.update(), summary=gr.update(), transcript=gr.update(),
         dialogue=gr.update(), profiles=gr.update(), interview=gr.update(),
         analytics=gr.update(), combined=gr.update(), dl_t=gr.update(),
         dl_s=gr.update(), dl_r=gr.update(), dl_c=gr.update(), dl_j=gr.update(),
         dl_p=gr.update(), dl_docx=gr.update(), dl_srt=gr.update(), dl_vtt=gr.update(),
         dl_acc=gr.update(), log=gr.update(), eta=gr.update(),
         rs=None):
    return (status, summary, transcript, dialogue, profiles, interview,
            analytics, combined, dl_t, dl_s, dl_r, dl_c, dl_j, dl_p,
            dl_docx, dl_srt, dl_vtt, dl_acc, log, eta, rs)


_PDF_LANGUAGES = [
    "Same as source",
    "English", "Spanish", "French", "German", "Portuguese",
    "Italian", "Dutch", "Russian", "Chinese (Simplified)",
    "Japanese", "Korean", "Arabic", "Hindi", "Turkish",
]


def _translate_combined_text(
    combined_text: str, target_language: str, api_key: str,
    provider: str = "anthropic", model: str = None, base_url: str = None,
) -> str:
    """Translate the combined report text to target_language using the selected provider."""
    _model = model or ("claude-sonnet-4-6" if provider == "anthropic" else "gpt-4o")
    client = LLMClient(provider=provider, api_key=api_key, model=_model, base_url=base_url)
    return client.chat(
        system="You are a professional translator.",
        user=(
            f"Translate the following transcript report to {target_language}. "
            "Preserve ALL formatting exactly — keep the section headers in ALL CAPS, "
            "keep divider lines (=== and ---), bullet characters (•, ☐), "
            "and timestamps in [HH:MM:SS] format unchanged. "
            "Only return the translated text, nothing else.\n\n"
            + combined_text
        ),
        max_tokens=8192,
    )


def generate_pdf_in_language(result_state, target_lang: str, api_key: str,
                              provider_name: str = "Claude (Anthropic)", model_name: str = None):
    """Generate (or re-generate) the PDF in the chosen language."""
    if not result_state:
        return gr.update()
    stem         = result_state["stem"]
    combined     = result_state["combined_text"]
    detected     = result_state.get("detected_language", "")
    out_dir      = Path(result_state["out_dir"])

    if target_lang and target_lang != "Same as source":
        cfg = _PROVIDERS.get(provider_name, _PROVIDERS["Claude (Anthropic)"])
        combined = _translate_combined_text(
            combined, target_lang, api_key,
            provider=cfg["type"], model=model_name, base_url=cfg["base_url"],
        )
        pdf_name = f"{stem}_report_{target_lang.replace(' ', '_')}.pdf"
    else:
        pdf_name = f"{stem}_report.pdf"

    pdf_path = out_dir / pdf_name
    try:
        _generate_pdf(f"{stem}  [{detected or target_lang}]", combined, pdf_path)
    except Exception as e:
        import traceback
        print(f"[PDF ERROR] {e}\n{traceback.format_exc()}")
        raise gr.Error(f"PDF generation failed: {e}")
    return gr.update(value=str(pdf_path), visible=True)


def _step_tracker_html(stage: str, done: bool = False) -> str:
    if done:
        states = ["done", "done", "done"]
    elif stage in ("loading",):
        states = ["active", "waiting", "waiting"]
    elif stage in ("extracting", "whisper"):
        states = ["done", "active", "waiting"]
    elif stage == "claude":
        states = ["done", "done", "active"]
    else:
        states = ["active", "waiting", "waiting"]

    labels = [("📁", "Upload"), ("🎤", "Transcribe"), ("🤖", "Analyze")]
    parts  = []
    for i, ((icon, label), state) in enumerate(zip(labels, states)):
        if state == "done":
            bg = "var(--ta-step-done-bg)"; bdr = "var(--ta-step-done-bdr)"
            clr = "var(--ta-step-done-clr)"; dot = "✓"
        elif state == "active":
            bg = "var(--ta-step-act-bg)"; bdr = "var(--ta-step-act-bdr)"
            clr = "var(--ta-step-act-clr)"; dot = "●"
        else:
            bg = "var(--ta-step-wait-bg)"; bdr = "var(--ta-step-wait-bdr)"
            clr = "var(--ta-step-wait-clr)"; dot = "○"
        parts.append(
            f'<div style="display:flex;flex-direction:column;align-items:center;gap:5px;flex:1;">'
            f'<div style="background:{bg};border:2px solid {bdr};border-radius:50%;'
            f'width:38px;height:38px;display:flex;align-items:center;justify-content:center;'
            f'font-size:1.05em;font-weight:800;color:{clr};transition:all 0.4s;">{dot}</div>'
            f'<div style="font-size:0.68em;font-weight:700;color:{clr};text-align:center;'
            f'text-transform:uppercase;letter-spacing:0.06em;">{icon} {label}</div>'
            f'</div>'
        )
        if i < 2:
            lc = "var(--ta-conn-line-done)" if states[i] == "done" else "var(--ta-conn-line-wait)"
            parts.append(
                f'<div style="flex:0.6;height:2px;background:{lc};margin-top:19px;'
                f'border-radius:2px;transition:background 0.4s;"></div>'
            )

    return (
        '<div style="display:flex;align-items:flex-start;padding:10px 16px 8px;'
        'background:var(--ta-card-bg);border:1px solid var(--ta-card-border);'
        'border-radius:12px;margin-bottom:10px;">'
        + "".join(parts) + "</div>"
    )


def _eta_panel_html(stage: str, pct: float = None, eta_secs: int = None,
                    elapsed: str = "", done: bool = False, tz_name: str = "") -> str:

    _slide_css = (
        "<style>@keyframes pgslide{0%{left:-45%}100%{left:110%}}</style>"
    )

    # ── Done ──────────────────────────────────────────────────────────────────
    tracker = _step_tracker_html(stage, done)

    if done:
        from datetime import datetime
        finished_at = datetime.now()
        hour = finished_at.hour % 12 or 12
        ampm = "AM" if finished_at.hour < 12 else "PM"
        finished_str = f"{hour}:{finished_at.minute:02d} {ampm}"
        return tracker + (
            '<div style="background:var(--ta-step-done-bg);'
            'border:2px solid var(--ta-step-done-bdr);border-radius:16px;padding:24px 28px;'
            'text-align:center;font-family:sans-serif;">'
            '<div style="font-size:2.5em;line-height:1;color:var(--ta-step-done-clr);">&#10003;</div>'
            '<div style="color:var(--ta-step-done-clr);font-size:1.4em;font-weight:800;margin-top:6px;">'
            'All Done!</div>'
            '<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;'
            'letter-spacing:0.1em;color:var(--ta-step-done-clr);margin-top:12px;">Finished at</div>'
            f'<div style="font-size:2.6em;font-weight:900;color:var(--ta-step-done-clr);'
            f'font-family:monospace;line-height:1.1;letter-spacing:-0.02em;">{finished_str}</div>'
            '<div style="display:flex;justify-content:center;gap:16px;margin-top:14px;flex-wrap:wrap;">'
            '<div style="background:var(--ta-stat-bg);border:1px solid var(--ta-step-done-bdr);'
            'border-radius:10px;padding:10px 20px;">'
            '<div style="font-size:0.68em;font-weight:700;text-transform:uppercase;'
            'letter-spacing:0.08em;color:var(--ta-step-done-clr);">Total Time</div>'
            f'<div style="font-size:1.5em;font-weight:800;color:var(--ta-card-val);'
            f'font-family:monospace;">{elapsed}</div>'
            '</div>'
            '<div style="background:var(--ta-stat-bg);border:1px solid var(--ta-step-done-bdr);'
            'border-radius:10px;padding:10px 20px;">'
            '<div style="font-size:0.68em;font-weight:700;text-transform:uppercase;'
            'letter-spacing:0.08em;color:var(--ta-step-done-clr);">Progress</div>'
            '<div style="font-size:1.5em;font-weight:800;color:var(--ta-card-val);">100%</div>'
            '</div>'
            '</div>'
            '</div>'
        )

    def _stat(label_txt, val_txt, label_var="--ta-stat-label", val_var="--ta-stat-val"):
        return (
            f'<div style="background:var(--ta-stat-bg);border-radius:8px;'
            f'padding:8px 14px;flex:1;min-width:90px;text-align:center;">'
            f'<div style="font-size:0.68em;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.08em;color:var({label_var});">{label_txt}</div>'
            f'<div style="font-size:1.3em;font-weight:800;color:var({val_var});'
            f'font-family:monospace;">{val_txt}</div></div>'
        )

    # ── shared "hero" time + ETA block ───────────────────────────────────────
    def _hero_time(finish_str, eta_str, border_color, label_color, approx=False):
        tilde = "~" if approx else ""
        finish_block = (
            f'<div style="text-align:center;margin-bottom:10px;">'
            f'<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.1em;color:{label_color};margin-bottom:2px;">Done by</div>'
            f'<div style="font-size:3em;font-weight:900;color:{label_color};'
            f'font-family:monospace;line-height:1;letter-spacing:-0.02em;">'
            f'{tilde}{finish_str}</div>'
            f'</div>'
        ) if finish_str else (
            f'<div style="text-align:center;margin-bottom:10px;">'
            f'<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.1em;color:{label_color};margin-bottom:2px;">Done by</div>'
            f'<div style="font-size:1.5em;font-weight:700;color:var(--ta-card-sub);">'
            f'calculating…</div></div>'
        )
        eta_block = (
            f'<div style="text-align:center;margin-bottom:14px;">'
            f'<div style="display:inline-block;background:var(--ta-stat-bg);'
            f'border:1px solid {border_color};border-radius:20px;'
            f'padding:5px 18px;font-size:1em;font-weight:700;color:{label_color};">'
            f'⏱ {eta_str}</div></div>'
        )
        return finish_block + eta_block

    # ── Whisper with real % ───────────────────────────────────────────────────
    if stage == "whisper" and pct is not None and pct > 0:
        pct_int    = int(pct * 100)
        bar_fill   = f"{pct_int}%"
        eta_str    = _fmt_eta(eta_secs) if (eta_secs and eta_secs > 0) else "calculating…"
        finish_str = _finish_time_str(eta_secs, tz_name) if (eta_secs and eta_secs > 0) else ""

        return tracker + (
            '<div style="background:var(--ta-card-bg);border:2px solid var(--ta-step-act-bdr);'
            'border-radius:16px;padding:20px 24px;font-family:sans-serif;">'
            '<div style="font-size:0.72em;font-weight:700;text-transform:uppercase;'
            'letter-spacing:0.1em;color:var(--ta-step-act-clr);margin-bottom:14px;">'
            'Step 1 of 2 &nbsp;&mdash;&nbsp; Transcribing Audio</div>'
            + _hero_time(finish_str, eta_str, "var(--ta-step-act-bdr)", "var(--ta-step-act-clr)") +
            '<div style="background:var(--ta-step-wait-bg);border-radius:8px;height:12px;'
            'overflow:hidden;margin-bottom:8px;">'
            f'<div style="width:{bar_fill};height:100%;'
            'background:linear-gradient(90deg,var(--ta-step-act-bdr),var(--ta-step-act-clr));'
            'border-radius:8px;transition:width 0.5s ease;"></div></div>'
            f'<div style="display:flex;justify-content:space-between;font-size:0.8em;'
            f'color:var(--ta-card-sub);">'
            f'<span>{pct_int}% complete</span>'
            f'<span>elapsed: {elapsed}</span></div>'
            '</div>'
        )

    # ── Claude AI stage — estimated ETA ──────────────────────────────────────
    if stage == "claude":
        eta_str    = _fmt_eta(eta_secs) if (eta_secs and eta_secs > 0) else "estimating…"
        finish_str = _finish_time_str(eta_secs, tz_name) if (eta_secs and eta_secs > 0) else ""
        return tracker + (
            '<div style="background:var(--ta-card-bg);border:2px solid #a855f7;'
            'border-radius:16px;padding:20px 24px;font-family:sans-serif;">'
            '<div style="font-size:0.72em;font-weight:700;text-transform:uppercase;'
            'letter-spacing:0.1em;color:#c4b5fd;margin-bottom:14px;">'
            'Step 2 of 2 &nbsp;&mdash;&nbsp; Analyzing with AI</div>'
            + _hero_time(finish_str, eta_str, "#a855f7", "#c4b5fd", approx=bool(finish_str)) +
            _slide_css +
            '<div style="background:var(--ta-step-wait-bg);border-radius:8px;height:12px;'
            'overflow:hidden;position:relative;margin-bottom:8px;">'
            '<div style="position:absolute;width:40%;height:100%;background:#a855f7;'
            'border-radius:8px;opacity:0.85;animation:pgslide 1.6s ease-in-out infinite;"></div>'
            '</div>'
            f'<div style="text-align:right;font-size:0.8em;color:var(--ta-card-sub);">'
            f'elapsed: {elapsed}</div>'
            '</div>'
        )

    # ── Other stages (loading / extracting / whisper indeterminate) ──────────
    stage_cfg = {
        "loading":    ("var(--ta-step-act-bdr)",  "Starting up…",        "Step 0 of 2", "var(--ta-step-act-clr)"),
        "extracting": ("var(--ta-step-done-bdr)", "Extracting audio…",   "Step 1 of 2", "var(--ta-step-done-clr)"),
        "whisper":    ("var(--ta-step-act-bdr)",  "Transcribing audio…", "Step 1 of 2", "var(--ta-step-act-clr)"),
        "claude":     ("#a855f7",                 "Analyzing with AI…",  "Step 2 of 2", "#c4b5fd"),
    }
    color, label, step, text_clr = stage_cfg.get(
        stage, ("var(--ta-card-border)", "Processing…", "", "var(--ta-card-sub)")
    )

    overlay_pct = ""
    if stage == "claude":
        overlay_pct = (
            '<div style="display:flex;align-items:flex-end;gap:4px;margin-bottom:14px;">'
            f'<div style="font-size:4.5em;font-weight:900;color:{text_clr};'
            f'font-family:monospace;line-height:1;letter-spacing:-0.04em;">50</div>'
            f'<div style="font-size:2em;font-weight:700;color:{color};margin-bottom:6px;">%+</div>'
            '</div>'
            f'<div style="font-size:0.82em;color:var(--ta-card-sub);margin-bottom:12px;">'
            'AI is reading the transcript and writing your report…</div>'
        )
    elif stage in ("loading", "extracting"):
        overlay_pct = (
            '<div style="display:flex;align-items:flex-end;gap:4px;margin-bottom:14px;">'
            f'<div style="font-size:4.5em;font-weight:900;color:{text_clr};'
            f'font-family:monospace;line-height:1;">—</div></div>'
        )

    return tracker + (
        f'<div style="background:var(--ta-card-bg);border:2px solid {color};'
        f'border-radius:16px;padding:20px 24px;font-family:sans-serif;">'
        f'<div style="font-size:0.72em;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.1em;color:{text_clr};margin-bottom:14px;">'
        f'{step} &nbsp;&mdash;&nbsp; {label}</div>'
        f'<div style="text-align:center;margin-bottom:10px;">'
        f'<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.1em;color:{text_clr};margin-bottom:2px;">Done by</div>'
        f'<div style="font-size:1.6em;font-weight:700;color:var(--ta-card-sub);">calculating…</div>'
        f'</div>'
        f'{overlay_pct}'
        f'{_slide_css}'
        f'<div style="background:var(--ta-step-wait-bg);border-radius:8px;height:12px;'
        f'overflow:hidden;position:relative;margin-bottom:8px;">'
        f'<div style="position:absolute;width:40%;height:100%;background:{color};'
        f'border-radius:8px;opacity:0.85;animation:pgslide 1.6s ease-in-out infinite;"></div>'
        f'</div>'
        f'<div style="text-align:right;font-size:0.8em;color:var(--ta-card-sub);">elapsed: {elapsed}</div>'
        f'</div>'
    )


def _err(msg: str) -> tuple:
    """Yield this tuple to display an inline error card instead of a popup."""
    html = (
        '<div style="background:linear-gradient(135deg,#fef2f2,#fee2e2);'
        'border:2px solid #ef4444;border-radius:12px;padding:18px 22px;'
        'display:flex;align-items:flex-start;gap:14px;font-family:sans-serif;">'
        '<div style="font-size:1.8em;line-height:1;flex-shrink:0;">❌</div>'
        '<div>'
        '<div style="color:#991b1b;font-weight:700;font-size:1em;">Something went wrong</div>'
        f'<div style="color:#b91c1c;font-size:0.88em;margin-top:5px;">{msg}</div>'
        '</div>'
        '</div>'
    )
    return _out(status=html, eta="")


def process_file(
    uploaded_file,
    path_input,
    panel_mode,
    speaker_names_raw,
    num_speakers,
    whisper_model,
    language_input,
    language_variant,
    report_style,
    inc_summary,
    inc_key_points,
    inc_action_items,
    inc_transcript,
    inc_profiles,
    inc_analytics,
    inc_interview=False,
    inc_interview_deep=False,
    resume_context="",
    user_api_key="",
    provider_name="Claude (Anthropic)",
    model_name="claude-sonnet-4-6",
    custom_base_url="",
    tz_name="",
    analysis_depth="balanced",
    stt_provider="Whisper (Local)",
    stt_cloud_key="",
    stt_cloud_model="",
):
    # ── validation (all errors shown inline, no popup) ────────────────────────
    api_key = (user_api_key or "").strip()
    provider_cfg = _PROVIDERS.get(provider_name, _PROVIDERS["Claude (Anthropic)"])
    no_key_ok = provider_name in ("Ollama (Local)", "Custom (OpenAI-compatible)")
    if not api_key and not no_key_ok:
        yield _err(f"Please enter your {provider_name} API key at the top of the page.")
        return
    provider_type = provider_cfg["type"]
    base_url = (custom_base_url or "").strip() if provider_name == "Custom (OpenAI-compatible)" else provider_cfg["base_url"]
    if provider_name == "Custom (OpenAI-compatible)" and not base_url:
        yield _err("Please enter the API base URL for your custom provider.")
        return

    _stt_cloud_key = (stt_cloud_key or "").strip()
    _stt_cloud_model = (stt_cloud_model or "").strip()
    _stt_id = _STT_PROVIDERS.get(stt_provider, {}).get("id", "whisper")
    _stt_is_cloud = _stt_id != "whisper"
    if _stt_is_cloud and not _stt_cloud_key:
        yield _err(f"Please enter your {stt_provider} API key in the Transcription Engine section.")
        return

    _prevent_sleep()

    # prefer pasted path/URL (no upload wait) over drag-and-drop
    pasted = (path_input or "").strip().strip('"').strip("'")
    if pasted:
        uploaded_file = pasted

    if not uploaded_file:
        yield _err("Please drag a file, paste a file path, or paste a URL above.")
        return

    # ── Log helpers (must be defined before download section uses them) ────────
    start_time    = time.time()
    log_entries   = []

    def _ts():
        secs = int(time.time() - start_time)
        m, s = divmod(secs, 60)
        return f"{m:02d}:{s:02d}"

    def _elapsed():
        secs = int(time.time() - start_time)
        m, s = divmod(secs, 60)
        return f"{m}m {s:02d}s" if m else f"{s}s"

    _KIND_COLORS = {
        'header':   ('#f8fafc', True),
        'info':     ('#94a3b8', False),
        'progress': ('#86efac', False),
        'warn':     ('#fbbf24', False),
        'error':    ('#f87171', False),
        'done':     ('#4ade80', True),
        'download': ('#22d3ee', False),
        'ai':       ('#c4b5fd', False),
    }

    def _render_log():
        parts = []
        for kind, ts, text in log_entries:
            color, bold = _KIND_COLORS.get(kind, ('#94a3b8', False))
            weight = 'font-weight:700;' if bold else ''
            if kind == 'header':
                parts.append(
                    f'<div style="color:var(--ta-log-hdr,#f8fafc);{weight}margin-top:8px;'
                    f'border-top:1px solid var(--ta-log-border,#1e3a5f);'
                    f'padding-top:6px;letter-spacing:0.05em;">{text}</div>'
                )
            else:
                # text-only line in log — the ETA panel owns the visual bar for 'progress'
                parts.append(
                    f'<div><span style="color:var(--ta-log-ts,#64748b);">[{ts}]</span> '
                    f'<span style="color:{color};{weight}">{text}</span></div>'
                )
        scroll = '<div id="ta-log-end"></div><script>document.getElementById("ta-log-end")?.scrollIntoView();</script>'
        inner = "".join(parts) + scroll if parts else '<span style="color:var(--ta-log-ts,#64748b);">Starting…</span>'
        return (
            '<div id="ta-log-wrap" style="background:var(--ta-log-bg,#0f172a);'
            'border:1px solid var(--ta-log-border,#1e3a5f);'
            'border-radius:10px;padding:12px 16px;min-height:120px;max-height:260px;'
            'overflow-y:auto;font-family:\'Courier New\',monospace;font-size:0.80em;line-height:1.7;">'
            + inner + '</div>'
        )

    def _add_log(text, kind='info'):
        log_entries.append((kind, _ts(), text))
        return _render_log()

    def _add_header(text):
        log_entries.append(('header', _ts(), text))
        return _render_log()

    # ── Download remote URL (threaded so we can stream progress to the log) ───
    if isinstance(uploaded_file, str) and (
        uploaded_file.startswith("http://") or uploaded_file.startswith("https://")
    ):
        _dl_log = _add_header("📥  DOWNLOADING FILE")
        yield _out(
            status=_status_compact("⬇️", "Downloading file from URL…", "0s"),
            eta=_eta_panel_html("loading", elapsed="0s", tz_name=tz_name),
            log=_dl_log,
        )
        _dl_dir  = Path(tempfile.mkdtemp(prefix="ta_dl_"))
        _dl_q    = Q.Queue()
        _dl_done = False

        def _dl_worker():
            def _on_progress(received, total):
                _dl_q.put(("dl_progress", received, total))
            try:
                path = _download_url(uploaded_file, _dl_dir, on_progress=_on_progress)
                _dl_q.put(("dl_done", str(path)))
            except Exception as exc:
                _dl_q.put(("dl_error", str(exc)))

        threading.Thread(target=_dl_worker, daemon=True).start()

        _dl_stall = 0
        while True:
            try:
                dmsg = _dl_q.get(timeout=1.0)
                _dl_stall = 0
                last_activity = time.time()
                if dmsg[0] == "dl_done":
                    uploaded_file = dmsg[1]
                    log = _add_log(f"✅ Download complete — {Path(uploaded_file).name}", "done")
                    yield _out(log=log)
                    break
                elif dmsg[0] == "dl_error":
                    yield _err(f"Download failed: {dmsg[1]}")
                    return
                elif dmsg[0] == "dl_progress":
                    recv, total = dmsg[1], dmsg[2]
                    if total:
                        pct = recv / total * 100
                        log = _add_log(f"⬇️  {recv//1024//1024} MB / {total//1024//1024} MB  ({pct:.0f}%)", "download")
                    else:
                        log = _add_log(f"⬇️  {recv//1024//1024} MB received…", "download")
                    yield _out(
                        status=_status_compact("⬇️", "Downloading…", _elapsed()),
                        log=log,
                    )
            except Q.Empty:
                _dl_stall += 1
                elapsed = _elapsed()
                if _dl_stall == 30:
                    log = _add_log("⚠️  Still downloading… (30s with no data). Large file or slow connection.", "warn")
                    yield _out(log=log)
                elif _dl_stall == 55:
                    log = _add_log("🚨  About to timeout (5s left). If this fails, download manually via browser.", "error")
                    yield _out(log=log)
                else:
                    yield _out(status=_status_compact("⬇️", "Downloading…", elapsed))

    from pathlib import Path as _P
    if not _P(uploaded_file).exists():
        yield _err(f"File not found: {uploaded_file}")
        return

    _fname = Path(uploaded_file).name
    log = _add_log(f"📂  File ready: {_fname}")
    # Tell the browser a job is starting — survives page refresh
    _job_js = f"<script>window.taJobStart && window.taJobStart({repr(_fname)})</script>"
    yield _out(
        status=_status_compact("⏳", "Starting…", _elapsed()) + _job_js,
        eta=_eta_panel_html("loading", elapsed=_elapsed(), tz_name=tz_name),
        log=log,
    )

    config = ReportConfig(
        style=report_style,
        include_summary=inc_summary,
        include_key_points=inc_key_points,
        include_action_items=inc_action_items,
        include_transcript=inc_transcript,
        include_speaker_profiles=inc_profiles,
        include_speech_analytics=inc_analytics,
        include_interview_mode=inc_interview,
        include_interview_deep=inc_interview_deep,
        interview_resume_context=resume_context or "",
        analysis_depth=analysis_depth or "balanced",
    )
    # Use names if provided, otherwise fall back to numeric count
    _names = (speaker_names_raw or "").strip()
    if _names:
        speaker_names = _names
    else:
        try:
            _n = int(num_speakers) if num_speakers not in (None, "", 0) else None
        except (ValueError, TypeError):
            _n = None
        speaker_names = f"{_n} speakers" if _n and _n >= 1 else None
    speakers = None  # WhisperX diarization disabled (requires HF_TOKEN)
    lang_code = language_input if language_input and language_input != "auto" else None
    lang_variant = (
        language_variant
        if language_variant and language_variant not in _VARIANT_AUTO_VALUES
        else None
    )
    stem     = Path(uploaded_file).stem
    job_id   = uuid.uuid4().hex[:8]
    job_dir  = OUT_DIR / f"{stem}_{job_id}"
    job_dir.mkdir(parents=True, exist_ok=True)
    ext      = Path(uploaded_file).suffix.lower()
    is_av    = ext in (AUDIO_EXTS | VIDEO_EXTS)

    _write_job_status("running", stem=stem, job_id=job_id, job_dir=str(job_dir))

    # ── Checkpoint: compute file hash and look for saved Whisper text ─────────
    _panel = bool(inc_profiles)
    _file_hash = _jdb.file_md5(uploaded_file) if _JDB_OK else ""
    _checkpoint = (_jdb.find_whisper_checkpoint(_file_hash, panel_mode=_panel)
                   if (_JDB_OK and is_av and _file_hash) else None)
    _ckpt_text  = _checkpoint[0] if _checkpoint else None
    _ckpt_json  = _checkpoint[1] if _checkpoint else None

    if _JDB_OK:
        try:
            import json as _j
            _jdb.create_job(
                job_id=job_id, stem=stem,
                original_filename=Path(uploaded_file).name,
                file_hash=_file_hash, job_dir=str(job_dir),
                config_json=_j.dumps({"panel_mode": _panel,
                                      "analysis_depth": analysis_depth}),
                panel_mode=_panel,
            )
        except Exception:
            pass

    # ── thread communication ──────────────────────────────────────────────────
    q = Q.Queue()

    def on_whisper_progress(pct):
        q.put(("pct", pct))

    def on_raw_transcript(text):
        q.put(("transcript", text))

    def on_stage_change(stage):
        q.put(("stage", stage))

    def on_log(msg):
        q.put(("log", msg))

    def _on_whisper_done(text, json_str):
        if _JDB_OK:
            try:
                _jdb.save_whisper_checkpoint(job_id, text, json_str)
            except Exception:
                pass

    def background():
        try:
            result = run(
                file_path=uploaded_file,
                output_dir=str(job_dir),
                whisper_model=whisper_model,
                panel_mode=bool(inc_profiles),
                num_speakers=None,
                config=config,
                api_key=api_key,
                provider=provider_type,
                model=model_name,
                base_url=base_url,
                language=lang_code,
                language_variant=lang_variant,
                speaker_names=speaker_names or None,
                on_whisper_progress=on_whisper_progress if (is_av and not _stt_is_cloud) else None,
                on_raw_transcript=on_raw_transcript if is_av else None,
                on_stage_change=on_stage_change if is_av else None,
                on_log=on_log,
                checkpoint_text=_ckpt_text,
                checkpoint_json=_ckpt_json,
                on_whisper_done=_on_whisper_done if is_av else None,
                stt_provider=_stt_id,
                stt_api_key=_stt_cloud_key or None,
                stt_model=_stt_cloud_model or None,
            )
            # Write status directly so it persists even if the browser disconnected
            import datetime as _dtt
            _write_job_status("done", stem=stem, job_id=job_id, job_dir=str(job_dir),
                              completed=_dtt.datetime.now().isoformat())
            q.put(("done", result))
        except Exception as e:
            # Write error status directly — generator may already be dead
            _write_job_status("error", stem=stem, job_id=job_id, job_dir=str(job_dir),
                              error=str(e)[:300])
            q.put(("error", str(e)))

    t = threading.Thread(target=background, daemon=True)
    t.start()

    # ── live update loop ──────────────────────────────────────────────────────
    whisper_pct     = 0.0
    raw_shown       = False
    claude_started  = False
    claude_start_time: float = 0.0
    stt_stage_start: float   = 0.0
    stt_elapsed_str: str     = ""
    stage           = "loading"
    last_activity   = time.time()
    stall_warned    = set()

    if _ckpt_text:
        log_entries.append(("info", _ts(),
            f"Whisper checkpoint found — skipping transcription for '{stem}' (same file). Passing saved text to AI."))


    def _eta_secs(pct):
        if pct <= 0.01:
            return None
        return max(0, int((time.time() - start_time) * (1.0 - pct) / pct))

    def _claude_eta_secs():
        if not claude_start_time:
            return None
        in_claude = time.time() - claude_start_time
        # Estimate: AI analysis typically takes 60-180s; scale up as time passes
        estimated_total = max(90.0, in_claude * 1.6)
        return max(0, int(estimated_total - in_claude))

    def _eta_str(pct):
        s = _eta_secs(pct)
        return _fmt_eta(s) if s is not None else ""

    try:
     while True:
        try:
            msg = q.get(timeout=1.0)
            last_activity = time.time()
        except Q.Empty:
            elapsed  = _elapsed()
            quiet    = int(time.time() - last_activity)
            eta_upd  = gr.update()

            # ── stall detection ──────────────────────────────────────────────
            stall_key = f"{stage}_{quiet}"
            if stage == "loading" and quiet == 60 and stall_key not in stall_warned:
                stall_warned.add(stall_key)
                log = _add_log("⚠️  Still preparing… (60s). Large file or slow start.", "warn")
                yield _out(log=log)
            elif stage == "extracting" and quiet == 120 and stall_key not in stall_warned:
                stall_warned.add(stall_key)
                log = _add_log("⚠️  Audio extraction taking 2+ min. Large video file — still working.", "warn")
                yield _out(log=log)
            elif stage == "whisper" and quiet == 300 and stall_key not in stall_warned:
                stall_warned.add(stall_key)
                log = _add_log("⚠️  Whisper running 5 min without a progress update. CPU transcription is slow for long recordings — still active.", "warn")
                yield _out(log=log)
            elif stage == "whisper" and quiet == 900 and stall_key not in stall_warned:
                stall_warned.add(stall_key)
                log = _add_log("🚨  15 min with no Whisper update. If progress bar is frozen, consider restarting and using a smaller Whisper model (tiny/base).", "error")
                yield _out(log=log)
            elif stage == "claude" and quiet == 120 and stall_key not in stall_warned:
                stall_warned.add(stall_key)
                log = _add_log("⚠️  AI analysis taking 2+ min. Long transcript or slow API — still waiting.", "warn")
                yield _out(log=log)
            elif stage == "claude" and quiet == 600 and stall_key not in stall_warned:
                stall_warned.add(stall_key)
                log = _add_log("🚨  10 min waiting for AI response. Check your API key quota or try a faster model.", "error")
                yield _out(log=log)

            # ── eta panel update ─────────────────────────────────────────────
            if stage in ("whisper",):
                eta_s   = _eta_secs(whisper_pct) if whisper_pct > 0 else None
                eta_upd = _eta_panel_html("whisper", pct=whisper_pct or None,
                                          eta_secs=eta_s, elapsed=elapsed, tz_name=tz_name)
                yield _out(status=_status_compact("🎤", "Transcribing audio…", elapsed), eta=eta_upd)
            elif stage == "extracting":
                yield _out(status=_status_compact("🎬", "Extracting audio…", elapsed),
                           eta=_eta_panel_html("extracting", elapsed=elapsed, tz_name=tz_name))
            elif stage in ("claude",) or claude_started:
                yield _out(status=_status_compact("🤖", "Analyzing with AI…", elapsed),
                           eta=_eta_panel_html("claude", elapsed=elapsed, eta_secs=_claude_eta_secs(), tz_name=tz_name))
            else:
                yield _out(status=_status_compact("⏳", "Loading…", elapsed),
                           eta=_eta_panel_html("loading", elapsed=elapsed, tz_name=tz_name))
            continue

        kind = msg[0]

        if kind == "log":
            log_text = _add_log(msg[1], "info")
            yield _out(log=log_text)

        elif kind == "stage":
            stage   = msg[1]
            elapsed = _elapsed()
            if stage == "extracting":
                log = _add_header("🎬  EXTRACTING AUDIO")
                yield _out(status=_status_compact("🎬", "Extracting audio from video…", elapsed),
                           eta=_eta_panel_html("extracting", elapsed=elapsed, tz_name=tz_name), log=log)
            elif stage == "whisper":
                stt_stage_start = time.time()
                log = _add_header(f"🎤  TRANSCRIBING AUDIO  ({stt_provider}) — Step 1 of 2")
                log = _add_log("Transcription in progress…", "info")
                yield _out(status=_status_compact("🎤", "Transcribing audio…", elapsed),
                           eta=_eta_panel_html("whisper", elapsed=elapsed, tz_name=tz_name), log=log)
            elif stage == "claude":
                if not claude_start_time:
                    claude_start_time = time.time()
                log = _add_header("🤖  AI ANALYSIS  (Step 2 of 2)")
                log = _add_log("Sending transcript to AI for analysis…", "ai")
                yield _out(status=_status_compact("🤖", "Analyzing with AI…", elapsed),
                           eta=_eta_panel_html("claude", elapsed=elapsed, eta_secs=_claude_eta_secs(), tz_name=tz_name), log=log)

        elif kind == "pct":
            whisper_pct = msg[1]
            elapsed     = _elapsed()
            eta_s       = _eta_secs(whisper_pct)
            eta_txt     = _eta_str(whisper_pct)
            pct_int     = int(whisper_pct * 100)
            # Log: clean text only — ETA panel owns the visual bar
            log_text = _add_log(
                f"🎤 {pct_int}%{('  —  ETA ' + eta_txt) if eta_txt else ''}",
                "progress"
            )
            yield _out(
                status=_status_compact("🎤", f"Transcribing…  {pct_int}%", elapsed),
                eta=_eta_panel_html("whisper", pct=whisper_pct, eta_secs=eta_s, elapsed=elapsed, tz_name=tz_name),
                log=log_text,
            )

        elif kind == "transcript":
            raw_shown = True
            elapsed   = _elapsed()
            if stt_stage_start:
                _stt_secs = int(time.time() - stt_stage_start)
                _m, _s = divmod(_stt_secs, 60)
                stt_elapsed_str = f"{_m}m {_s:02d}s" if _m else f"{_s}s"
            _stt_done_msg = f"✅ Transcription complete! ({stt_provider} took {stt_elapsed_str})" if stt_elapsed_str else "✅ Transcription complete!"
            log_text  = _add_log(_stt_done_msg, "done")
            log_text  = _add_header("🤖  AI ANALYSIS  (Step 2 of 2)")
            log_text  = _add_log("Sending transcript to AI for analysis…", "ai")
            if not claude_start_time:
                claude_start_time = time.time()
            yield _out(
                status=_status_compact("🤖", "Analyzing with AI…", elapsed),
                eta=_eta_panel_html("claude", elapsed=elapsed, eta_secs=_claude_eta_secs(), tz_name=tz_name),
                transcript=msg[1],
                log=log_text,
            )
            claude_started = True
            stage = "claude"

        elif kind == "done":
            result = msg[1]

            summary_md = f"## Summary\n\n{result.summary}"
            if inc_key_points and result.key_points:
                summary_md += "\n\n## Key Points\n" + "\n".join(f"- {p}" for p in result.key_points)
            if inc_action_items and result.action_items:
                summary_md += "\n\n## Action Items\n" + "\n".join(f"- [ ] {a}" for a in result.action_items)
            if stt_elapsed_str:
                summary_md += f"\n\n---\n\n_⏱ Transcription engine: **{stt_provider}** — completed in **{stt_elapsed_str}**_"
            if result.clean_transcript:
                summary_md += "\n\n---\n\n## 📄 Transcript\n\n" + result.clean_transcript
            if result.speaker_dialogue:
                summary_md += "\n\n---\n\n## 🗣️ Speaker Dialogue\n\n" + result.speaker_dialogue

            # ── Speaker profiles ──────────────────────────────────────────────
            if result.speaker_profiles:
                profiles_md = "\n\n---\n\n".join(
                    f"### {name}\n\n{profile}" for name, profile in result.speaker_profiles.items()
                )
                if result.speaker_map:
                    mapping = "\n".join(f"- `{k}` → **{v}**" for k, v in result.speaker_map.items())
                    profiles_md = f"## Speaker Map\n\n{mapping}\n\n---\n\n{profiles_md}"
            else:
                profiles_md = "_No speaker profiles found. The AI may not have detected multiple distinct speakers in this recording._"

            # ── Interview Q&A — dedicated Interview tab content ──────────────────
            interview_md = ""
            if inc_interview and result.interview_questions:
                _vicon  = {"strong": "✅", "acceptable": "🟡", "weak": "⚠️", "missed": "❌"}
                _vlabel = {"strong": "Great", "acceptable": "Good", "weak": "Needs Improvement", "missed": "Missed"}
                _ilines = ["## 🎤 Interview Q&A\n"]

                if result.round_advance_probability >= 0:
                    _prob = result.round_advance_probability
                    _plabel = ("Strong" if _prob >= 80 else "Competitive" if _prob >= 60
                               else "Borderline" if _prob >= 40 else "Unlikely")
                    _pemoji = "🟢" if _prob >= 80 else "🟡" if _prob >= 60 else "🟠" if _prob >= 40 else "🔴"
                    _ilines.append(f"**Overall likelihood of advancing: {_pemoji} {_prob}% — {_plabel}**\n")

                for _qi, _q in enumerate(result.interview_questions, 1):
                    _verdict = (_q.get("verdict") or "").lower()
                    _icon = _vicon.get(_verdict, "•")
                    _ilines += [
                        f"### Q{_qi}. {_q.get('question', '')}",
                        "",
                        f"**Verdict:** {_icon} {_vlabel.get(_verdict, (_q.get('verdict') or '').upper())}",
                        "",
                    ]
                    if _q.get("your_answer_summary"):
                        _ilines += [
                            "**📝 What you said:**",
                            "",
                            f"> {_q['your_answer_summary']}",
                            "",
                        ]
                    if _q.get("deflection_detected"):
                        _ilines += [f"**⚡ Deflection detected:** {_q.get('deflection_note', 'Candidate stalled or deflected')}", ""]
                    if _q.get("ideal_answer"):
                        _ilines += [
                            "**💬 How you could have answered it:**",
                            "",
                            f"> {_q['ideal_answer']}",
                            "",
                        ]
                    if _q.get("feedback"):
                        _ilines += [f"**🎯 Coaching tip:** {_q['feedback']}", ""]
                    _ilines.append("---\n")

                if inc_interview_deep and result.prep_guide:
                    _ilines += [
                        "## 📋 Prep Guide — Questions to Practise\n",
                        "_Review these before your next interview:_\n",
                    ]
                    for _pi, _pg in enumerate(result.prep_guide, 1):
                        _ilines += [
                            f"### P{_pi}. {_pg.get('question', '')}",
                            "",
                        ]
                        if _pg.get("why_it_matters"):
                            _ilines += [f"**Why they ask it:** {_pg['why_it_matters']}", ""]
                        if _pg.get("suggested_answer"):
                            _ilines += [
                                "**💬 Suggested answer:**",
                                "",
                                f"> {_pg['suggested_answer']}",
                                "",
                            ]
                        _ilines.append("---\n")

                interview_md = "\n".join(_ilines)

            analytics_md  = stats_to_markdown(result.speaker_stats)
            combined_text = build_combined_report(result, config)

            f_t = str(job_dir / f"{stem}_transcript.txt")
            f_s = str(job_dir / f"{stem}_speakers.txt")
            f_r = str(job_dir / f"{stem}_report.md")
            f_c = str(job_dir / f"{stem}_combined.txt")
            f_j = str(job_dir / f"{stem}_full.json")
            f_p_path = job_dir / f"{stem}_report.pdf"
            try:
                _generate_pdf(stem, combined_text, f_p_path)
                f_p = str(f_p_path)
            except Exception:
                f_p = None

            f_docx = _generate_docx(stem, combined_text, job_dir / f"{stem}_report.docx")

            srt_content = _transcript_to_srt(result.clean_transcript)
            f_srt = None
            if srt_content:
                _srt_path = job_dir / f"{stem}_transcript.srt"
                _srt_path.write_text(srt_content, encoding="utf-8")
                f_srt = str(_srt_path)

            vtt_content = _transcript_to_vtt(result.clean_transcript)
            f_vtt = None
            if vtt_content:
                _vtt_path = job_dir / f"{stem}_transcript.vtt"
                _vtt_path.write_text(vtt_content, encoding="utf-8")
                f_vtt = str(_vtt_path)

            total_elapsed = _elapsed()
            import datetime as _dtt
            _rd = {"summary": summary_md, "transcript": result.clean_transcript,
                   "dialogue": result.speaker_dialogue, "profiles": profiles_md,
                   "interview": interview_md,
                   "analytics": analytics_md, "combined": combined_text}
            _write_job_status("done", stem=stem, job_id=job_id, job_dir=str(job_dir),
                              completed=_dtt.datetime.now().isoformat(), result=_rd)
            if _JDB_OK:
                try:
                    _jdb.complete_job(job_id, _rd)
                except Exception:
                    pass
            log_text = _add_header("✅  COMPLETE")
            log_text = _add_log(f"All done in {total_elapsed}. Results ready in all tabs.", "done")
            yield _out(
                status=_status_compact("✅", "Done! All tabs are ready.", total_elapsed)
                      + "<script>window.taJobEnd && window.taJobEnd()</script>",
                eta=_eta_panel_html("done", elapsed=total_elapsed, done=True, tz_name=tz_name),
                summary=summary_md,
                transcript=result.clean_transcript,
                dialogue=result.speaker_dialogue,
                profiles=profiles_md,
                interview=interview_md,
                analytics=analytics_md,
                combined=combined_text,
                dl_t=gr.update(value=f_t, visible=True),
                dl_s=gr.update(value=f_s, visible=True),
                dl_r=gr.update(value=f_r, visible=True),
                dl_c=gr.update(value=f_c, visible=True),
                dl_j=gr.update(value=f_j, visible=True),
                dl_p=gr.update(value=f_p, visible=f_p is not None),
                dl_docx=gr.update(value=f_docx, visible=f_docx is not None),
                dl_srt=gr.update(value=f_srt, visible=f_srt is not None),
                dl_vtt=gr.update(value=f_vtt, visible=f_vtt is not None),
                dl_acc=gr.update(open=True),
                rs={"stem": stem, "combined_text": combined_text,
                    "detected_language": result.detected_language,
                    "out_dir": str(job_dir)},
                log=log_text,
            )
            break

        elif kind == "error":
            _write_job_status("error", stem=stem, job_id=job_id, job_dir=str(job_dir),
                              error=str(msg[1])[:200])
            if _JDB_OK:
                try:
                    _jdb.fail_job(job_id, str(msg[1]))
                except Exception:
                    pass
            log_text = _add_log(f"🚨 {msg[1]}", "error")
            yield _out(log=log_text)
            yield _err(f"Processing failed: {msg[1]}")
            break
    finally:
        _allow_sleep()


def toggle_speakers(is_panel):
    return gr.update(visible=is_panel)


# ── bandwidth monitor ──────────────────────────────────────────────────────────
_bw_state: dict = {}


def _get_bandwidth_html() -> str:
    if not _PSUTIL_OK:
        return ""
    now = time.time()
    counters = _psutil.net_io_counters()
    sent  = counters.bytes_sent
    recv  = counters.bytes_recv

    if _bw_state:
        dt   = now - _bw_state["ts"]
        ds   = (sent - _bw_state["sent"]) / max(dt, 0.001)
        dr   = (recv - _bw_state["recv"]) / max(dt, 0.001)
    else:
        ds = dr = 0.0

    _bw_state.update({"ts": now, "sent": sent, "recv": recv,
                      "session_sent":  _bw_state.get("session_sent",  sent),
                      "session_recv":  _bw_state.get("session_recv",  recv)})

    def _fmt(bps: float) -> str:
        if bps >= 1_048_576:  return f"{bps/1_048_576:.1f} MB/s"
        if bps >= 1_024:      return f"{bps/1_024:.0f} KB/s"
        return f"{bps:.0f} B/s"

    def _fmt_total(b: int) -> str:
        if b >= 1_073_741_824: return f"{b/1_073_741_824:.2f} GB"
        if b >= 1_048_576:     return f"{b/1_048_576:.1f} MB"
        if b >= 1_024:         return f"{b/1_024:.0f} KB"
        return f"{b} B"

    active = ds > 500 or dr > 500
    dot_color = "#22c55e" if active else "#64748b"

    total_sent = sent - _bw_state["session_sent"]
    total_recv = recv - _bw_state["session_recv"]

    return (
        '<div style="background:var(--ta-card-bg);border:1px solid var(--ta-card-border);'
        'border-radius:10px;padding:10px 14px;margin-top:10px;">'
        '<div style="display:flex;align-items:center;gap:7px;margin-bottom:8px;">'
        f'<span style="width:9px;height:9px;border-radius:50%;background:{dot_color};'
        f'display:inline-block;flex-shrink:0;{"animation:ta-pulse 1.2s infinite;" if active else ""}"></span>'
        '<span style="font-size:0.72em;font-weight:700;text-transform:uppercase;'
        'letter-spacing:0.08em;color:var(--ta-card-sub);">Network Activity</span>'
        '</div>'
        '<div style="display:grid;grid-template-columns:1fr 1fr;gap:6px;">'
        # Upload
        '<div style="background:var(--ta-stat-bg,rgba(255,255,255,0.7));border-radius:7px;padding:7px 10px;">'
        '<div style="font-size:0.68em;color:var(--ta-stat-label,#4338ca);font-weight:600;margin-bottom:2px;">↑ Upload</div>'
        f'<div style="font-size:0.95em;font-weight:700;color:var(--ta-stat-val,#1d4ed8);">{_fmt(ds)}</div>'
        f'<div style="font-size:0.68em;color:var(--ta-card-sub);">session: {_fmt_total(total_sent)}</div>'
        '</div>'
        # Download
        '<div style="background:var(--ta-stat-bg,rgba(255,255,255,0.7));border-radius:7px;padding:7px 10px;">'
        '<div style="font-size:0.68em;color:var(--ta-stat-label,#4338ca);font-weight:600;margin-bottom:2px;">↓ Download</div>'
        f'<div style="font-size:0.95em;font-weight:700;color:var(--ta-stat-val,#1d4ed8);">{_fmt(dr)}</div>'
        f'<div style="font-size:0.68em;color:var(--ta-card-sub);">session: {_fmt_total(total_recv)}</div>'
        '</div>'
        '</div>'
        '</div>'
    )


_VARIANT_LABELS = {
    "en": "English regional variant",
    "es": "Spanish regional variant",
    "fr": "French regional variant",
    "pt": "Portuguese regional variant",
    "de": "German regional variant",
    "it": "Italian regional variant",
    "zh": "Chinese regional variant",
    "ja": "Japanese dialect",
    "ko": "Korean dialect",
    "ar": "Arabic regional variant",
    "ru": "Russian regional variant",
    "hi": "Hindi dialect",
    "bn": "Bengali regional variant",
    "ta": "Tamil regional variant",
    "te": "Telugu dialect",
    "gu": "Gujarati regional variant",
    "kn": "Kannada dialect",
    "ml": "Malayalam dialect",
    "mr": "Marathi dialect",
    "pa": "Punjabi regional variant",
    "ur": "Urdu regional variant",
    "nl": "Dutch regional variant",
    "tr": "Turkish regional variant",
    "vi": "Vietnamese regional variant",
    "sv": "Swedish regional variant",
    "no": "Norwegian variant",
    "pl": "Polish regional variant",
    "th": "Thai regional variant",
    "el": "Greek regional variant",
    "he": "Hebrew regional variant",
    "ro": "Romanian regional variant",
    "hu": "Hungarian regional variant",
    "cs": "Czech regional variant",
    "fi": "Finnish regional variant",
    "da": "Danish regional variant",
    "uk": "Ukrainian regional variant",
    "id": "Indonesian regional variant",
}

def toggle_language_variant(lang):
    variants = LANGUAGE_VARIANTS.get(lang, [])
    if variants:
        return gr.update(choices=variants, value=variants[0][1], visible=True)
    return gr.update(choices=[], value=None, visible=False)


# ── build theme ────────────────────────────────────────────────────────────────
_THEME = gr.themes.Soft(
    primary_hue=gr.themes.colors.indigo,
    secondary_hue=gr.themes.colors.violet,
    neutral_hue=gr.themes.colors.zinc,
    font=[gr.themes.GoogleFont("Inter"), "ui-sans-serif", "system-ui", "sans-serif"],
    font_mono=[gr.themes.GoogleFont("JetBrains Mono"), "ui-monospace", "Courier New", "monospace"],
).set(
    body_background_fill="#f4f4f8",
    body_text_color="#111827",
    button_primary_background_fill="#6366f1",
    button_primary_background_fill_hover="#4f46e5",
    button_primary_text_color="white",
    button_primary_border_color="transparent",
    block_background_fill="#ffffff",
    block_border_color="#e5e7eb",
    block_border_width="1px",
    block_shadow="0 1px 4px 0 rgba(0,0,0,0.06)",
    block_radius="16px",
    block_label_text_weight="600",
    block_label_text_color="#6b7280",
    block_label_text_size="*text_sm",
    input_background_fill="#ffffff",
    input_border_color="#e5e7eb",
    panel_background_fill="#ffffff",
    panel_border_color="#e5e7eb",
)

# ── HTML snippets ───────────────────────────────────────────────────────────────
_HERO = """
<style>
@keyframes ta-hero-shift {
  0%{background-position:0% 50%} 50%{background-position:100% 50%} 100%{background-position:0% 50%}
}
@keyframes ta-hero-glow {
  0%,100%{opacity:0.5;transform:scale(1)} 50%{opacity:1;transform:scale(1.08)}
}
@keyframes ta-badge-in {
  from{opacity:0;transform:translateY(6px)} to{opacity:1;transform:translateY(0)}
}
</style>
<div id="ta-hero" style="position:relative;overflow:hidden;
     background:linear-gradient(135deg,#060614 0%,#150d3d 25%,#2d1b69 55%,#1e3a8a 80%,#0d1f4a 100%);
     background-size:400% 400%;animation:ta-hero-shift 14s ease infinite;
     border-radius:24px;padding:38px 48px 34px;color:#fff;margin-bottom:10px;
     box-shadow:0 12px 48px rgba(99,102,241,0.30),0 4px 12px rgba(0,0,0,0.35);">

  <!-- dot grid overlay -->
  <div style="position:absolute;inset:0;opacity:0.06;pointer-events:none;
       background-image:radial-gradient(circle,#fff 1px,transparent 1px);
       background-size:28px 28px;"></div>

  <!-- orb glows -->
  <div style="position:absolute;top:-50px;right:-30px;width:260px;height:260px;
       background:radial-gradient(circle,rgba(139,92,246,0.40) 0%,transparent 70%);
       animation:ta-hero-glow 7s ease-in-out infinite;pointer-events:none;"></div>
  <div style="position:absolute;bottom:-80px;left:25%;width:320px;height:320px;
       background:radial-gradient(circle,rgba(99,102,241,0.28) 0%,transparent 70%);
       animation:ta-hero-glow 9s ease-in-out infinite 2s;pointer-events:none;"></div>
  <div style="position:absolute;bottom:10px;right:15%;width:160px;height:160px;
       background:radial-gradient(circle,rgba(56,189,248,0.15) 0%,transparent 70%);
       animation:ta-hero-glow 11s ease-in-out infinite 4s;pointer-events:none;"></div>

  <!-- title row -->
  <div style="position:relative;display:flex;align-items:center;gap:20px;margin-bottom:20px;">
    <div style="width:60px;height:60px;border-radius:18px;flex-shrink:0;
         background:rgba(255,255,255,0.10);backdrop-filter:blur(12px);
         border:1px solid rgba(255,255,255,0.20);
         display:flex;align-items:center;justify-content:center;
         font-size:1.9em;box-shadow:0 4px 16px rgba(99,102,241,0.40);">🎙️</div>
    <div>
      <div style="font-size:2.1em;font-weight:900;letter-spacing:-0.05em;line-height:1;
           color:#fff!important;text-shadow:0 0 40px rgba(139,92,246,0.60);">Transcript Agent</div>
      <div style="color:#c4b5fd!important;font-size:0.88em;font-weight:500;margin-top:7px;
           letter-spacing:0.02em;display:flex;align-items:center;gap:8px;">
        <span style="display:inline-block;width:6px;height:6px;border-radius:50%;
              background:#4ade80;box-shadow:0 0 8px rgba(74,222,128,0.8);flex-shrink:0;"></span>
        AI-powered transcription &amp; analysis &mdash; Whisper &amp; Deepgram + Claude
      </div>
    </div>
  </div>

  <!-- feature badges -->
  <div style="position:relative;display:flex;gap:8px;flex-wrap:wrap;">
    <span style="background:rgba(255,255,255,0.09);backdrop-filter:blur(10px);border:1px solid rgba(255,255,255,0.18);border-radius:22px;padding:6px 15px;font-size:0.74em;font-weight:700;letter-spacing:0.03em;color:#fff!important;animation:ta-badge-in 0.5s ease 0.1s both;">🎵 Audio &amp; Video</span>
    <span style="background:rgba(255,255,255,0.09);backdrop-filter:blur(10px);border:1px solid rgba(255,255,255,0.18);border-radius:22px;padding:6px 15px;font-size:0.74em;font-weight:700;letter-spacing:0.03em;color:#fff!important;animation:ta-badge-in 0.5s ease 0.2s both;">📄 Documents</span>
    <span style="background:rgba(255,255,255,0.09);backdrop-filter:blur(10px);border:1px solid rgba(255,255,255,0.18);border-radius:22px;padding:6px 15px;font-size:0.74em;font-weight:700;letter-spacing:0.03em;color:#fff!important;animation:ta-badge-in 0.5s ease 0.3s both;">🗣️ Speaker Diarization</span>
    <span style="background:rgba(255,255,255,0.09);backdrop-filter:blur(10px);border:1px solid rgba(255,255,255,0.18);border-radius:22px;padding:6px 15px;font-size:0.74em;font-weight:700;letter-spacing:0.03em;color:#fff!important;animation:ta-badge-in 0.5s ease 0.4s both;">📊 Speech Analytics</span>
    <span style="background:rgba(99,102,241,0.25);backdrop-filter:blur(10px);border:1px solid rgba(139,92,246,0.40);border-radius:22px;padding:6px 15px;font-size:0.74em;font-weight:700;letter-spacing:0.03em;color:#e0d9ff!important;animation:ta-badge-in 0.5s ease 0.5s both;">🌐 37+ Languages</span>
  </div>
</div>
"""

_API_BANNER = """
<div id="api-banner" style="background:linear-gradient(135deg,rgba(255,251,235,0.95),rgba(254,249,238,0.95));
     backdrop-filter:blur(12px);border:1.5px solid #fbbf24;
     border-radius:18px;padding:14px 20px;display:flex;align-items:center;gap:14px;
     box-shadow:0 4px 16px rgba(251,191,36,0.14),0 1px 3px rgba(0,0,0,0.05);
     transition:all 0.35s ease;">
  <div id="api-banner-icon" style="font-size:1.5em;transition:all 0.3s;flex-shrink:0;">🔑</div>
  <div style="flex:1;min-width:0;">
    <div id="api-banner-title" style="font-weight:800;color:#92400e;font-size:0.85em;
         text-transform:uppercase;letter-spacing:0.05em;transition:color 0.3s;">API Key Required</div>
    <div id="api-banner-sub" style="color:#a16207;font-size:0.79em;margin-top:3px;transition:color 0.3s;line-height:1.5;">
      Enter your <strong>AI provider key</strong> below (Anthropic, OpenAI, Gemini, Groq, etc.).
      Billed to your account — nothing stored here.
    </div>
  </div>
  <div id="api-banner-badge" style="display:none;background:linear-gradient(135deg,#10b981,#059669);color:#fff;
       font-size:0.70em;font-weight:800;padding:5px 14px;border-radius:20px;letter-spacing:0.06em;
       box-shadow:0 3px 10px rgba(16,185,129,0.40);white-space:nowrap;flex-shrink:0;">✓ READY</div>
</div>
"""

_THEME_TOGGLE = """
<div id="ta-widget"
  style="position:fixed;top:14px;right:18px;z-index:9999;display:flex;align-items:center;
         background:rgba(255,255,255,0.92);backdrop-filter:blur(16px);-webkit-backdrop-filter:blur(16px);
         border:1px solid rgba(99,102,241,0.2);border-radius:32px;padding:4px;
         box-shadow:0 4px 16px rgba(0,0,0,0.10),0 0 0 1px rgba(255,255,255,0.8);gap:2px;">
  <button id="ta-btn-light" title="Light mode"
    style="display:flex;align-items:center;gap:5px;padding:5px 13px;border-radius:26px;
           border:none;cursor:pointer;font-size:0.82em;font-weight:600;
           background:#6366f1;color:#fff;transition:all 0.2s;
           box-shadow:0 2px 8px rgba(99,102,241,0.35);">
    ☀️ Light
  </button>
  <button id="ta-btn-dark" title="Dark mode"
    style="display:flex;align-items:center;gap:5px;padding:5px 13px;border-radius:26px;
           border:none;cursor:pointer;font-size:0.82em;font-weight:600;
           background:transparent;color:#6b7280;transition:all 0.2s;">
    🌙 Dark
  </button>
</div>
"""

# ── Theme JS — injected via gr.Blocks(js=...) which is the guaranteed execution
# path. gr.HTML uses Svelte {#html} which deliberately does NOT run <script> tags.
_THEME_JS = """
(function(){
  var _dark = false;

  /* ── PERMANENT static CSS injected directly into <head> ─────────────────────
     Gradio 6 embeds css=CSS as JSON data in <script> tags and injects it later
     via its own pipeline — we can't rely on it being in the DOM at toggle time.
     We inject all our static CSS here so it's guaranteed to be real CSS. */
  if (!document.getElementById('ta-static')) {
    var ps = document.createElement('style');
    ps.id = 'ta-static';
    ps.textContent = [
      /* Light mode page base */
      'body{background:#f4f4f8!important}',
      /* Custom checkboxes — visible in both modes */
      'input[type=checkbox]{-webkit-appearance:none!important;appearance:none!important;width:18px!important;height:18px!important;min-width:18px!important;border:2px solid #6366f1!important;border-radius:4px!important;background:#fff!important;cursor:pointer!important;position:relative!important;vertical-align:middle!important;flex-shrink:0!important}',
      'input[type=checkbox]:checked{background:#6366f1!important;border-color:#6366f1!important}',
      'input[type=checkbox]:checked::after{content:""!important;position:absolute!important;left:4px!important;top:1px!important;width:6px!important;height:10px!important;border:2px solid #fff!important;border-top:none!important;border-left:none!important;transform:rotate(45deg)!important;display:block!important}',
      'html.dark input[type=checkbox]{background:#1e1e2a!important;border-color:#a5b4fc!important}',
      'html.dark input[type=checkbox]:checked{background:#6366f1!important;border-color:#6366f1!important}',
      '.checkbox-wrap{align-items:center!important;gap:8px!important}',
      /* CSS vars — light defaults for step tracker + ETA panel */
      ':root{--ta-card-bg:#ffffff;--ta-card-border:#e0e0f0;--ta-card-text:#1e1e2a;--ta-card-sub:#6b6b80;--ta-card-val:#1e1e2a;',
      '--ta-step-done-bg:#ede9fe;--ta-step-done-bdr:#7c3aed;--ta-step-done-clr:#5b21b6;',
      '--ta-step-act-bg:#eef2ff;--ta-step-act-bdr:#6366f1;--ta-step-act-clr:#4338ca;',
      '--ta-step-wait-bg:#f4f4f8;--ta-step-wait-bdr:#e0e0f0;--ta-step-wait-clr:#9090a8;',
      '--ta-conn-line-done:#7c3aed;--ta-conn-line-wait:#e0e0f0;--ta-stat-bg:rgba(255,255,255,0.85);',
      '--ta-stat-label:#4338ca;--ta-stat-val:#6366f1;',
      '--ta-log-bg:#f8f8fc;--ta-log-border:#e0e0f0;--ta-log-ts:#64748b;--ta-log-hdr:#334155}',
      /* CSS vars — dark overrides */
      'html.dark{--ta-card-bg:#1e1e2a;--ta-card-border:#2e2e42;--ta-card-text:#e8e8f0;--ta-card-sub:#9090a8;--ta-card-val:#e8e8f0;',
      '--ta-step-done-bg:#1e1040;--ta-step-done-bdr:#7c3aed;--ta-step-done-clr:#c4b5fd;',
      '--ta-step-act-bg:#1e1e40;--ta-step-act-bdr:#a5b4fc;--ta-step-act-clr:#c7d2fe;',
      '--ta-step-wait-bg:#0a0a12;--ta-step-wait-bdr:#2e2e42;--ta-step-wait-clr:#9090a8;',
      '--ta-conn-line-done:#7c3aed;--ta-conn-line-wait:#2e2e42;--ta-stat-bg:rgba(10,10,18,0.85);',
      '--ta-stat-label:#a5b4fc;--ta-stat-val:#e8e8f0;',
      '--ta-log-bg:#0a0a12;--ta-log-border:#1e1e40;--ta-log-ts:#64748b;--ta-log-hdr:#e2e8f0}',
      /* Process button */
      '.big-btn button{background:linear-gradient(135deg,#059669,#10b981)!important;color:#fff!important;font-size:1.08em!important;font-weight:700!important;border:none!important;border-radius:10px!important;padding:15px!important;width:100%!important;box-shadow:0 4px 14px rgba(16,185,129,0.45)!important}',
      /* Scrollable dropdowns */
      '[role=listbox]{max-height:220px!important;overflow-y:auto!important}',
      '#provider-sel [role=listbox],#model-sel [role=listbox]{max-height:280px!important;overflow-y:auto!important}',
      /* Live log terminal */
      '#live-log textarea{background:#0f172a!important;color:#86efac!important;font-family:"Courier New",monospace!important;font-size:0.80em!important;border-color:#1e3a5f!important}',
      /* Bandwidth pulse animation */
      '@keyframes ta-pulse{0%,100%{opacity:1}50%{opacity:0.3}}',
    ].join('');
    document.head.appendChild(ps);
  }

  /* ── Dynamic dark/light override tag ─────────────────────────────────────────
     Re-appended to END of <head> on every toggle so it always wins cascade ties
     against Gradio's own stylesheets (last stylesheet wins when both !important). */
  var st = document.getElementById('ta-override') || document.createElement('style');
  st.id = 'ta-override';

  /* Gradio 6 CSS variable overrides + element rules. Every color/bg has !important
     so Gradio inline styles can't win. */
  var DARK_RULES = [
    /* CSS variables — set on html.dark so they cascade into all Gradio components */
    'html.dark{color-scheme:dark;color:#e8e8f0!important;background:#0a0a12!important;'
      +'--body-background-fill:#0a0a12;--background-fill-primary:#0a0a12;'
      +'--background-fill-secondary:#1e1e2a;--block-background-fill:#1e1e2a;'
      +'--block-border-color:#2e2e42;--block-label-text-color:#9090a8;'
      +'--input-background-fill:#14141e;--input-border-color:#2e2e42;'
      +'--panel-background-fill:#1e1e2a;--panel-border-color:#2e2e42;'
      +'--border-color-primary:#2e2e42;--body-text-color:#e8e8f0;'
      +'--body-text-color-subdued:#9090a8;--neutral-100:#1e1e2a;--neutral-200:#28283a;'
      +'--neutral-700:#9090a8;--neutral-800:#c4c4d4;--neutral-900:#e8e8f0;}',
    /* page & containers */
    'html.dark body,html.dark .gradio-container,html.dark .main,html.dark .contain{background:#0a0a12!important;color:#e8e8f0!important}',
    /* blocks */
    'html.dark .block,html.dark .form,html.dark .wrap,html.dark .panel-full-width,html.dark .compact,html.dark .upload-container,html.dark .padded{background:#1e1e2a!important;border-color:#2e2e42!important}',
    /* text */
    'html.dark span,html.dark p,html.dark div,html.dark h1,html.dark h2,html.dark h3,html.dark h4,html.dark li,html.dark td,html.dark th,html.dark strong,html.dark em{color:#e8e8f0!important}',
    /* labels */
    'html.dark .label-wrap span,html.dark .block-label,html.dark label>span,html.dark .info,html.dark .file-name{color:#9090a8!important}',
    /* inputs */
    'html.dark input,html.dark textarea,html.dark select,[role=combobox]{background:#14141e!important;color:#e8e8f0!important;border-color:#2e2e42!important}',
    'html.dark input::placeholder,html.dark textarea::placeholder{color:#9090a8!important;opacity:1!important}',
    /* tabs */
    'html.dark .tabs>.tab-nav button{color:#9090a8!important;background:#1e1e2a!important;border-color:#2e2e42!important}',
    'html.dark .tabs>.tab-nav button.selected{color:#fff!important;background:#6366f1!important;border-color:#6366f1!important}',
    'html.dark .tabitem{background:#0a0a12!important}',
    /* markdown */
    'html.dark .prose,html.dark .markdown{color:#e8e8f0!important;background:transparent!important}',
    'html.dark .prose *,html.dark .markdown *{color:#e8e8f0!important}',
    'html.dark .prose a,html.dark .markdown a{color:#a5b4fc!important}',
    'html.dark .prose code,html.dark .markdown code{background:#14141e!important;color:#86efac!important}',
    /* dropdowns */
    'html.dark [role=listbox]{background:#1e1e2a!important;border-color:#2e2e42!important}',
    'html.dark [role=option]{color:#e8e8f0!important;background:#1e1e2a!important}',
    'html.dark [role=option]:hover,html.dark [role=option][aria-selected=true]{background:#28283a!important;color:#fff!important}',
    /* accordion */
    'html.dark .accordion,html.dark details{background:#1e1e2a!important;border-color:#2e2e42!important}',
    'html.dark .accordion .label-wrap,html.dark details summary{color:#e8e8f0!important}',
    'html.dark .checkbox-group label span,html.dark .radio-group label span{color:#e8e8f0!important}',
    'html.dark .file-preview{background:#1e1e2a!important;color:#e8e8f0!important}',
    'html.dark .dropdown-arrow svg{fill:#9090a8!important}',
    /* buttons */
    'html.dark button{background:#1e1e2a!important;border-color:#2e2e42!important;color:#e8e8f0!important}',
    'html.dark button.selected{background:#28283a!important}',
    'html.dark .big-btn button{background:linear-gradient(135deg,#047857,#10b981)!important;color:#fff!important;border:none!important}',
    /* theme toggle — restore correct colors */
    'html.dark #ta-btn-light{background:transparent!important;color:#9090a8!important}',
    'html.dark #ta-btn-dark{background:#6366f1!important;color:#fff!important}',
    /* scrollbars */
    'html.dark ::-webkit-scrollbar-track{background:#0a0a12!important}',
    'html.dark ::-webkit-scrollbar-thumb{background:#2e2e42!important}',
    'html.dark ::-webkit-scrollbar-thumb:hover{background:#6366f1!important}',
  ].join('');

  /* ── DOM patcher ─────────────────────────────────────────────────────────────
     Sets/removes inline styles directly on elements Gradio styles inline.
     Uses removeProperty() in light mode — setProperty('x','') is a no-op. */
  function _sp(el, prop, val) {
    if (val) el.style.setProperty(prop, val, 'important');
    else     el.style.removeProperty(prop);
  }

  /* _SKIP: elements whose colors we manage separately */
  function _skip(el) {
    return el.closest('#ta-widget') || el.closest('#api-banner') || el.closest('#ta-wake-notice');
  }

  function patchDOM(dark) {
    var bg0 = dark ? '#0a0a12' : null;
    var bg1 = dark ? '#1e1e2a' : null;
    var fg  = dark ? '#e8e8f0' : null;
    var fg2 = dark ? '#9090a8' : null;
    var bd  = dark ? '#2e2e42' : null;

    /* Page containers */
    document.querySelectorAll('.gradio-container,.main,.contain,body').forEach(function(el){
      _sp(el,'background',bg0); _sp(el,'color',fg);
    });

    /* Blocks, forms, accordions — backgrounds */
    document.querySelectorAll('.block,.form,.wrap,.panel-full-width,.compact,details,summary').forEach(function(el){
      _sp(el,'background',bg1); _sp(el,'border-color',bd);
      if(fg) _sp(el,'color',fg);
    });

    /* Inputs */
    document.querySelectorAll('input,textarea,select').forEach(function(el){
      _sp(el,'background',dark?'#14141e':null); _sp(el,'color',fg); _sp(el,'border-color',dark?'#2e2e42':null);
    });

    /* Nuclear text patch — every text node inside .gradio-container gets the
       correct color. We skip: images, buttons, our own widgets, and the banner.
       This ensures accordion labels, info text, step numbers, etc. are visible. */
    var SKIP_TAGS = {IMG:1,VIDEO:1,CANVAS:1,SVG:1,PATH:1,INPUT:1,TEXTAREA:1,SELECT:1};
    document.querySelectorAll('.gradio-container *').forEach(function(el){
      if (SKIP_TAGS[el.tagName]) return;
      if (_skip(el)) return;
      /* Labels/info get subdued color; everything else gets full white */
      var isLabel = el.classList.contains('block-label') ||
                    (el.parentElement && (el.parentElement.classList.contains('label-wrap') ||
                                          el.parentElement.classList.contains('info')));
      _sp(el, 'color', dark ? (isLabel ? '#9090a8' : '#e8e8f0') : null);
    });

    /* Dropdowns specifically */
    document.querySelectorAll('[role=listbox],[role=option],[role=combobox]').forEach(function(el){
      _sp(el,'background',bg1); _sp(el,'color',fg); _sp(el,'border-color',bd);
    });

    /* Big button — keep it indigo */
    document.querySelectorAll('.big-btn button').forEach(function(el){
      _sp(el,'background',dark?'linear-gradient(135deg,#4f46e5,#6366f1)':null);
      _sp(el,'color',dark?'#fff':null);
    });
  }

  /* ── Gradio CSS variable names — set as inline props on <html> ──────────────
     Inline custom properties on documentElement beat ALL :root CSS rules,
     which is how Gradio reads them. This is the only approach that reliably
     overrides Gradio's Soft theme variables regardless of specificity. */
  var DARK_VARS = {
    '--body-background-fill':      '#0a0a12',
    '--background-fill-primary':   '#0a0a12',
    '--background-fill-secondary': '#1e1e2a',
    '--block-background-fill':     '#1e1e2a',
    '--input-background-fill':     '#14141e',
    '--panel-background-fill':     '#1e1e2a',
    '--chatbot-background-fill':   '#1e1e2a',
    '--body-text-color':           '#e8e8f0',
    '--block-label-text-color':    '#9090a8',
    '--block-title-text-color':    '#e8e8f0',
    '--block-info-text-color':     '#9090a8',
    '--block-border-color':        '#2e2e42',
    '--block-border-width':        '1px',
    '--input-border-color':        '#2e2e42',
    '--border-color-primary':      '#2e2e42',
    '--border-color-accent':       '#a5b4fc',
    '--neutral-100':               '#1e1e2a',
    '--neutral-200':               '#28283a',
    '--neutral-300':               '#2e2e42',
    '--neutral-400':               '#6b6b80',
    '--neutral-500':               '#9090a8',
    '--neutral-600':               '#c4c4d4',
    '--neutral-700':               '#e8e8f0',
    '--neutral-800':               '#f0f0ff',
    '--neutral-900':               '#f8f8fc',
    '--color-accent':              '#a5b4fc',
    '--link-text-color':           '#a5b4fc',
    '--shadow-drop':               '0 1px 3px rgba(0,0,0,0.6)',
  };

  function setGradioVars(dark) {
    var root = document.documentElement;
    Object.keys(DARK_VARS).forEach(function(k) {
      if (dark) root.style.setProperty(k, DARK_VARS[k]);
      else      root.style.removeProperty(k);
    });
  }

  /* ── Core apply function ─────────────────────────────────────────────────── */
  function applyTheme(dark) {
    _dark = dark;

    /* 1. Set Gradio CSS vars inline on <html> — beats all :root theme rules */
    setGradioVars(dark);

    /* 2. Toggle .dark class for our html.dark CSS rules */
    document.documentElement.classList.toggle('dark', dark);
    document.body.classList.toggle('dark', dark);

    /* 3. Re-append override sheet last in <head> */
    if (st.parentNode) st.parentNode.removeChild(st);
    document.head.appendChild(st);
    st.textContent = dark ? DARK_RULES : '';

    /* 4. Patch inline styles (handles elements Gradio styles inline) */
    patchDOM(dark);

    /* Direct body/html inline styles — these beat everything */
    _sp(document.body, 'background', dark ? '#0a0a12' : null);
    _sp(document.body, 'color',      dark ? '#e8e8f0' : null);
    _sp(document.documentElement, 'background', dark ? '#0a0a12' : null);

    localStorage.setItem('ta-dark',      dark ? 'true'  : 'false');
    localStorage.setItem('theme',        dark ? 'dark'  : 'light');
    localStorage.setItem('gradio-theme', dark ? 'dark'  : 'light');

    /* Update button visuals */
    var bl = document.getElementById('ta-btn-light');
    var bd = document.getElementById('ta-btn-dark');
    var wg = document.getElementById('ta-widget');
    if (bl && bd) {
      bl.style.background = dark ? 'transparent' : '#6366f1';
      bl.style.color      = dark ? '#9090a8'     : '#fff';
      bd.style.background = dark ? '#6366f1'     : 'transparent';
      bd.style.color      = dark ? '#fff'        : '#6b6b80';
    }
    if (wg) {
      wg.style.background  = dark ? 'rgba(10,10,18,0.95)' : 'rgba(255,255,255,0.95)';
      wg.style.borderColor = dark ? '#2e2e42' : '#e0e0f0';
    }

    /* API banner */
    var banner = document.getElementById('api-banner');
    var title  = document.getElementById('api-banner-title');
    var sub    = document.getElementById('api-banner-sub');
    if (banner && !banner.dataset.state) {
      banner.style.background  = dark ? 'linear-gradient(135deg,#292107,#3b2d00)' : 'linear-gradient(135deg,#fffbeb,#fef3c7)';
      banner.style.borderColor = dark ? '#d97706' : '#f59e0b';
      if (title) title.style.color = dark ? '#fbbf24' : '#92400e';
      if (sub)   sub.style.color   = dark ? '#fcd34d' : '#a16207';
    }
  }

  /* ── MutationObservers ───────────────────────────────────────────────────────
     1. Prevent Gradio stripping .dark off <html>
     2. Watch <head> — when Gradio injects a new <style> after ta-override,
        immediately move ta-override back to the END so our rules win cascade
     3. Re-patch DOM when Gradio adds new body nodes */
  new MutationObserver(function(muts) {
    if (!_dark) return;
    muts.forEach(function(m) {
      if (m.attributeName === 'class' && !m.target.classList.contains('dark'))
        m.target.classList.add('dark');
    });
  }).observe(document.documentElement, { attributes: true, attributeFilter: ['class'] });

  /* Watch <head> for new <style> injections — move ta-override to end immediately */
  new MutationObserver(function(muts) {
    if (!_dark) return;
    var newStyle = muts.some(function(m){
      return Array.prototype.some.call(m.addedNodes, function(n){
        return n.nodeType === 1 && n.tagName === 'STYLE' && n.id !== 'ta-override' && n.id !== 'ta-static';
      });
    });
    if (newStyle && st.parentNode) {
      st.parentNode.removeChild(st);
      document.head.appendChild(st);
    }
  }).observe(document.head, { childList: true });

  var _patchTimer = null;
  new MutationObserver(function(muts) {
    if (!_dark) return;
    var hasNodes = muts.some(function(m){ return m.addedNodes.length > 0; });
    if (!hasNodes) return;
    if (_patchTimer) return;
    _patchTimer = setTimeout(function(){ _patchTimer = null; setGradioVars(true); patchDOM(true); }, 50);
  }).observe(document.body || document.documentElement, { childList: true, subtree: true });

  /* Periodic re-apply in dark mode — catches any Gradio re-renders we miss */
  setInterval(function() {
    if (!_dark) return;
    /* Re-append ta-override to ensure it stays last */
    if (st.parentNode && st.parentNode.lastChild !== st) {
      st.parentNode.removeChild(st);
      document.head.appendChild(st);
    }
    setGradioVars(true);
  }, 1000);

  /* ── Init — single init guard prevents double event-listener bug ─────────── */
  var _inited = false;
  function init() {
    if (_inited) return;
    var bl = document.getElementById('ta-btn-light');
    var bd = document.getElementById('ta-btn-dark');
    if (!bl || !bd) { setTimeout(init, 250); return; }
    _inited = true;
    applyTheme(localStorage.getItem('ta-dark') === 'true');
    bl.addEventListener('click', function() { applyTheme(false); });
    bd.addEventListener('click', function() { applyTheme(true);  });
  }
  /* Use event delegation as belt-and-suspenders so re-mounts can't break it */
  document.addEventListener('click', function(e) {
    var t = e.target.closest('#ta-btn-light,#ta-btn-dark');
    if (!t) return;
    applyTheme(t.id === 'ta-btn-dark');
  });
  document.readyState === 'loading'
    ? document.addEventListener('DOMContentLoaded', function() { setTimeout(init, 150); })
    : setTimeout(init, 150);

  /* ── 👁 Password show/hide eye ─────────────────────────────────────────── */
  function addEyes() {
    document.querySelectorAll('input[type="password"]').forEach(function(inp) {
      if (inp.dataset.eye) return;
      inp.dataset.eye = '1';
      var eye = document.createElement('button');
      eye.type = 'button';
      eye.textContent = '👁';
      eye.style.cssText = 'position:absolute;right:10px;top:50%;transform:translateY(-50%);'
        + 'background:none;border:none;cursor:pointer;font-size:1em;opacity:0.5;z-index:20;padding:2px;';
      eye.addEventListener('click', function() {
        var show = inp.type === 'password';
        inp.type = show ? 'text' : 'password';
        eye.textContent  = show ? '🙈' : '👁';
        eye.style.opacity = show ? '0.9' : '0.5';
      });
      inp.parentElement.style.position = 'relative';
      inp.parentElement.appendChild(eye);
    });
  }
  setTimeout(addEyes, 600);
  setTimeout(addEyes, 1800);
  setTimeout(addEyes, 3500);

  /* ── 🔑 API key banner — multi-provider aware ───────────────────────────── */
  var KEY_PROVIDERS = [
    { prefix: 'sk-ant',  name: 'Anthropic'   },
    { prefix: 'sk-',     name: 'OpenAI'      },
    { prefix: 'AIzaSy',  name: 'Google Gemini' },
    { prefix: 'gsk_',    name: 'Groq'        },
    { prefix: 'pplx-',   name: 'Perplexity'  },
    { prefix: 'dg-',     name: 'Deepgram'    },
    { prefix: 'ollama',  name: 'Ollama'      },
  ];

  function detectProvider(v) {
    for (var i = 0; i < KEY_PROVIDERS.length; i++) {
      if (v.startsWith(KEY_PROVIDERS[i].prefix)) return KEY_PROVIDERS[i].name;
    }
    return null;
  }

  function watchApiKey() {
    var inp = document.querySelector('input[type="password"]');
    if (!inp) { setTimeout(watchApiKey, 600); return; }

    /* Use setProperty(...,'important') so banner colors win over our broad
       html.dark div{color:...!important} rule that would otherwise make
       light text invisible on light-green backgrounds. */
    function _bs(el, prop, val) {
      if (el) el.style.setProperty(prop, val, 'important');
    }

    function refreshBanner() {
      var v        = inp.value.trim();
      var banner   = document.getElementById('api-banner');
      var icon     = document.getElementById('api-banner-icon');
      var title    = document.getElementById('api-banner-title');
      var sub      = document.getElementById('api-banner-sub');
      var badge    = document.getElementById('api-banner-badge');
      var provider = detectProvider(v);

      if (provider || v.length >= 20) {
        /* ── Approved ── */
        banner.dataset.state = 'approved';
        var label = provider ? (provider + ' Key Approved') : 'API Key Set';
        var detail = provider
          ? ('Your <strong>' + provider + '</strong> key is set. Ready to <strong>Analyze File</strong>.')
          : 'Key entered. Ready to <strong>Analyze File</strong>.';
        if (_dark) {
          _bs(banner, 'background', 'linear-gradient(135deg,#052e16,#14532d)');
          _bs(banner, 'border-color', '#22c55e');
          if (title) { title.textContent = label; _bs(title, 'color', '#4ade80'); }
          if (sub)   { sub.innerHTML = detail;    _bs(sub,   'color', '#86efac'); }
        } else {
          _bs(banner, 'background', 'linear-gradient(135deg,#f0fdf4,#dcfce7)');
          _bs(banner, 'border-color', '#22c55e');
          if (title) { title.textContent = label; _bs(title, 'color', '#166534'); }
          if (sub)   { sub.innerHTML = detail;    _bs(sub,   'color', '#15803d'); }
        }
        if (icon)  icon.textContent = '✅';
        if (badge) badge.style.display = 'block';

      } else if (v.length > 0) {
        /* ── Too short ── */
        banner.dataset.state = 'error';
        _bs(banner, 'background',   'linear-gradient(135deg,#fef2f2,#fee2e2)');
        _bs(banner, 'border-color', '#ef4444');
        if (icon)  icon.textContent = '⚠️';
        if (title) { title.textContent = 'Key Too Short'; _bs(title, 'color', '#991b1b'); }
        if (sub)   { sub.innerHTML = 'Paste your full API key (Anthropic, OpenAI, Gemini, Groq, etc.)'; _bs(sub, 'color', '#b91c1c'); }
        if (badge) badge.style.display = 'none';

      } else {
        /* ── Empty ── */
        delete banner.dataset.state;
        if (_dark) {
          _bs(banner, 'background',   'linear-gradient(135deg,#292107,#3b2d00)');
          _bs(banner, 'border-color', '#d97706');
          if (title) { title.textContent = 'API Key Required'; _bs(title, 'color', '#fbbf24'); }
          if (sub)   { sub.innerHTML = 'Enter your API key below. Billed to your account — nothing stored here.'; _bs(sub, 'color', '#fcd34d'); }
        } else {
          _bs(banner, 'background',   'linear-gradient(135deg,#fffbeb,#fef3c7)');
          _bs(banner, 'border-color', '#f59e0b');
          if (title) { title.textContent = 'API Key Required'; _bs(title, 'color', '#92400e'); }
          if (sub)   { sub.innerHTML = 'Enter your API key below (Anthropic, OpenAI, Gemini, Groq, etc.). Billed to your account — nothing stored here.'; _bs(sub, 'color', '#a16207'); }
        }
        if (icon)  icon.textContent = '🔑';
        if (badge) badge.style.display = 'none';
      }
    }

    /* Re-run banner colors whenever dark mode changes */
    var _origApply = applyTheme;
    applyTheme = function(dark) { _origApply(dark); setTimeout(refreshBanner, 50); };

    inp.addEventListener('input',  refreshBanner);
    inp.addEventListener('change', refreshBanner);
    refreshBanner();
  }
  setTimeout(watchApiKey, 800);
  setTimeout(watchApiKey, 2200);

  /* ── 💤 Wake detection — passive indicator, NO auto-reload ─────────────────
     Server-side sleep prevention keeps the CPU running and the connection alive.
     If the system DID suspend (overriding our prevention), we just show a
     "Processing continued in background" notice — no forced reload. */
  (function(){
    var _last = Date.now();
    setInterval(function(){
      var now = Date.now();
      var gap = now - _last;
      _last = now;
      if (gap > 10000) {
        /* Computer was suspended/resumed — show a non-intrusive notice */
        var existing = document.getElementById('ta-wake-notice');
        if (existing) return;
        var n = document.createElement('div');
        n.id = 'ta-wake-notice';
        n.style.cssText = 'position:fixed;bottom:20px;right:20px;z-index:99999;'
          + 'background:#0f172a;color:#e2e8f0;padding:12px 18px;border-radius:10px;'
          + 'font-family:sans-serif;font-size:0.85em;border:1px solid #334155;'
          + 'box-shadow:0 4px 20px rgba(0,0,0,0.5);max-width:280px;';
        n.innerHTML = '💤 Woke from sleep — processing continued in background.'
          + '<br><span style="color:#64748b;font-size:0.85em;">Results will appear when complete.</span>'
          + '<button onclick="this.parentElement.remove()" style="float:right;margin-left:8px;'
          + 'background:none;border:none;color:#94a3b8;cursor:pointer;font-size:1.1em;">✕</button>';
        document.body.appendChild(n);
        /* auto-hide after 8s */
        setTimeout(function(){ if(n.parentElement) n.remove(); }, 8000);
      }
    }, 3000);
  })();

  /* ── 🔄 Background job tracker ──────────────────────────────────────────────
     Saves the current job state to localStorage so if the page crashes or
     the user navigates away, a banner shows on return: "Job still running".
     app.py Python side sets 'ta-job-active' when processing starts/ends. */
  (function(){
    function showJobBanner(msg, sub) {
      if (document.getElementById('ta-job-banner')) return;
      var b = document.createElement('div');
      b.id = 'ta-job-banner';
      b.style.cssText = 'position:fixed;bottom:20px;left:50%;transform:translateX(-50%);'
        + 'z-index:99998;background:#1e1e40;color:#e8e8f0;padding:14px 22px;border-radius:12px;'
        + 'font-family:sans-serif;font-size:0.88em;font-weight:600;border:1px solid #6366f1;'
        + 'box-shadow:0 4px 20px rgba(0,0,0,0.5);display:flex;align-items:center;gap:12px;max-width:380px;';
      b.innerHTML = '<span style="font-size:1.3em;">⚙️</span>'
        + '<div><div>' + msg + '</div>'
        + '<div style="color:#93c5fd;font-size:0.82em;margin-top:2px;">' + sub + '</div></div>'
        + '<button onclick="this.parentElement.remove()" style="margin-left:auto;background:none;'
        + 'border:none;color:#94a3b8;cursor:pointer;font-size:1.1em;flex-shrink:0;">✕</button>';
      document.body.appendChild(b);
      setTimeout(function(){ if(b.parentElement) b.remove(); }, 12000);
    }

    /* On page load: check if a job was running when the page was last closed */
    var jobInfo = localStorage.getItem('ta-job-running');
    if (jobInfo) {
      try {
        var j = JSON.parse(jobInfo);
        var elapsed = Math.round((Date.now() - j.started) / 1000);
        var mins = Math.floor(elapsed / 60), secs = elapsed % 60;
        showJobBanner(
          'Processing continued in background',
          (j.file || 'File') + ' — ' + mins + 'm ' + secs + 's elapsed. Check outputs or wait for results.'
        );
      } catch(e) {}
    }

    /* ── Screen Wake Lock ─────────────────────────────────────────────────────
       Prevents the display from sleeping during transcription.
       While the screen stays on, Windows will not enter deep sleep.
       Browser auto-releases the lock when the tab is hidden — we re-acquire
       when the tab becomes visible again if a job is still running. */
    var _taWakeLock = null;

    window._taAcquireWakeLock = async function() {
      if (!('wakeLock' in navigator)) return;
      try {
        _taWakeLock = await navigator.wakeLock.request('screen');
        var n = document.getElementById('ta-wake-notice');
        if (n) n.style.display = 'flex';
      } catch(e) {}
    };

    window._taReleaseWakeLock = async function() {
      if (_taWakeLock) { try { await _taWakeLock.release(); } catch(e) {} _taWakeLock = null; }
      var n = document.getElementById('ta-wake-notice');
      if (n) n.style.display = 'none';
    };

    /* Watch status bar — release wake lock when job finishes or errors */
    var _taStatusObs = new MutationObserver(function() {
      var bar = document.querySelector('.status-bar, [data-testid="status-bar"], #status-bar');
      if (!bar) bar = document.querySelector('.eta-panel, .gradio-html');
      var txt = document.body.innerText || '';
      if ((txt.includes('All tabs are ready') || txt.includes('Processing failed')) && _taWakeLock) {
        window._taReleaseWakeLock();
        localStorage.removeItem('ta-job-running');
      }
    });
    setTimeout(function() {
      _taStatusObs.observe(document.body, { childList: true, subtree: true, characterData: true });
    }, 2000);

    document.addEventListener('visibilitychange', function() {
      if (document.visibilityState === 'visible' && !_taWakeLock) {
        if (localStorage.getItem('ta-job-running')) window._taAcquireWakeLock();
      }
    });

    window.taJobStart = function(filename) {
      localStorage.setItem('ta-job-running', JSON.stringify({ file: filename, started: Date.now() }));
      window._taAcquireWakeLock();
    };
    window.taJobEnd = function() {
      localStorage.removeItem('ta-job-running');
      window._taReleaseWakeLock();
      var b = document.getElementById('ta-job-banner');
      if (b) b.remove();
    };

    /* ── History tab: one-click row Load ── */
    window.taLoadJob = function(jid) {
      var box = document.querySelector('#history-job-id-input input, #history-job-id-input textarea');
      if (!box) return;
      var proto = (box.tagName === 'TEXTAREA')
                  ? window.HTMLTextAreaElement.prototype
                  : window.HTMLInputElement.prototype;
      var desc = Object.getOwnPropertyDescriptor(proto, 'value');
      if (desc && desc.set) desc.set.call(box, jid);
      else box.value = jid;
      box.dispatchEvent(new Event('input', {bubbles: true}));
      setTimeout(function() {
        var btn = document.querySelector('#history-load-btn button');
        if (btn) btn.click();
      }, 120);
    };
  })();
})();
"""

_IDLE_STATUS = """
<div style="background:linear-gradient(135deg,#1a1440 0%,#2d1b69 50%,#1e3a8a 100%);
     border-radius:20px;padding:24px 28px;display:flex;align-items:center;gap:20px;
     box-shadow:0 8px 32px rgba(99,102,241,0.30),0 2px 8px rgba(0,0,0,0.20);
     position:relative;overflow:hidden;">
  <div style="position:absolute;top:-30px;right:-20px;width:160px;height:160px;
       background:radial-gradient(circle,rgba(139,92,246,0.30) 0%,transparent 70%);
       pointer-events:none;"></div>
  <div style="width:52px;height:52px;border-radius:16px;flex-shrink:0;
       background:rgba(255,255,255,0.12);backdrop-filter:blur(8px);
       display:flex;align-items:center;justify-content:center;font-size:1.6em;
       border:1px solid rgba(255,255,255,0.15);">📂</div>
  <div style="position:relative;">
    <div style="color:#fff;font-size:1.05em;font-weight:800;letter-spacing:-0.01em;">Ready to process</div>
    <div style="color:#c4b5fd;font-size:0.83em;margin-top:5px;line-height:1.5;">
      Upload a file on the left, then click
      <span style="color:#fff;background:rgba(255,255,255,0.18);backdrop-filter:blur(4px);
            padding:2px 10px;border-radius:6px;font-weight:700;border:1px solid rgba(255,255,255,0.25);">
        Analyze File
      </span>
    </div>
  </div>
</div>
"""

_FORMATS = """
<div style="background:var(--ta-card-bg);border:1px solid var(--ta-card-border);
     border-radius:10px;padding:14px 16px;font-family:sans-serif;">
  <div style="display:flex;flex-direction:column;gap:8px;">
    <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
      <span style="font-size:0.78em;font-weight:700;text-transform:uppercase;
            letter-spacing:0.06em;color:var(--ta-card-sub);min-width:38px;">🎵 Audio</span>
      <span style="background:var(--ta-step-act-bg);color:var(--ta-step-act-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">mp3</span>
      <span style="background:var(--ta-step-act-bg);color:var(--ta-step-act-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">wav</span>
      <span style="background:var(--ta-step-act-bg);color:var(--ta-step-act-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">m4a</span>
      <span style="background:var(--ta-step-act-bg);color:var(--ta-step-act-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">flac</span>
      <span style="background:var(--ta-step-act-bg);color:var(--ta-step-act-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">ogg</span>
      <span style="background:var(--ta-step-act-bg);color:var(--ta-step-act-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">aac</span>
    </div>
    <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
      <span style="font-size:0.78em;font-weight:700;text-transform:uppercase;
            letter-spacing:0.06em;color:var(--ta-card-sub);min-width:38px;">🎬 Video</span>
      <span style="background:var(--ta-step-done-bg);color:var(--ta-step-done-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">mp4</span>
      <span style="background:var(--ta-step-done-bg);color:var(--ta-step-done-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">mov</span>
      <span style="background:var(--ta-step-done-bg);color:var(--ta-step-done-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">avi</span>
      <span style="background:var(--ta-step-done-bg);color:var(--ta-step-done-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">mkv</span>
      <span style="background:var(--ta-step-done-bg);color:var(--ta-step-done-clr);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">webm</span>
    </div>
    <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
      <span style="font-size:0.78em;font-weight:700;text-transform:uppercase;
            letter-spacing:0.06em;color:var(--ta-card-sub);min-width:38px;">📄 Docs</span>
      <span style="background:var(--ta-step-wait-bg);color:var(--ta-card-text);
            border:1px solid var(--ta-card-border);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">pdf</span>
      <span style="background:var(--ta-step-wait-bg);color:var(--ta-card-text);
            border:1px solid var(--ta-card-border);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">docx</span>
      <span style="background:var(--ta-step-wait-bg);color:var(--ta-card-text);
            border:1px solid var(--ta-card-border);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">txt</span>
      <span style="background:var(--ta-step-wait-bg);color:var(--ta-card-text);
            border:1px solid var(--ta-card-border);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">md</span>
      <span style="background:var(--ta-step-wait-bg);color:var(--ta-card-text);
            border:1px solid var(--ta-card-border);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">srt</span>
      <span style="background:var(--ta-step-wait-bg);color:var(--ta-card-text);
            border:1px solid var(--ta-card-border);
            font-size:0.75em;font-weight:600;padding:2px 8px;border-radius:5px;">vtt</span>
    </div>
  </div>
</div>
"""

def _SECTION(label: str) -> str:
    import re as _re
    m = _re.match(r"Step\s+(\d+)\s*[—-]\s*(.*)", label, _re.IGNORECASE)
    if m:
        num, title = m.group(1), m.group(2).strip()
        return (
            f'<div style="display:flex;align-items:center;gap:12px;margin:20px 0 10px;">'
            f'<div style="width:30px;height:30px;border-radius:50%;flex-shrink:0;'
            f'background:linear-gradient(135deg,#6366f1,#8b5cf6);'
            f'display:flex;align-items:center;justify-content:center;'
            f'color:#fff;font-size:0.78em;font-weight:800;'
            f'box-shadow:0 3px 10px rgba(99,102,241,0.4);">{num}</div>'
            f'<span style="font-size:0.78em;font-weight:800;text-transform:uppercase;'
            f'letter-spacing:0.10em;color:#6366f1;">{title}</span>'
            f'<div style="flex:1;height:1px;background:linear-gradient(90deg,rgba(99,102,241,0.25),transparent);"></div>'
            f'</div>'
        )
    return (
        f'<div style="display:flex;align-items:center;gap:10px;margin:16px 0 8px;">'
        f'<div style="width:6px;height:6px;border-radius:50%;background:#6366f1;flex-shrink:0;"></div>'
        f'<span style="font-size:0.78em;font-weight:800;text-transform:uppercase;'
        f'letter-spacing:0.10em;color:#6366f1;">{label}</span>'
        f'<div style="flex:1;height:1px;background:linear-gradient(90deg,rgba(99,102,241,0.25),transparent);"></div>'
        f'</div>'
    )

# ── UI ──────────────────────────────────────────────────────────────────────────

with gr.Blocks(title="Transcript Agent") as demo:

    gr.HTML(_THEME_TOGGLE)
    gr.HTML(_HERO)
    gr.HTML(_API_BANNER)

    # ── Wake lock notice (shown while transcription runs) ────────────────────
    gr.HTML(
        '<div id="ta-wake-notice" style="display:none;background:#1e3a5f;border:1px solid #3b82f6;'
        'border-radius:8px;padding:8px 14px;margin-bottom:8px;align-items:center;gap:8px;font-size:0.85em;">'
        '<span>☕</span>'
        '<span style="color:#93c5fd"><strong>Screen kept awake</strong> — your computer will not sleep during transcription. '
        'Safe to step away.</span></div>'
    )

    # ── Version / update banner (always visible) ─────────────────────────────
    _init_ub = _get_update_banner()
    update_banner = gr.HTML(value=_init_ub, visible=bool(_init_ub.strip()))
    with gr.Row(visible=False) as update_row:
        update_btn    = gr.Button("Download & Install Update", variant="primary", size="sm")
        update_status = gr.Markdown(value="", visible=True)

    # ── Job status banner (updates on page load) ──────────────────────────────
    _init_jb = get_job_banner()
    job_banner = gr.HTML(value=_init_jb, visible=bool(_init_jb.strip()))
    with gr.Row():
        load_last_btn = gr.Button("📂 Load Last Result", size="sm", variant="secondary")
        load_last_msg = gr.Markdown(visible=False)

    provider_dropdown = gr.Dropdown(
        label="AI Provider",
        choices=list(_PROVIDERS.keys()),
        value="Claude (Anthropic)",
        elem_id="provider-sel",
    )
    model_dropdown = gr.Dropdown(
        label="Model",
        choices=_PROVIDERS["Claude (Anthropic)"]["models"],
        value=_PROVIDERS["Claude (Anthropic)"]["models"][0],
        allow_custom_value=True,
        elem_id="model-sel",
    )

    user_api_key = gr.Textbox(
        label="Claude (Anthropic) API Key",
        placeholder="sk-ant-api03-…",
        type="password",
        info="console.anthropic.com → API keys → Create key",
    )
    custom_base_url = gr.Textbox(
        label="Custom API Base URL",
        placeholder="http://localhost:1234/v1",
        info="e.g. LM Studio, vLLM, Azure OpenAI, or any OpenAI-compatible endpoint",
        visible=False,
    )

    with gr.Row(equal_height=False):

        # ── left sidebar ──────────────────────────────────────────────────────
        with gr.Column(scale=1, min_width=320):

            gr.HTML(_SECTION("Step 1 — Upload"))
            file_input = gr.File(
                label="Drag & drop a file or click to browse",
                file_types=SUPPORTED,
                type="filepath",
            )
            path_input = gr.Textbox(
                label="Or paste a file path or URL (large files / no upload wait)",
                placeholder='e.g.  C:\\Videos\\interview.mp4  or  https://example.com/recording.webm',
            )
            gr.HTML(_FORMATS)

            gr.HTML(_SECTION("Step 2 — Configure"))
            with gr.Accordion("Transcription Engine", open=True):
                stt_radio = gr.Dropdown(
                    choices=list(_STT_PROVIDERS.keys()),
                    value="Whisper (Local)",
                    label="Engine",
                    info="Whisper runs locally (free, private). Cloud providers are faster and support more languages.",
                    allow_custom_value=False,
                )
                # Whisper-only options
                whisper_input = gr.Dropdown(
                    label="Whisper model",
                    choices=_STT_PROVIDERS["Whisper (Local)"]["models"],
                    value="base",
                    info="tiny = fastest   |   large = most accurate",
                    visible=True,
                )
                # Cloud STT options (hidden when Whisper selected; label/choices update dynamically)
                stt_cloud_key_input = gr.Textbox(
                    label="Deepgram API Key",
                    placeholder="dg-…",
                    type="password",
                    info="console.deepgram.com → Create API Key",
                    visible=False,
                )
                stt_cloud_model_input = gr.Dropdown(
                    label="Model",
                    choices=_STT_PROVIDERS["Deepgram (Cloud)"]["models"],
                    value="nova-2",
                    info="nova-2 = best accuracy   |   nova = fast & accurate   |   enhanced/base = cheaper",
                    visible=False,
                )

            with gr.Accordion("Processing Options", open=True):
                _COMMON_TZ = [
                    "America/New_York", "America/Chicago", "America/Denver",
                    "America/Los_Angeles", "America/Anchorage", "Pacific/Honolulu",
                    "America/Phoenix", "America/Toronto", "America/Vancouver",
                    "Europe/London", "Europe/Paris", "Europe/Berlin", "Europe/Moscow",
                    "Asia/Tokyo", "Asia/Shanghai", "Asia/Kolkata", "Asia/Dubai",
                    "Australia/Sydney", "Pacific/Auckland",
                ]
                _ALL_TZ = _COMMON_TZ + sorted(
                    z for z in zoneinfo.available_timezones() if z not in _COMMON_TZ
                )
                tz_input = gr.Dropdown(
                    label="Timezone",
                    choices=_ALL_TZ,
                    value="",
                    allow_custom_value=True,
                    info="Controls the 'Done By' finish time. Auto-filled on load; change to any IANA timezone.",
                )
                speakers_name_input = gr.Textbox(visible=False, value="")
                speakers_count_input = gr.Number(
                    label="Number of speakers (optional)",
                    value=None,
                    step=1,
                    info="How many people are speaking? AI will label them Speaker 1, Speaker 2, etc.",
                )
                analysis_depth = gr.Dropdown(
                    label="AI analysis depth",
                    choices=[
                        ("Balanced — full report (recommended)", "balanced"),
                        ("Fast — quick summary, fewer details", "fast"),
                        ("Deep — extended analysis, more detail", "deep"),
                    ],
                    value="balanced",
                    info="Fast: 4K tokens, quick. Balanced: 16K, standard. Deep: 24K + extended thinking (higher cost).",
                )
                # panel_toggle kept as hidden dummy so existing wiring doesn't break
                panel_toggle = gr.Checkbox(value=False, visible=False)

            with gr.Accordion("Language", open=False):
                language_input = gr.Dropdown(
                    label="Transcript language",
                    choices=LANGUAGES,
                    value="auto",
                )
                # Pre-load every possible variant value so Gradio never rejects
                # a value update because the new value isn't in the current choices.
                _all_variant_choices = [
                    (label, val)
                    for variants in LANGUAGE_VARIANTS.values()
                    for label, val in variants
                ]
                language_variant = gr.Dropdown(
                    label="Regional variant / dialect",
                    choices=_all_variant_choices,
                    value=None,
                    visible=False,
                )

            with gr.Accordion("Report Format", open=False):
                report_style = gr.Dropdown(
                    label="Writing style",
                    choices=[
                        ("Formal — professional, structured",  "formal"),
                        ("Casual — conversational, friendly",  "casual"),
                        ("Executive brief — ultra concise",    "executive"),
                        ("Bullet-heavy — minimal prose",       "bullet"),
                    ],
                    value="formal",
                )
                gr.HTML("""
<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;
     letter-spacing:0.1em;color:var(--ta-card-sub);margin:8px 0 4px;">
  Include in report
</div>
<div style="background:var(--ta-card-bg);border:1px solid var(--ta-card-border);
     border-radius:8px;padding:10px 14px;font-size:0.88em;">
  <div style="color:var(--ta-card-sub);font-size:0.78em;margin-bottom:6px;">
    Tick what to include — untick to skip
  </div>
</div>""")
                with gr.Row():
                    with gr.Column(min_width=130):
                        inc_summary    = gr.Checkbox(label="Summary",          value=True)
                        inc_key_points = gr.Checkbox(label="Key points",       value=True)
                        inc_action     = gr.Checkbox(label="Action items",     value=True)
                    with gr.Column(min_width=130):
                        inc_transcript = gr.Checkbox(label="Full transcript",  value=True)
                        inc_profiles   = gr.Checkbox(label="Speaker profiles", value=True)
                        inc_analytics  = gr.Checkbox(label="Speech analytics", value=True)
                with gr.Row():
                    inc_interview = gr.Checkbox(
                        label="🎤 Interview mode — extract questions, ideal answers & score responses",
                        value=False,
                    )

                with gr.Column(visible=False) as interview_advanced_col:
                    gr.HTML('<div style="font-size:0.75em;color:var(--ta-card-sub);margin:4px 0 2px;">Advanced — optional</div>')
                    inc_interview_deep = gr.Checkbox(
                        label="Deep analysis — deflection detection, % likelihood of advancing & prep guide",
                        value=False,
                    )
                    resume_context_input = gr.Textbox(
                        label="Your resume / narratives (optional — personalizes ideal answers and prep guide)",
                        placeholder="Paste a summary of your background, key accomplishments, or STAR stories…",
                        lines=4,
                        max_lines=10,
                        visible=True,
                    )

                inc_interview.change(
                    fn=lambda v: gr.update(visible=v),
                    inputs=inc_interview,
                    outputs=interview_advanced_col,
                )

            gr.HTML("""
<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;
     letter-spacing:0.1em;color:var(--ta-card-sub);margin:8px 0 4px;">
  Step 3 — Run
</div>""")
            process_btn = gr.Button(
                "Analyze File",
                variant="primary", size="lg",
                elem_classes=["big-btn"],
            )

            result_state = gr.State(value=None)

            download_accordion = gr.Accordion("📥 Download Outputs", open=False)
            with download_accordion:
                gr.HTML("""
<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;
     letter-spacing:0.1em;color:var(--ta-card-sub);margin:4px 0 6px;">
  Generate in a different language
</div>""")
                with gr.Row():
                    pdf_lang_input = gr.Dropdown(
                        label="Output language",
                        choices=_PDF_LANGUAGES,
                        value="Same as source",
                        scale=3,
                        info="Translate the report before generating — applies to PDF",
                    )
                    pdf_regen_btn = gr.Button("🔄 Regenerate PDF", scale=1, size="sm")
                gr.HTML("<hr style='margin:10px 0 8px;opacity:0.25;'>")
                gr.HTML("""
<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;
     letter-spacing:0.1em;color:var(--ta-card-sub);margin:4px 0 6px;">
  📄 Report files
</div>""")
                with gr.Row():
                    dl_pdf      = gr.DownloadButton("📄 PDF",      visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-pdf"])
                    dl_docx     = gr.DownloadButton("📝 DOCX",     visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-docx"])
                    dl_report   = gr.DownloadButton("📋 Markdown", visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-md"])
                    dl_combined = gr.DownloadButton("📃 Full TXT", visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-txt"])
                gr.HTML("""
<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;
     letter-spacing:0.1em;color:var(--ta-card-sub);margin:10px 0 6px;">
  📝 Transcript files
</div>""")
                with gr.Row():
                    dl_transcript = gr.DownloadButton("📄 Transcript",  visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-txt"])
                    dl_speakers   = gr.DownloadButton("👥 Speakers",    visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-txt"])
                    dl_srt        = gr.DownloadButton("🎬 SRT",         visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-srt"])
                    dl_vtt        = gr.DownloadButton("🎬 VTT",         visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-vtt"])
                gr.HTML("""
<div style="font-size:0.7em;font-weight:700;text-transform:uppercase;
     letter-spacing:0.1em;color:var(--ta-card-sub);margin:10px 0 6px;">
  📊 Raw data
</div>""")
                with gr.Row():
                    dl_json = gr.DownloadButton("📊 JSON", visible=False, size="sm", variant="secondary", elem_classes=["dl-btn", "dl-json"])

            bw_display = gr.HTML(value=_get_bandwidth_html(), visible=_PSUTIL_OK)
            bw_timer   = gr.Timer(value=2, active=True)

        # ── results panel ─────────────────────────────────────────────────────
        with gr.Column(scale=2):

            status_bar = gr.HTML(value=_IDLE_STATUS, elem_id="ta-status-bar")
            eta_panel  = gr.HTML(value="")
            log_out    = gr.HTML(
                value='<div id="ta-log-wrap" style="background:var(--ta-log-bg,#0f172a);'
                      'border:1px solid var(--ta-log-border,#1e3a5f);'
                      'border-radius:10px;padding:12px 16px;min-height:120px;max-height:260px;'
                      'overflow-y:auto;font-family:\'Courier New\',monospace;font-size:0.80em;'
                      'line-height:1.7;">'
                      '<span style="color:var(--ta-log-ts,#64748b);">Progress appears here once you click Analyze File…</span>'
                      '</div>',
                elem_id="live-log",
                label="Live Processing Log",
            )

            with gr.Tabs():
                with gr.TabItem("Summary"):
                    summary_out = gr.Markdown(value="")

                with gr.TabItem("Transcript"):
                    transcript_out = gr.Textbox(
                        lines=28, buttons=["copy"], label="",
                        placeholder="Clean transcript will appear here…",
                    )

                with gr.TabItem("Speaker Dialogue"):
                    dialogue_out = gr.Textbox(
                        lines=28, buttons=["copy"], label="",
                        placeholder="Speaker-labelled dialogue will appear here…",
                    )

                with gr.TabItem("Profiles"):
                    profiles_out = gr.Markdown(
                        value="_Speaker profiles will appear here after processing. Enable **Speaker profiles** in Report Format settings._"
                    )

                with gr.TabItem("Interview Q&A"):
                    interview_out = gr.Markdown(
                        value="_Interview Q&A will appear here after processing. Enable **Interview mode** in Report Format settings._"
                    )

                with gr.TabItem("Speech Analytics"):
                    analytics_out = gr.Markdown(
                        value="_Speech rate and accent analysis will appear here after processing._"
                    )

                with gr.TabItem("Copy All"):
                    combined_out = gr.Textbox(
                        lines=30, buttons=["copy"], label="",
                        placeholder="All sections combined will appear here…",
                    )

                with gr.TabItem("History"):
                    history_refresh_btn = gr.Button("🔄 Refresh", size="sm", variant="secondary")
                    history_table = gr.HTML(value=_build_history_html())
                    with gr.Row():
                        history_job_id_box = gr.Textbox(
                            label="Job ID",
                            placeholder="Click 📂 Load on any row above",
                            scale=3,
                            elem_id="history-job-id-input",
                        )
                        history_load_btn = gr.Button("📂 Load", size="sm", variant="secondary", scale=1, elem_id="history-load-btn")
                    history_msg = gr.Markdown(visible=False)

    gr.HTML("""
    <div style="text-align:center;padding:20px 0 4px;font-size:0.76em;color:#94a3b8;">
      Transcript Agent &nbsp;&bull;&nbsp; Transcription by OpenAI Whisper
      &nbsp;&bull;&nbsp; Analysis by Anthropic Claude
      &nbsp;&bull;&nbsp; Files processed privately on your machine
    </div>
    """)

    # ── event wiring ──────────────────────────────────────────────────────────
    def on_provider_change(provider):
        cfg = _PROVIDERS.get(provider, _PROVIDERS["Claude (Anthropic)"])
        is_custom = provider == "Custom (OpenAI-compatible)"
        return (
            gr.update(choices=cfg["models"], value=cfg["models"][0]),
            gr.update(label=f"{provider} API Key", placeholder=cfg["placeholder"], info=cfg["info"]),
            gr.update(visible=is_custom),
        )

    provider_dropdown.change(
        fn=on_provider_change,
        inputs=[provider_dropdown],
        outputs=[model_dropdown, user_api_key, custom_base_url],
    )

    # ── Transcription Engine toggle ───────────────────────────────────────────
    # STT providers whose API key can be reused from the analysis provider at the top
    _STT_KEY_REUSE = {
        "Groq Whisper (Cloud)": "Groq",
        "OpenAI Whisper API (Cloud)": "OpenAI",
    }

    def on_stt_change(stt, current_api_key, current_provider):
        is_local = stt.startswith("Whisper")
        cfg = _STT_PROVIDERS.get(stt, {})

        # Auto-fill key when the STT provider matches the analysis provider at top
        key_kwargs = {
            "visible": not is_local,
            "label": f"{stt} API Key",
            "placeholder": cfg.get("key_placeholder", ""),
            "info": cfg.get("key_info", ""),
        }
        matched = _STT_KEY_REUSE.get(stt)
        if matched and matched == current_provider and current_api_key:
            key_kwargs["value"] = current_api_key

        return (
            gr.update(visible=is_local),
            gr.update(**key_kwargs),
            gr.update(
                visible=not is_local,
                choices=cfg.get("models", []),
                value=cfg.get("default_model"),
                info=cfg.get("model_info", ""),
            ),
        )

    stt_radio.change(
        fn=on_stt_change,
        inputs=[stt_radio, user_api_key, provider_dropdown],
        outputs=[whisper_input, stt_cloud_key_input, stt_cloud_model_input],
    )

    # panel_toggle is hidden dummy — no change handler needed
    language_input.change(
        fn=toggle_language_variant,
        inputs=language_input,
        outputs=language_variant,
    )

    process_btn.click(
        fn=process_file,
        js="(...args) => { if(window._taAcquireWakeLock) window._taAcquireWakeLock(); var el = document.getElementById('ta-status-bar'); if(el) el.scrollIntoView({behavior:'smooth', block:'start'}); return args; }",
        inputs=[
            file_input, path_input,
            panel_toggle, speakers_name_input, speakers_count_input, whisper_input,
            language_input, language_variant,
            report_style,
            inc_summary, inc_key_points, inc_action,
            inc_transcript, inc_profiles, inc_analytics,
            inc_interview, inc_interview_deep, resume_context_input,
            user_api_key,
            provider_dropdown, model_dropdown,
            custom_base_url,
            tz_input,
            analysis_depth,
            stt_radio, stt_cloud_key_input, stt_cloud_model_input,
        ],
        outputs=[
            status_bar,
            summary_out, transcript_out, dialogue_out,
            profiles_out, interview_out, analytics_out, combined_out,
            dl_transcript, dl_speakers, dl_report, dl_combined, dl_json, dl_pdf,
            dl_docx, dl_srt, dl_vtt,
            download_accordion,
            log_out,
            eta_panel,
            result_state,
        ],
    )

    pdf_regen_btn.click(
        fn=generate_pdf_in_language,
        inputs=[result_state, pdf_lang_input, user_api_key, provider_dropdown, model_dropdown],
        outputs=[dl_pdf],
    )

    # ── Load Last Result ──────────────────────────────────────────────────────
    load_last_btn.click(
        fn=load_last_result,
        inputs=[],
        outputs=[
            summary_out, transcript_out, dialogue_out, profiles_out, interview_out,
            analytics_out, combined_out, dl_transcript, dl_speakers, dl_report,
            dl_combined, dl_json, dl_pdf, dl_docx, dl_srt, dl_vtt, download_accordion, load_last_msg,
        ],
    )

    # ── History tab ───────────────────────────────────────────────────────────
    history_refresh_btn.click(
        fn=_build_history_html,
        inputs=[],
        outputs=[history_table],
    )

    history_load_btn.click(
        fn=load_job_from_history,
        inputs=[history_job_id_box],
        outputs=[
            summary_out, transcript_out, dialogue_out, profiles_out, interview_out,
            analytics_out, combined_out, dl_transcript, dl_speakers, dl_report,
            dl_combined, dl_json, dl_pdf, dl_docx, dl_srt, dl_vtt, download_accordion, history_msg,
        ],
    )

    # ── Update button ─────────────────────────────────────────────────────────
    update_btn.click(
        fn=_do_update,
        inputs=[],
        outputs=[update_status],
    ).then(fn=None)  # keep generator alive until done

    # ── Bandwidth timer ───────────────────────────────────────────────────────
    if _PSUTIL_OK:
        bw_timer.tick(fn=_get_bandwidth_html, inputs=[], outputs=[bw_display])

    # ── Refresh banners on page load ──────────────────────────────────────────
    def _on_load(browser_tz=""):
        tz_val = browser_tz if browser_tz else ""
        btn_label = "Restart to Apply Update" if _update_downloaded.is_set() else "Download & Install Update"
        jb = get_job_banner()
        ub = _get_update_banner()
        return (
            gr.update(value=jb, visible=bool(jb.strip())),
            gr.update(value=ub, visible=bool(ub.strip())),
            gr.update(visible=bool(_update_info)),
            gr.update(value=tz_val),
            gr.update(value=btn_label),
        )

    demo.load(
        fn=_on_load,
        inputs=[tz_input],
        outputs=[job_banner, update_banner, update_row, tz_input, update_btn],
        js="() => [Intl.DateTimeFormat().resolvedOptions().timeZone]",
    )


def main():
    _host     = os.environ.get("GRADIO_SERVER_NAME", "0.0.0.0")
    _port     = int(os.environ.get("GRADIO_SERVER_PORT", 7860))
    _docker   = _host == "0.0.0.0"
    _windowed = bool(os.environ.get("TRANSCRIPT_AGENT_WINDOWED"))
    demo.queue(max_size=5, default_concurrency_limit=1)
    demo.launch(
        server_name=_host,
        server_port=_port,
        js=_THEME_JS,
        theme=_THEME,
        css=CSS,
        allowed_paths=[str(OUT_DIR), tempfile.gettempdir()],
        max_file_size="4gb",
        inbrowser=not _docker and not _windowed,
        show_error=True,
        share=not _docker and not _windowed,
        strict_cors=not _docker,
    )

if __name__ == "__main__":
    main()
